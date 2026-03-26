#!/usr/bin/env python3
"""
Write It Great - Book Proposal Evaluation System
Flask application with database, admin dashboard, status tracking, and email notifications
"""

import os
import json
import uuid
import hashlib
import smtplib
import traceback
import threading
from io import BytesIO
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, send_file, session, abort
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import text
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from xhtml2pdf import pisa
import pyotp
import qrcode
import base64
import openai
from docx import Document
import PyPDF2

# ============================================================================
# FLASK APP CONFIGURATION
# ============================================================================

app = Flask(__name__)

# ── Reverse-proxy trust (required on Heroku / any platform behind a load balancer)
# Without this Flask never sees the HTTPS scheme from the browser, which breaks:
#   - url_for() generating http:// redirect targets after login
#   - Session cookies missing the Secure flag (browser silently drops them on HTTPS)
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///proposals.db')
if app.config['SQLALCHEMY_DATABASE_URI'].startswith('postgres://'):
    app.config['SQLALCHEMY_DATABASE_URI'] = app.config['SQLALCHEMY_DATABASE_URI'].replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# ── APP_BASE_URL — must be defined before cookie config (used as the HTTPS signal)
# Set APP_URL (or APP_BASE_URL) in Heroku Config Vars:
#   APP_URL = https://authors.writeitgreat.com
APP_BASE_URL = (
    os.environ.get('APP_BASE_URL') or
    os.environ.get('APP_URL') or
    'http://localhost:5000'
).rstrip('/')

# ── Production detection: trust APP_BASE_URL, not DATABASE_URL.
# If the public URL is https://, we're on a live deployment — enable secure cookies.
# This works correctly on Heroku as soon as APP_URL is set to the custom domain.
_is_production = APP_BASE_URL.startswith('https://')

# ── Session / cookie security ──────────────────────────────────────────────────
app.config['SESSION_COOKIE_SECURE']   = _is_production   # only transmit over HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True              # JS cannot read it
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'            # CSRF protection; allows normal nav
app.config['REMEMBER_COOKIE_SECURE']  = _is_production
app.config['REMEMBER_COOKIE_HTTPONLY']= True
# SESSION_COOKIE_DOMAIN is intentionally NOT set — Flask derives it from the
# incoming Host header, which means it works on any domain without code changes.

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'author_login'
login_manager.login_message = None  # Disable "Please log in" message

# Custom Jinja filter for parsing JSON strings in templates
@app.template_filter('fromjson')
def fromjson_filter(s):
    try:
        return json.loads(s) if s else []
    except (json.JSONDecodeError, TypeError):
        return []

@app.template_filter('format_prompt')
def format_prompt_filter(text):
    """Convert markdown-style homework prompt text to clean HTML."""
    import re, html as htmllib
    if not text:
        return ''
    escaped = htmllib.escape(text)
    # Bold: **text** → <strong>text</strong>
    escaped = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', escaped)
    # Split into paragraphs on double newlines
    paragraphs = re.split(r'\n{2,}', escaped)
    parts = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # Convert single newlines to <br> within a paragraph
        para = para.replace('\n', '<br>')
        parts.append(f'<p>{para}</p>')
    return ''.join(parts)

# OpenAI client
client = openai.OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))

# Email configuration
SMTP_HOST = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
SMTP_PORT = int(os.environ.get('SMTP_PORT', 587))
SMTP_USER = os.environ.get('SMTP_USER', '')
SMTP_PASSWORD = os.environ.get('SMTP_PASSWORD', '')
FROM_EMAIL = os.environ.get('FROM_EMAIL', '') or SMTP_USER
TEAM_EMAILS = (os.environ.get('TEAM_EMAIL') or os.environ.get('TEAM_EMAILS') or 'anna@writeitgreat.com').split(',')

# External API configuration (Wix integration)
API_KEY = os.environ.get('API_KEY', '')

# CORS allowed origins for the /api/submit Wix endpoint.
# APP_BASE_URL is always added so authors.writeitgreat.com works automatically.
# CORS_ORIGIN env var may contain additional comma-separated origins.
# NOTE: No Google OAuth or third-party auth is used in this codebase, so no
#       OAuth callback URL changes are required.
_extra_origins = [o.strip() for o in os.environ.get('CORS_ORIGIN', '').split(',') if o.strip()]
CORS_ORIGINS = list({
    'https://www.writeitgreat.com',
    'https://authors.writeitgreat.com',
    APP_BASE_URL,
    *_extra_origins,
} - {'http://localhost:5000'} | ({'http://localhost:5000'} if not _is_production else set()))

# Keep CORS_ORIGIN as a backwards-compat alias used by _cors_headers()
CORS_ORIGIN = ','.join(CORS_ORIGINS)

# Status options for proposals
STATUS_OPTIONS = [
    ('submitted', 'Submitted'),
    ('read', 'Read'),
    ('author_call_scheduled', 'Author Call Scheduled'),
    ('author_call_completed', 'Author Call Completed'),
    ('contract_sent', 'Contract Sent'),
    ('contract_signed', 'Contract Signed'),
    ('shopping', 'Shopping with Publishers'),
    ('publisher_interest', 'Publisher Interest'),
    ('offer_received', 'Offer Received'),
    ('deal_closed', 'Deal Closed'),
    ('declined', 'Declined'),
    ('on_hold', 'On Hold'),
]

# Friendly status labels for authors (internal → public-facing)
AUTHOR_STATUS_LABELS = {
    'submitted': 'Submitted',
    'processing': 'Being Evaluated',
    'read': 'Under Review',
    'author_call_scheduled': 'Call Scheduled',
    'author_call_completed': 'Call Completed',
    'contract_sent': 'Contract Sent',
    'contract_signed': 'Contract Signed',
    'shopping': 'Being Presented to Publishers',
    'publisher_interest': 'Publisher Interest',
    'offer_received': 'Offer Received',
    'deal_closed': 'Deal Closed',
    'declined': 'Declined',
    'on_hold': 'On Hold',
    'error': 'Being Evaluated',
}

# Statuses that trigger email to the author
AUTHOR_EMAIL_MILESTONES = {
    'author_call_scheduled', 'contract_sent', 'contract_signed',
    'shopping', 'publisher_interest', 'offer_received',
    'deal_closed', 'declined',
}

# Publisher-specific statuses (set by the publisher on shared proposals)
PUBLISHER_STATUS_OPTIONS = [
    ('new', 'New'),
    ('proposal_read', 'Proposal Read'),
    ('interested', 'Interested'),
    ('ready_to_discuss', 'Ready to Discuss Offer'),
    ('deal_sent', 'Deal Sent'),
    ('deal_signed', 'Deal Signed'),
]

PUBLISHER_STATUS_LABELS = {key: label for key, label in PUBLISHER_STATUS_OPTIONS}

# Genre options for publisher profiles
GENRE_OPTIONS = [
    'Literary Fiction', 'Commercial Fiction', 'Mystery & Thriller', 'Romance',
    'Science Fiction', 'Fantasy', 'Horror', 'Historical Fiction',
    'Young Adult', 'Middle Grade', 'Children\'s', 'Memoir',
    'Biography', 'Self-Help', 'Business', 'Health & Wellness',
    'Science & Technology', 'History', 'Politics & Current Affairs',
    'True Crime', 'Travel', 'Cookbooks & Food', 'Religion & Spirituality',
    'Poetry', 'Graphic Novels', 'Other Nonfiction',
]

# Coaching program module definitions (ordered curriculum)
COACHING_MODULES = [
    {
        'order': 1,
        'title': 'Book Concept & Hook',
        'subtitle': 'Craft a compelling 1-2 sentence pitch',
        'description': (
            "Every great book starts with a hook — a clear, compelling statement that tells "
            "an editor exactly what the book is, who it's for, and why it matters now.\n\n"
            "Use the WIG framework as a starting point:\n"
            "\"Most people believe [common misconception]. But actually, [reframe]. "
            "And it's a much bigger deal than you think because [stakes/consequence].\"\n\n"
            "Adapt it freely — a strong hook sounds like something an editor would quote "
            "back to a colleague."
        ),
        'homework_prompt': (
            "Write your book hook: 1-2 sentences that describe your book, who it's for, "
            "and why it matters now. Try the WIG framework:\n"
            "\"Most people believe [misconception]. But actually, [reframe]. And it's a much "
            "bigger deal than you think because [stakes].\"\n\n"
            "Aim for clarity and specificity. A strong hook sounds like something an editor "
            "would quote back to a colleague."
        ),
        'homework_label': 'Book Hook',
        'icon': '🎯',
        'chat_context': 'developing the book concept and hook',
    },
    {
        'order': 2,
        'title': 'Target Reader',
        'subtitle': 'Define your ideal reader with precision',
        'description': (
            "Publishers need to know exactly who they're selling to. Vague audiences kill "
            "proposals. Specific, well-defined readers win deals.\n\n"
            "Not 'adults interested in business' — but 'mid-career professionals facing "
            "their first leadership role who have never managed a team before.'\n\n"
            "Describe the type of reader — their occupation, life stage, specific problem, "
            "and motivation. Do not invent a fictional named character; describe the real "
            "category of person this book is written for."
        ),
        'homework_prompt': (
            "Write a 1-2 paragraph description of your target reader. Include:\n\n"
            "• Their occupation, life stage, or professional context\n"
            "• The specific problem, frustration, or aspiration that brings them to your book\n"
            "• Why they would choose YOUR book over alternatives\n"
            "• Why this book is especially relevant for them right now\n\n"
            "Describe the type of reader — the real category of person who will buy this book. "
            "Use demographic and psychographic language, not a fictional named character."
        ),
        'homework_label': 'Target Reader Profile',
        'icon': '👤',
        'chat_context': 'defining the target reader and audience',
    },
    {
        'order': 3,
        'title': 'Comparative Analysis',
        'subtitle': 'Show where your book fits in the market',
        'description': (
            "Comp titles show publishers where your book lives in the market. "
            "You're not competing with these books — you're joining a conversation.\n\n"
            "📌 Pro Tip: Go to Amazon, search your genre, then filter by \"New Releases\" or "
            "sort by publication date to find recent titles. Check Goodreads ratings and "
            "review counts to gauge popularity. Sales figures are never required — focus on "
            "audience overlap, tone, and the gap your book fills.\n\n"
            "Consider pitching your book as \u201cTitle A meets Title B.\u201d\n\n"
            "All primary comps should be published within the last 3 years."
        ),
        'homework_prompt': (
            "List 3-5 comparable titles published in the last 3 years. "
            "For each one include:\n\n"
            "**[Book Title] by [Author] ([Year])**\n"
            "What it covers: a brief description of the book's focus.\n"
            "What it does well: what makes it resonate with readers?\n"
            "How your book is different: what specific gap does YOUR book fill?\n\n"
            "Then add a **Market Gap** paragraph: what unique angle, methodology, or "
            "perspective do you bring that none of these titles offer?\n\n"
            "Note: sales figures are NOT required and should not be included. "
            "Titles older than 3 years may be used as secondary comps — flag them as such."
        ),
        'homework_label': 'Comparative Analysis',
        'icon': '📚',
        'chat_context': 'comparative titles, market positioning, and market gap',
    },
    {
        'order': 4,
        'title': 'Author Bio',
        'subtitle': 'Why you are the right person to write this book',
        'description': (
            "Nonfiction publishers buy the author first and the book idea second. "
            "Your bio needs to establish credibility, humanity, and a clear "
            "throughline to your book's topic. Aim for 200-300 words (one page maximum).\n\n"
            "This section is about your STORY and CREDENTIALS — why you are uniquely "
            "qualified to write this book. It is NOT about your platform size or social "
            "media following (those belong in the Marketing & Platform section)."
        ),
        'homework_prompt': (
            "Write your Author Bio (200-300 words, one page maximum). Include:\n\n"
            "• Your professional credentials and expertise directly relevant to this topic\n"
            "• Relevant life experiences or turning points that inform the book\n"
            "• Previous publications, media appearances, or speaking credentials\n"
            "• The specific perspective or methodology only you can bring\n\n"
            "Keep this focused on WHO YOU ARE and WHY YOU — not your follower counts "
            "or platform statistics (save those for the Marketing & Platform section).\n\n"
            "Suggested structure:\n"
            "Paragraph 1: [Your Name] is a [role/expertise] who [relevant experience]. "
            "Through [your work], you have [demonstrated expertise or insight].\n"
            "Paragraph 2: [Media appearances, publications, speaking, or notable credentials].\n"
            "Paragraph 3: Your approach to [topic] comes from [unique perspective], making "
            "you the right person to write this book for [audience].\n\n"
            "Write in third person."
        ),
        'homework_label': 'Author Bio',
        'icon': '🏆',
        'chat_context': 'author bio, credentials, and unique positioning',
    },
    {
        'order': 5,
        'title': 'Book Outline',
        'subtitle': "Map your book's narrative arc chapter by chapter",
        'description': (
            "A strong chapter outline shows publishers you've thought through how the book "
            "actually works — not just what it's about, but how it delivers on its promise "
            "from beginning to end. Structure is argument.\n\n"
            "Each chapter summary should be 25-100 words. Show the logical flow from one "
            "chapter to the next — the structure should tell a story with momentum."
        ),
        'homework_prompt': (
            "Write a chapter-by-chapter outline (2+ pages). For each chapter include:\n\n"
            "**Introduction or Chapter 1: [Title]**\n"
            "Summarise what this chapter covers. What story will you tell? What key concepts "
            "will you introduce? What will readers learn or feel by the end?\n\n"
            "**Chapter 2: [Title]**\n"
            "Continue with chapter summaries (25-100 words each). Show the logical flow "
            "from one chapter to the next.\n\n"
            "Continue through all chapters. Include an introduction and conclusion. "
            "Aim for at least 8-10 chapters unless your structure calls for fewer."
        ),
        'homework_label': 'Book Outline',
        'icon': '📋',
        'chat_context': 'book outline and chapter structure',
    },
    {
        'order': 6,
        'title': 'Sample Writing',
        'subtitle': 'Show editors your voice in action',
        'description': (
            "Your writing sample is often the last thing editors read — but it can override "
            "everything else. Weak writing kills strong proposals. Strong writing rescues "
            "weak ones. This section gives you no room to hide.\n\n"
            "Typically include Chapter 1 plus 1-2 others that showcase different aspects "
            "of the book."
        ),
        'homework_prompt': (
            "Submit 3,000-10,000 words of your strongest writing for this book. "
            "Include 1-3 sample chapters — typically Chapter 1 plus one or two others "
            "that showcase different aspects of the book. Choose writing that immediately "
            "draws the reader in and reflects the tone of the full book."
        ),
        'homework_label': 'Sample Writing',
        'icon': '✍️',
        'chat_context': 'sample writing and authorial voice',
    },
    {
        'order': 7,
        'title': 'Marketing & Platform',
        'subtitle': "Show publishers you can help sell this book",
        'description': (
            "Publishers want to know there will be a return on their investment. "
            "Breakeven for most publishers is around 3,000 copies sold — 10,000 means "
            "they've made their money back and then some.\n\n"
            "This section is distinct from your Author Bio. Here, publishers want NUMBERS "
            "and a concrete marketing strategy — not credentials or story.\n\n"
            "Three sections to cover:\n"
            "1. Your Current Platform — real numbers (email list, social, website, speaking)\n"
            "2. Marketing Opportunities — media, speaking, bulk sales, endorsements\n"
            "3. Your Marketing Commitment — what YOU will personally do to sell this book\n\n"
            "No huge platform yet? Lead with engagement quality and growth trajectory."
        ),
        'homework_prompt': (
            "Write your Marketing & Platform section (2+ pages). Use these three headings:\n\n"
            "**Your Current Platform**\n"
            "Email list: ___ subscribers\n"
            "Instagram: ___ followers\n"
            "TikTok: ___ followers\n"
            "LinkedIn: ___ connections/followers\n"
            "Facebook: ___ followers\n"
            "YouTube / Podcast: ___ subscribers / ___ monthly listeners\n"
            "Website: ___ monthly visitors\n"
            "Speaking: ___ events per year, typical audience size ___\n"
            "(No huge platform yet? Note your engagement rate and growth trajectory.)\n\n"
            "**Marketing Opportunities**\n"
            "Personal network: who in your network can help promote this? "
            "Name specific influencers, professional associations, potential foreword "
            "writers or blurb providers.\n"
            "Media: specific podcasts, publications, or outlets that are natural fits.\n"
            "Speaking & events: named conferences, workshops, or events you could speak at.\n"
            "Corporate / bulk sales: organisations or companies that might buy in bulk.\n\n"
            "**Your Marketing Commitment**\n"
            "Describe specifically what YOU will do to market this book — concrete actions, "
            "not vague promises. Podcast tour, social media launch, speaking engagements, "
            "newsletter campaign — be specific and enthusiastic. Publishers invest in "
            "authors who invest in their own books."
        ),
        'homework_label': 'Marketing & Platform',
        'icon': '📣',
        'chat_context': 'marketing platform, audience reach, and marketing commitment',
    },
]

# Per-module knowledge base: tips, examples, and templates shown in the
# collapsible drawer on each module page.
COACHING_MODULE_RESOURCES = {
    1: {
        'tips': [
            'A great hook survives the "so what?" test — read it aloud to someone unfamiliar with your topic.',
            'Specificity beats cleverness. "Habits compound" is weaker than "1% better every day adds up to 37x improvement in a year."',
            'The best hooks name both the reader AND the payoff in one breath.',
        ],
        'examples': [
            '**Atomic Habits**: "Tiny changes, remarkable results — the proven system for building good habits and breaking bad ones."',
            '**The Body Keeps the Score**: "Trauma reshapes the brain and body. Here is the science — and the path to healing."',
            '**Dare to Lead**: "Courage is a skill, not a trait. Here is how to build it in yourself and your team."',
        ],
        'templates': [
            'WIG Framework: "Most people believe [misconception]. But actually, [reframe]. And it\'s a much bigger deal than you think because [stakes]."',
            'Simple format: "[This book] helps [specific reader] [achieve outcome] by [unique method/insight]."',
        ],
    },
    2: {
        'tips': [
            'The more specific the reader, the more an editor trusts you understand the market.',
            'Ask yourself: "What was happening in my reader\'s life the week before they picked up this book?"',
            'Avoid "anyone who…" — if your reader is everyone, your reader is no one.',
        ],
        'examples': [
            '**Narrow example**: "Mid-career professionals in their first leadership role who were never trained to manage people — they\'re good at their craft but overwhelmed by the human side of the job."',
            '**Motivation example**: "Parents of children with ADHD who are exhausted by school conflict and desperate for a framework that works at home, not just in a clinic."',
        ],
        'templates': [
            'Reader template: "[Job/life stage] who [specific problem/frustration] and want [desired outcome] — especially [timely reason this matters now]."',
        ],
    },
    3: {
        'tips': [
            'Go to Amazon → Books → your genre → filter "New Releases" to find titles from the last 3 years.',
            'Check Goodreads review counts and star ratings to understand which comps actually connected with readers.',
            'The "meets" format is a shortcut: "Think [Title A] meets [Title B] — but for [specific audience]."',
            'Sales figures are never required. Publishers know their own market.',
        ],
        'examples': [
            '**Good comp entry**: "**Never Split the Difference** by Chris Voss (2016) — covers negotiation psychology for high-stakes situations. My book applies the same principles specifically to salary and freelance contract conversations."',
            '**Market gap**: "No current title addresses this for the freelance economy specifically — the existing books assume a corporate context."',
        ],
        'templates': [
            'Comp entry: "**[Title]** by [Author] ([Year]) — [1 sentence: what it covers and why it sells]. My book differs because [1 sentence: specific gap you fill]."',
        ],
    },
    4: {
        'tips': [
            'Write in third person — "Jane Smith is a…" not "I am a…"',
            'Credentials relevant to the TOPIC matter most. A surgeon writing about leadership needs different credentials than a surgeon writing about surgery.',
            'Personal story > job title. The "why I had to write this book" moment is often the strongest credential of all.',
            'Keep it to 200-300 words — editors skim long bios.',
        ],
        'examples': [
            '**Opening example**: "Sarah Chen is a cognitive neuroscientist who spent 12 years studying decision fatigue at Stanford before burning out at 34. That experience became the catalyst for *The Clarity Method*."',
            '**Credentials example**: "Featured in Harvard Business Review, The Guardian, and NPR\'s Hidden Brain. Her TEDx talk on burnout has 2.1M views."',
        ],
        'templates': [
            'Bio template: "[Name] is a [role/expertise] who [key experience relevant to topic]. Through [work/platform], [what they\'ve demonstrated or achieved]. [Media/publications]. [Why uniquely positioned to write this book]."',
        ],
    },
    5: {
        'tips': [
            'Your chapter structure IS your argument. Read the titles alone — does the arc make sense?',
            'Each chapter summary should answer: what does the reader know/feel at the END of this chapter that they didn\'t at the start?',
            'Aim for 25-75 words per chapter summary. Vague summaries signal underdeveloped thinking.',
            'Show momentum: each chapter should make the next one feel inevitable.',
        ],
        'examples': [
            '**Strong chapter summary**: "Chapter 3: The Permission Trap. Most people wait for external validation before pursuing their goals. This chapter exposes the neuroscience of approval-seeking, then introduces the Permission Audit — a 20-minute exercise that breaks the pattern for good."',
        ],
        'templates': [
            'Chapter entry: "**Chapter [N]: [Title]**. [What problem/question this chapter addresses]. [Key concept, story, or exercise]. [What reader understands/can do by the end]."',
        ],
    },
    6: {
        'tips': [
            'Your first page is everything. If it doesn\'t pull an editor in, nothing else will.',
            'Show range: if possible, include one narrative chapter and one more practical/concept-driven chapter.',
            'Voice is not style — it\'s the feeling that a specific, irreplaceable person is speaking.',
            '3,000–5,000 words is usually enough to demonstrate your range. More is rarely better.',
        ],
        'examples': [
            '**Strong opening**: Drop the reader into a scene or moment before explaining anything. Let them feel the stakes before they understand the thesis.',
        ],
        'templates': [],
    },
    7: {
        'tips': [
            'Publishers care more about list quality than social follower counts. 5,000 engaged email subscribers beats 50,000 passive Instagram followers.',
            'Be specific about WHO in your network can help — name organisations, podcasts, or influencers.',
            'Bulk sales potential (corporate training, conferences, universities) can be more compelling than individual sales reach.',
            'Growth trajectory matters: "500 subscribers, growing 20% month-on-month" is a strong signal.',
        ],
        'examples': [
            '**Platform example**: "Email list: 4,200 subscribers (38% open rate). LinkedIn: 12,000 followers. Speaks at 6-8 HR conferences per year, avg audience 300. Podcast guest: appeared on WorkLife with Adam Grant, Dare to Lead, and HBR IdeaCast."',
            '**Marketing commitment example**: "Will launch a 6-week podcast tour (15+ shows confirmed), run a pre-order campaign to her list, and pitch a 3-part series to Harvard Business Review online."',
        ],
        'templates': [],
    },
}


# ============================================================================
# DATABASE MODELS
# ============================================================================

ROLE_ADMIN = 'admin'
ROLE_MEMBER = 'member'
ROLE_CHOICES = [(ROLE_ADMIN, 'Admin'), (ROLE_MEMBER, 'Member')]


class AdminUser(UserMixin, db.Model):
    """Admin user model for secure dashboard access"""
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    role = db.Column(db.String(20), default=ROLE_MEMBER)
    is_active_account = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    password_reset_token = db.Column(db.String(100))
    password_reset_expires = db.Column(db.DateTime)

    # TOTP 2FA
    totp_secret = db.Column(db.String(64))
    totp_enabled = db.Column(db.Boolean, default=False)

    # Login security
    failed_login_attempts = db.Column(db.Integer, default=0)
    locked_until = db.Column(db.DateTime)

    @property
    def is_admin(self):
        return self.role == ROLE_ADMIN

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def is_locked(self):
        if self.locked_until and datetime.utcnow() < self.locked_until:
            return True
        return False

    def record_failed_login(self):
        self.failed_login_attempts = (self.failed_login_attempts or 0) + 1
        if self.failed_login_attempts >= 5:
            self.locked_until = datetime.utcnow() + timedelta(minutes=15)

    def record_successful_login(self):
        self.failed_login_attempts = 0
        self.locked_until = None

    def generate_reset_token(self):
        self.password_reset_token = uuid.uuid4().hex
        self.password_reset_expires = datetime.utcnow() + timedelta(hours=1)
        return self.password_reset_token

    def verify_reset_token(self, token):
        if not self.password_reset_token or not self.password_reset_expires:
            return False
        if self.password_reset_token != token:
            return False
        if datetime.utcnow() > self.password_reset_expires:
            return False
        return True

    def setup_totp(self):
        self.totp_secret = pyotp.random_base32()
        return self.totp_secret

    def get_totp_uri(self):
        return pyotp.totp.TOTP(self.totp_secret).provisioning_uri(
            name=self.email, issuer_name='Write It Great')

    # Markers to distinguish user types in user loader
    is_author = False
    is_team_member = True
    is_publisher = False

    def verify_totp(self, code):
        if not self.totp_secret:
            return False
        # Strip any non-digit characters from input
        code = ''.join(c for c in str(code) if c.isdigit())
        if len(code) != 6:
            return False
        totp = pyotp.TOTP(self.totp_secret)
        return totp.verify(code, valid_window=2)


class Author(UserMixin, db.Model):
    """Author account for the public portal"""
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(200), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    password_reset_token = db.Column(db.String(100))
    password_reset_expires = db.Column(db.DateTime)

    # Markers to distinguish user types in user loader
    is_author = True
    is_team_member = False
    is_publisher = False

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def generate_reset_token(self):
        self.password_reset_token = uuid.uuid4().hex
        self.password_reset_expires = datetime.utcnow() + timedelta(hours=1)
        return self.password_reset_token

    def verify_reset_token(self, token):
        if not self.password_reset_token or not self.password_reset_expires:
            return False
        if self.password_reset_token != token:
            return False
        if datetime.utcnow() > self.password_reset_expires:
            return False
        return True

    # Admin-created account fields
    pending_setup = db.Column(db.Boolean, default=False)   # True until they set their own password
    admin_created = db.Column(db.Boolean, default=False)   # Created by admin on their behalf
    assigned_path = db.Column(db.String(30))               # 'full_proposal' | 'one_pager' | None
    last_login_at = db.Column(db.DateTime)

    # These properties keep templates simple
    @property
    def is_admin(self):
        return False

    proposals = db.relationship('Proposal', backref='author', lazy='dynamic')


class Proposal(db.Model):
    """Book proposal submission model"""
    id = db.Column(db.Integer, primary_key=True)
    submission_id = db.Column(db.String(50), unique=True, nullable=False)

    # Author info
    author_id = db.Column(db.Integer, db.ForeignKey('author.id'))
    author_name = db.Column(db.String(200), nullable=False)
    author_email = db.Column(db.String(200), nullable=False)
    book_title = db.Column(db.String(500), nullable=False)
    
    # Submission details
    proposal_type = db.Column(db.String(50), default='full')
    ownership_confirmed = db.Column(db.Boolean, default=True)
    content_hash = db.Column(db.String(64), index=True)  # SHA-256 of proposal text + type
    
    # Evaluation results
    tier = db.Column(db.String(10))
    overall_score = db.Column(db.Float)
    evaluation_json = db.Column(db.Text)
    proposal_text = db.Column(db.Text)

    # Original uploaded file
    original_filename = db.Column(db.String(500))
    original_file = db.Column(db.LargeBinary)
    
    # Status tracking
    status = db.Column(db.String(50), default='submitted')
    is_archived = db.Column(db.Boolean, default=False)
    notes = db.Column(db.Text)
    
    # Timestamps
    submitted_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Email tracking
    author_email_sent = db.Column(db.Boolean, default=False)
    team_email_sent = db.Column(db.Boolean, default=False)

    # Structured platform data from Update 3 form (JSON string)
    platform_data = db.Column(db.Text)
    marketing_strategy = db.Column(db.Text)


class ProposalNote(db.Model):
    """Activity log entry for a proposal — notes, status changes, etc."""
    id = db.Column(db.Integer, primary_key=True)
    proposal_id = db.Column(db.Integer, db.ForeignKey('proposal.id'), nullable=False)
    user_name = db.Column(db.String(200))
    action = db.Column(db.String(50))  # 'note', 'status_change', 'created', 'email_sent'
    old_value = db.Column(db.String(100))
    new_value = db.Column(db.String(100))
    content = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    proposal = db.relationship('Proposal', backref=db.backref('activity_log', lazy='dynamic', order_by='ProposalNote.created_at.desc()'))


class Publisher(UserMixin, db.Model):
    """External publisher/editor account"""
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(200), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    company = db.Column(db.String(200))
    is_approved = db.Column(db.Boolean, default=False)
    is_active_account = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    password_reset_token = db.Column(db.String(100))
    password_reset_expires = db.Column(db.DateTime)

    # Profile fields (visible to admin team, not authors)
    bio = db.Column(db.Text)  # Short professional bio
    preferred_genres = db.Column(db.Text)  # JSON list of genres
    preferred_topics = db.Column(db.Text)  # Free-text topics/interests
    website = db.Column(db.String(300))

    is_author = False
    is_team_member = False
    is_publisher = True

    shared_proposals = db.relationship('PublisherProposal', backref='publisher', lazy='dynamic')

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def generate_reset_token(self):
        self.password_reset_token = uuid.uuid4().hex
        self.password_reset_expires = datetime.utcnow() + timedelta(hours=1)
        return self.password_reset_token

    def verify_reset_token(self, token):
        if not self.password_reset_token or not self.password_reset_expires:
            return False
        if self.password_reset_token != token:
            return False
        if datetime.utcnow() > self.password_reset_expires:
            return False
        return True

    @property
    def is_admin(self):
        return False


class PublisherProposal(db.Model):
    """Tracks which proposals are shared with which publishers"""
    id = db.Column(db.Integer, primary_key=True)
    publisher_id = db.Column(db.Integer, db.ForeignKey('publisher.id'), nullable=False)
    proposal_id = db.Column(db.Integer, db.ForeignKey('proposal.id'), nullable=False)
    shared_at = db.Column(db.DateTime, default=datetime.utcnow)
    shared_by = db.Column(db.String(200))  # Name of admin who shared
    publisher_status = db.Column(db.String(50), default='new')  # Publisher's status on this proposal
    status_updated_at = db.Column(db.DateTime)  # When publisher last changed status

    proposal = db.relationship('Proposal', backref=db.backref('shared_with', lazy='dynamic'))

    __table_args__ = (db.UniqueConstraint('publisher_id', 'proposal_id'),)


# ── Coaching Platform Models ─────────────────────────────────────────────────

class CoachingEnrollment(db.Model):
    """Tracks an author's enrollment in the coaching program"""
    __tablename__ = 'coaching_enrollment'
    id = db.Column(db.Integer, primary_key=True)
    author_id = db.Column(db.Integer, db.ForeignKey('author.id'), nullable=False)
    book_title = db.Column(db.String(500))
    enrolled_at = db.Column(db.DateTime, default=datetime.utcnow)
    status = db.Column(db.String(20), default='active')  # active / completed / paused
    completed_at = db.Column(db.DateTime)
    current_module = db.Column(db.Integer, default=1)
    welcome_email_sent = db.Column(db.Boolean, default=False)
    complete_email_sent = db.Column(db.Boolean, default=False)

    author = db.relationship('Author', backref=db.backref('coaching_enrollments', lazy='dynamic'))
    module_progress = db.relationship('AuthorModuleProgress', backref='enrollment',
                                      lazy='dynamic', order_by='AuthorModuleProgress.module_order')
    chat_messages = db.relationship('CoachingChatMessage', backref='enrollment',
                                    lazy='dynamic', order_by='CoachingChatMessage.created_at')
    homework_submissions = db.relationship('HomeworkSubmission', backref='enrollment', lazy='dynamic')


class AuthorModuleProgress(db.Model):
    """Tracks an author's status for each coaching module"""
    __tablename__ = 'author_module_progress'
    id = db.Column(db.Integer, primary_key=True)
    enrollment_id = db.Column(db.Integer, db.ForeignKey('coaching_enrollment.id'), nullable=False)
    module_order = db.Column(db.Integer, nullable=False)  # 1-7
    status = db.Column(db.String(30), default='locked')   # locked / in_progress / revision_requested / approved
    unlocked_at = db.Column(db.DateTime)
    completed_at = db.Column(db.DateTime)
    admin_notes = db.Column(db.Text)
    module_unlock_email_sent = db.Column(db.Boolean, default=False)
    homework_reminder_sent_at = db.Column(db.DateTime)


class CoachingChatMessage(db.Model):
    """Server-side storage of module coaching chat messages"""
    __tablename__ = 'coaching_chat_message'
    id = db.Column(db.Integer, primary_key=True)
    enrollment_id = db.Column(db.Integer, db.ForeignKey('coaching_enrollment.id'), nullable=False)
    module_order = db.Column(db.Integer, nullable=False)
    role = db.Column(db.String(20), nullable=False)  # 'user' or 'assistant'
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class HomeworkSubmission(db.Model):
    """Author homework submission for a coaching module"""
    __tablename__ = 'homework_submission'
    id = db.Column(db.Integer, primary_key=True)
    enrollment_id = db.Column(db.Integer, db.ForeignKey('coaching_enrollment.id'), nullable=False)
    module_order = db.Column(db.Integer, nullable=False)
    content = db.Column(db.Text, nullable=False)
    submitted_at = db.Column(db.DateTime, default=datetime.utcnow)
    revision_number = db.Column(db.Integer, default=1)
    # AI review
    ai_feedback = db.Column(db.Text)
    ai_approved = db.Column(db.Boolean)
    ai_reviewed_at = db.Column(db.DateTime)
    # Admin review
    admin_feedback = db.Column(db.Text)
    admin_reviewed_by = db.Column(db.String(200))
    admin_reviewed_at = db.Column(db.DateTime)
    status = db.Column(db.String(30), default='pending_review')  # pending_review / revision_requested / approved
    review_email_sent = db.Column(db.Boolean, default=False)


class CoachingModuleContent(db.Model):
    """Autosaved draft content for each coaching module — the living proposal sections"""
    __tablename__ = 'coaching_module_content'
    id = db.Column(db.Integer, primary_key=True)
    enrollment_id = db.Column(db.Integer, db.ForeignKey('coaching_enrollment.id'), nullable=False)
    module_order = db.Column(db.Integer, nullable=False)
    content = db.Column(db.Text)
    word_count = db.Column(db.Integer, default=0)
    last_saved_at = db.Column(db.DateTime, default=datetime.utcnow)

    __table_args__ = (db.UniqueConstraint('enrollment_id', 'module_order'),)


class OnePagerSubmission(db.Model):
    """Stores quick one-pager submissions (answers + AI-generated summary)"""
    __tablename__ = 'one_pager_submission'
    id = db.Column(db.Integer, primary_key=True)
    author_id = db.Column(db.Integer, db.ForeignKey('author.id'), nullable=False)
    book_title = db.Column(db.String(500))
    answers_json = db.Column(db.Text)          # JSON: {problem, reader, different, why_you, marketing}
    summary_text = db.Column(db.Text)          # AI-generated summary
    status = db.Column(db.String(20), default='draft')   # 'draft' | 'submitted'
    submitted_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    admin_notes = db.Column(db.Text)

    author = db.relationship('Author', backref=db.backref('one_pager_submissions', lazy='dynamic'))


class KnowledgeBaseDocument(db.Model):
    """Admin-uploaded training/reference documents per module"""
    __tablename__ = 'knowledge_base_document'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    filename = db.Column(db.String(300))
    content_text = db.Column(db.Text)          # Extracted plain text for AI context
    file_data = db.Column(db.LargeBinary)      # Raw file bytes
    file_type = db.Column(db.String(10))       # 'pdf', 'docx', 'txt'
    module_order = db.Column(db.Integer)       # NULL = applies to all modules
    doc_type = db.Column(db.String(30), default='resource')  # 'example' | 'template' | 'resource'
    uploaded_by = db.Column(db.String(200))
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)


class AuthorEngagementEmail(db.Model):
    """Tracks automated re-engagement emails sent to authors (hard cap: 3 total)"""
    __tablename__ = 'author_engagement_email'
    id = db.Column(db.Integer, primary_key=True)
    author_id = db.Column(db.Integer, db.ForeignKey('author.id'), nullable=False)
    email_type = db.Column(db.String(60))      # see REENGAGEMENT_TYPES
    sent_at = db.Column(db.DateTime, default=datetime.utcnow)

    author = db.relationship('Author', backref=db.backref('engagement_emails', lazy='dynamic'))


@login_manager.user_loader
def load_user(user_id):
    """Load user from session — checks user_type to pick the right model"""
    from flask import session as flask_session
    user_type = flask_session.get('user_type', 'admin')
    if user_type == 'author':
        return Author.query.get(int(user_id))
    elif user_type == 'publisher':
        pub = Publisher.query.get(int(user_id))
        if pub and (not pub.is_active_account or not pub.is_approved):
            return None
        return pub
    else:
        user = AdminUser.query.get(int(user_id))
        if user and not user.is_active_account:
            return None
        return user


from functools import wraps
from urllib.parse import urlparse

def _safe_next(next_url):
    """Return next_url only if it is a relative path on this app.
    Rejects absolute URLs (open-redirect prevention) and any URL pointing
    to a different host — including stale references to the old Heroku domain."""
    if not next_url:
        return None
    parsed = urlparse(next_url)
    # Accept only paths with no scheme/netloc (i.e. relative URLs)
    if parsed.scheme or parsed.netloc:
        return None
    return next_url


def admin_required(f):
    """Decorator: requires login + admin role"""
    @wraps(f)
    @login_required
    def decorated(*args, **kwargs):
        if not current_user.is_admin:
            flash('You need admin privileges to access this page.', 'error')
            return redirect(url_for('admin_dashboard'))
        return f(*args, **kwargs)
    return decorated


def team_required(f):
    """Decorator: requires login as a team member (AdminUser)"""
    @wraps(f)
    @login_required
    def decorated(*args, **kwargs):
        if not getattr(current_user, 'is_team_member', False):
            flash('Team access required.', 'error')
            return redirect(url_for('author_dashboard'))
        return f(*args, **kwargs)
    return decorated


def author_login_required(f):
    """Decorator: requires login as an author"""
    @wraps(f)
    @login_required
    def decorated(*args, **kwargs):
        if not getattr(current_user, 'is_author', False):
            return redirect(url_for('author_login'))
        return f(*args, **kwargs)
    return decorated


def publisher_login_required(f):
    """Decorator: requires login as an approved publisher"""
    @wraps(f)
    @login_required
    def decorated(*args, **kwargs):
        if not getattr(current_user, 'is_publisher', False):
            return redirect(url_for('publisher_login'))
        return f(*args, **kwargs)
    return decorated


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def generate_submission_id():
    """Generate unique submission ID"""
    date_str = datetime.now().strftime('%Y%m%d')
    random_str = uuid.uuid4().hex[:5].upper()
    return f"WIG-{date_str}-{random_str}"


def extract_text_from_pdf(file):
    """Extract text from PDF file"""
    try:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def convert_docx_to_html(file_bytes):
    """Convert DOCX bytes to formatted HTML preserving headings, bold, italic, lists"""
    try:
        doc = Document(BytesIO(file_bytes))
        html_parts = []
        for para in doc.paragraphs:
            style_name = (para.style.name or '').lower()
            # Build inline runs with bold/italic
            runs_html = ''
            for run in para.runs:
                text = run.text
                if not text:
                    continue
                # Escape HTML
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if run.bold and run.italic:
                    text = f'<strong><em>{text}</em></strong>'
                elif run.bold:
                    text = f'<strong>{text}</strong>'
                elif run.italic:
                    text = f'<em>{text}</em>'
                if run.underline:
                    text = f'<u>{text}</u>'
                runs_html += text

            if not runs_html.strip():
                html_parts.append('<br>')
                continue

            if 'heading 1' in style_name:
                html_parts.append(f'<h2>{runs_html}</h2>')
            elif 'heading 2' in style_name:
                html_parts.append(f'<h3>{runs_html}</h3>')
            elif 'heading 3' in style_name or 'heading 4' in style_name:
                html_parts.append(f'<h4>{runs_html}</h4>')
            elif 'list' in style_name or 'bullet' in style_name:
                html_parts.append(f'<li>{runs_html}</li>')
            elif 'title' in style_name:
                html_parts.append(f'<h1>{runs_html}</h1>')
            elif 'subtitle' in style_name:
                html_parts.append(f'<h3 style="color: var(--gray-500);">{runs_html}</h3>')
            else:
                html_parts.append(f'<p>{runs_html}</p>')

        # Wrap consecutive <li> tags in <ul>
        result = '\n'.join(html_parts)
        import re
        result = re.sub(r'((?:<li>.*?</li>\s*)+)', r'<ul>\1</ul>', result)
        return result
    except Exception as e:
        print(f"DOCX to HTML conversion error: {e}")
        return None


# Scoring weights for full proposals
FULL_WEIGHTS = {
    'marketing': 0.30,
    'overview': 0.20,
    'credentials': 0.15,
    'comps': 0.10,
    'writing': 0.15,
    'outline': 0.05,
    'completeness': 0.05
}

MARKETING_ONLY_WEIGHTS = {
    'marketing': 1.00, 'overview': 0.00, 'credentials': 0.00,
    'comps': 0.00, 'writing': 0.00, 'outline': 0.00, 'completeness': 0.00
}

NO_MARKETING_WEIGHTS = {
    'marketing': 0.00, 'overview': 0.29, 'credentials': 0.21,
    'comps': 0.14, 'writing': 0.21, 'outline': 0.07, 'completeness': 0.08
}


def get_weights_for_type(proposal_type):
    if proposal_type == 'marketing_only':
        return MARKETING_ONLY_WEIGHTS
    elif proposal_type == 'no_marketing':
        return NO_MARKETING_WEIGHTS
    return FULL_WEIGHTS


def compute_content_hash(proposal_text, proposal_type):
    """SHA-256 hash of proposal content + type for dedup/caching"""
    content = f"{proposal_type}::{proposal_text.strip()}"
    return hashlib.sha256(content.encode('utf-8')).hexdigest()


def bucket_score(score, step=5):
    """Round a score to the nearest step (e.g. 5) to reduce LLM scoring noise"""
    s = int(round(float(score or 0) / step) * step)
    return max(0, min(100, s))


def calculate_weighted_score(scores, proposal_type):
    weights = get_weights_for_type(proposal_type)
    total = 0
    for category, weight in weights.items():
        score_data = scores.get(category, 0)
        score = score_data.get('score', 0) if isinstance(score_data, dict) else score_data
        total += float(score or 0) * weight
    return round(total, 2)


def determine_tier(score):
    if score >= 85:
        return 'A'
    elif score >= 70:
        return 'B'
    elif score >= 60:
        return 'C'
    return 'D'


def get_tier_description(tier):
    return {
        'A': 'Exceptional - Your proposal demonstrates strong potential for top-tier publishers.',
        'B': 'Strong Foundation - With targeted improvements, your proposal could reach A-tier status.',
        'C': 'Developing - Your proposal shows promise but needs significant strengthening in key areas.',
        'D': 'Early Stage - Your proposal needs substantial work before submission to publishers.'
    }.get(tier, '')


def calculate_advance_from_platform(tier, platform_data, marketing_text=''):
    """
    Calculate advance estimate from structured platform numbers.

    Formula: projected_first_year_copies × $4 royalty / 2 = advance ceiling
    Conversion rates per channel:
        email_list            × 3%
        instagram + tiktok    × 0.7%
        linkedin              × 1%
        youtube               × 1.5%
        podcast               × 2%
        speaking × avg_audience × 7%
        bulk_orders           × 1 (direct)

    Tier caps:
        C / D → $0, not viable
        B     → capped at $10,000; viable only if ceiling > $2,000
        A     → floor $10,000, ceiling $250,000

    marketing_text is accepted for future use but is not part of the formula.
    The internal math is NOT exposed to users — only lowRange/highRange/confidence.
    """
    if tier in ('D', 'C'):
        return {
            'lowRange': 0,
            'highRange': 0,
            'viable': False,
            'confidence': 'Low',
            'reasoning': 'Proposal needs significant development before it could attract a traditional publishing advance.'
        }

    pd = platform_data or {}

    def safe_int(key):
        val = pd.get(key)
        if val is None:
            return None
        try:
            v = int(val)
            return v if v >= 0 else None
        except (ValueError, TypeError):
            return None

    email_list   = safe_int('email_list')
    instagram    = safe_int('instagram_followers')
    tiktok       = safe_int('tiktok_followers')
    linkedin     = safe_int('linkedin_followers')
    youtube      = safe_int('youtube_subscribers')
    podcast      = safe_int('podcast_audience')
    speaking     = safe_int('speaking_engagements')
    avg_audience = safe_int('avg_audience_per_talk')
    bulk         = safe_int('bulk_orders')

    speaking_valid = (speaking is not None and avg_audience is not None)
    individual_fields = [email_list, instagram, tiktok, linkedin, youtube, podcast, bulk]
    populated = sum(1 for f in individual_fields if f is not None) + (1 if speaking_valid else 0)

    # Confidence based on how many fields were filled
    if populated >= 3:
        confidence = 'High'
    elif populated >= 1:
        confidence = 'Medium'
    else:
        confidence = 'Low'

    # All-empty fallback — no platform data provided
    if populated == 0:
        if tier == 'B':
            return {
                'lowRange': 0, 'highRange': 5000, 'viable': True,
                'confidence': 'Low',
                'reasoning': 'No platform data provided. Range based on tier assessment alone.'
            }
        else:  # A
            return {
                'lowRange': 10000, 'highRange': 25000, 'viable': True,
                'confidence': 'Low',
                'reasoning': 'No platform data provided. Range based on tier assessment alone.'
            }

    # Project first-year copy sales
    copies = 0.0
    if email_list:       copies += email_list   * 0.03
    if instagram:        copies += instagram    * 0.007
    if tiktok:           copies += tiktok       * 0.007
    if linkedin:         copies += linkedin     * 0.01
    if youtube:          copies += youtube      * 0.015
    if podcast:          copies += podcast      * 0.02
    if speaking_valid:   copies += speaking * avg_audience * 0.07
    if bulk:             copies += bulk  # 1:1

    total_copies = max(0, int(copies))

    # Publisher formula: advance ceiling = (copies × $4 royalty) / 2
    ceiling = (total_copies * 4) / 2

    def round_to(n, base):
        return int(round(n / base) * base)

    if tier == 'B':
        capped = min(ceiling, 10000)
        if capped <= 2000:
            return {
                'lowRange': 0, 'highRange': 0, 'viable': False,
                'confidence': confidence,
                'reasoning': 'Platform reach is not sufficient for a traditional publishing advance at this tier.'
            }
        high = round_to(capped, 500)
        low  = round_to(high * 0.5, 500)
        return {
            'lowRange': low, 'highRange': high, 'viable': True,
            'confidence': confidence,
            'reasoning': 'Based on estimated first-year platform reach.'
        }
    else:  # A tier
        floored = max(ceiling, 10000)
        capped  = min(floored, 250000)
        high = round_to(capped, 1000)
        low  = max(10000, round_to(high * 0.6, 1000))
        return {
            'lowRange': low, 'highRange': high, 'viable': True,
            'confidence': confidence,
            'reasoning': 'Based on estimated first-year platform reach.'
        }


def compute_advance_estimate(evaluation):
    """Compute advance estimate using platform numbers from the evaluation dict.
    Always call this on any evaluation dict to ensure correct ranges.
    Expects evaluation['platform_data'] to be set before calling (embedded by
    process_evaluation_background); falls back gracefully if absent."""
    tier           = evaluation.get('tier', 'D')
    platform_data  = evaluation.get('platform_data') or {}
    marketing_text = evaluation.get('marketing_text', '')
    adv            = evaluation.get('advanceEstimate') or {}

    result = calculate_advance_from_platform(tier, platform_data, marketing_text)

    adv['lowRange']   = result['lowRange']
    adv['highRange']  = result['highRange']
    adv['viable']     = result['viable']
    adv['confidence'] = result['confidence']
    # Preserve any AI-generated reasoning unless the proposal is not viable
    if not adv.get('reasoning') or not result['viable']:
        adv['reasoning'] = result['reasoning']

    evaluation['advanceEstimate'] = adv
    evaluation['advance_estimate'] = {
        'low':   adv['lowRange'],
        'high':  adv['highRange'],
        'notes': adv.get('reasoning', '')
    }
    return evaluation


def evaluate_proposal(proposal_text, proposal_type='full', author_name='', book_title='',
                      platform_data=None):
    """Evaluate proposal using OpenAI with comprehensive analysis"""

    if proposal_type == 'marketing_only':
        evaluation_focus = "You are evaluating ONLY the Marketing & Platform section. Score all other categories as 0."
    elif proposal_type == 'no_marketing':
        evaluation_focus = "This proposal does NOT include a Marketing section. Score Marketing as 0. Evaluate all other sections normally."
    else:
        evaluation_focus = "This is a FULL proposal submission. Evaluate all categories comprehensively."

    system_prompt = """You are an elite literary agent with 25+ years of experience evaluating book proposals for major publishers. You have placed hundreds of books with advances ranging from $50,000 to $2 million+. Your evaluations are known for being thorough, actionable, and honest.

Your evaluation style:
- Be specific and cite examples from the actual proposal text
- Provide actionable feedback that authors can implement immediately
- Be encouraging but honest about weaknesses
- Think like a publisher evaluating commercial viability
- Score distribution: use the full 0-100 range meaningfully. A score of 50 means average, 70 means solid, 85+ means exceptional. If your scores for different sections all fall within a 10-point range of each other, you are not differentiating enough — push yourself to find what truly excels and what truly needs work."""

    # Build a human-readable platform data block (only include populated fields)
    pd = platform_data or {}
    def _pv(key): return pd.get(key)
    platform_lines = []
    if _pv('email_list'):          platform_lines.append(f"  Email list: {_pv('email_list'):,}")
    if _pv('instagram_followers'): platform_lines.append(f"  Instagram followers: {_pv('instagram_followers'):,}")
    if _pv('linkedin_followers'):  platform_lines.append(f"  LinkedIn followers: {_pv('linkedin_followers'):,}")
    if _pv('youtube_subscribers'): platform_lines.append(f"  YouTube subscribers: {_pv('youtube_subscribers'):,}")
    if _pv('tiktok_followers'):    platform_lines.append(f"  TikTok followers: {_pv('tiktok_followers'):,}")
    if _pv('podcast_audience'):    platform_lines.append(f"  Podcast audience: {_pv('podcast_audience'):,}")
    if _pv('speaking_engagements') and _pv('avg_audience_per_talk'):
        platform_lines.append(f"  Speaking: {_pv('speaking_engagements')} engagements/year, avg {_pv('avg_audience_per_talk'):,} audience")
    if _pv('bulk_orders'):         platform_lines.append(f"  Confirmed bulk orders: {_pv('bulk_orders'):,}")
    if platform_lines:
        platform_block = ("VERIFIED AUTHOR PLATFORM DATA (structured input submitted alongside the proposal —\n"
                          "use this when evaluating the credibility of marketing claims and when checking\n"
                          "comp/sales alignment; do NOT flag claims as unrealistic if the platform numbers\n"
                          "below support them):\n" + "\n".join(platform_lines))
    else:
        platform_block = "VERIFIED AUTHOR PLATFORM DATA: None provided."

    user_prompt = f"""{evaluation_focus}

Evaluate this book proposal comprehensively.

AUTHOR: {author_name}
BOOK TITLE: {book_title}

{platform_block}

PROPOSAL TEXT:
{proposal_text[:50000]}

---

Provide your evaluation as a JSON object with this EXACT structure:

{{
    "executiveSummary": "<3-5 sentence executive summary of strengths and areas for improvement>",

    "redFlags": ["<each entry must be a complete explanatory sentence describing a specific problem and why it matters to a publisher — e.g. 'The author cites no verifiable platform numbers, which makes it impossible for a publisher to assess the book's commercial reach.' or 'The comparative titles are all 10+ years old, suggesting the author has not researched the current market.' Do NOT use code slugs or one-word labels. Return an empty array if there are no genuine red flags.>"],

    "scores": {{
        "marketing": {{"score": <0-100>, "weight": 30}},
        "overview": {{"score": <0-100>, "weight": 20}},
        "credentials": {{"score": <0-100>, "weight": 15}},
        "comps": {{"score": <0-100>, "weight": 10}},
        "writing": {{"score": <0-100>, "weight": 15}},
        "outline": {{"score": <0-100>, "weight": 5}},
        "completeness": {{"score": <0-100>, "weight": 5}}
    }},

    "detailedAnalysis": {{
        "marketing": {{
            "currentState": "<2-3 sentences describing current state>",
            "strengths": "<what's working well>",
            "gaps": "<what's missing or weak>",
            "exampleOfExcellence": "<specific example of what A-tier looks like for this category>",
            "actionItems": ["<specific action 1>", "<specific action 2>", "<specific action 3>"]
        }},
        "overview": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"]
        }},
        "credentials": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"]
        }},
        "comps": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"]
        }},
        "writing": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"],
            "writingExamples": {{
                "strongPassage": "<quote a strong passage from the proposal if available>",
                "improvementExample": "<quote a passage that could be improved>"
            }}
        }},
        "outline": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"]
        }},
        "completeness": {{
            "currentState": "<2-3 sentences>",
            "strengths": "<what's working>",
            "gaps": "<what's missing>",
            "exampleOfExcellence": "<A-tier example>",
            "actionItems": ["<action 1>", "<action 2>", "<action 3>"]
        }}
    }},

    "strengths": ["<top strength 1>", "<top strength 2>", "<top strength 3>"],
    "improvements": ["<top improvement 1>", "<top improvement 2>", "<top improvement 3>"],

    "priorityActionPlan": [
        {{"priority": 1, "action": "<most important action>", "timeline": "<e.g., This week>", "impact": "<why this matters>"}},
        {{"priority": 2, "action": "<second action>", "timeline": "<e.g., Next 2 weeks>", "impact": "<why this matters>"}},
        {{"priority": 3, "action": "<third action>", "timeline": "<e.g., Next month>", "impact": "<why this matters>"}}
    ],

    "pathToATier": "<2-3 sentences describing the specific path this author needs to take to reach A-tier status>",

    "advanceEstimate": {{
        "viable": <true or false>,
        "lowRange": <number or 0>,
        "highRange": <number or 0>,
        "confidence": "<Low, Medium, or High>",
        "reasoning": "<2-3 sentences explaining the estimate>"
    }},

    "recommendedNextSteps": ["<step 1>", "<step 2>", "<step 3>", "<step 4>", "<step 5>"],

    "contradictions": ["<full-sentence description of a cross-section inconsistency a real editor would notice — e.g., 'Marketing plan promises 50 speaking engagements annually but the credentials section lists no prior speaking history.' Leave array empty if none found.>"]
}}

SCORING GUIDELINES:
- A-Tier (85-100): Exceptional, publisher-ready proposal with strong platform
- B-Tier (70-84): Strong foundation, improvements needed to be publisher-ready
- C-Tier (60-69): Developing, shows promise but needs significant strengthening
- D-Tier (Below 60): Early stage, needs substantial work before submission

IMPORTANT SCORING RULES:
1. Score each category in multiples of 5 (e.g. 60, 65, 70, 75, 80, 85, 90, 95). Never use scores like 72 or 83.
2. Your scores MUST align with the tier. If a proposal deserves a B, score it 70-84. If it deserves a C, score it 60-69. Do NOT give a score of 70 and call it C-tier - that would be B-tier.
3. Differentiate meaningfully: a weak section on an otherwise strong proposal should score 40-55, not 65. Do not cluster all scores within 15 points of each other — identify what truly excels and what truly needs work and score accordingly.

COMPARATIVE TITLES ("comps") SCORING RUBRIC:
When scoring the comps category, explicitly check all three of the following:
- Recency: Are the comps published within the last 5 years? Titles older than 5 years are a red flag — flag them and reduce the score.
- Genre and audience alignment: Are the comps in the same genre and targeting the same readership as this book? Mismatched comps signal the author does not understand the market.
- Platform-to-comp-sales alignment: Do the cited comp sales figures match the author's current platform size? If the author has fewer than 5,000 combined followers/subscribers but cites books that sold 500,000+ copies, score this 30-45 and explain in the gaps field that editors immediately reject this framing — strong comps should reflect realistic sales given the author's reach, not aspirational outliers.

CROSS-SECTION CONSISTENCY (the contradictions field):
After completing all section scores, scan for contradictions between sections that a real editor would notice. Write each as a concrete, specific sentence. Examples of what to look for:
- Marketing plan promises a large number of speaking engagements but credentials list no prior speaking history
- Comparative titles all sold 200K+ copies but the author's platform has no demonstrated reach
- Author claims deep niche expertise in credentials but the overview targets a mass-market general audience
- Writing sample quality is inconsistent with the professional credentials claimed
If no meaningful contradictions exist, return an empty array. Do not manufacture contradictions.

ADVANCE ESTIMATE: The advance ranges will be calculated automatically. For the advanceEstimate field, just provide your reasoning about commercial viability. Set lowRange and highRange to 0 — they will be overridden by the system.

Return ONLY the JSON object, no other text."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0,
            seed=42,
            max_tokens=6000
        )

        response_text = response.choices[0].message.content.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]

        evaluation = json.loads(response_text.strip())

        # Bucket individual category scores to nearest 5 to eliminate LLM noise
        scores = evaluation.get('scores', {})
        for cat_key, cat_data in scores.items():
            if isinstance(cat_data, dict) and 'score' in cat_data:
                cat_data['score'] = bucket_score(cat_data['score'])

        # Calculate weighted total score from bucketed scores
        evaluation['total_score'] = calculate_weighted_score(scores, proposal_type)
        evaluation['tier'] = determine_tier(evaluation['total_score'])
        evaluation['tierDescription'] = get_tier_description(evaluation['tier'])
        evaluation['proposal_type'] = proposal_type

        # Compute advance estimate deterministically from score and tier
        compute_advance_estimate(evaluation)

        # Backward-compat aliases so old templates/emails still work
        evaluation['overall_score'] = evaluation['total_score']
        evaluation['summary'] = evaluation.get('executiveSummary', '')
        evaluation['red_flags'] = evaluation.get('redFlags', [])
        evaluation['next_steps'] = evaluation.get('recommendedNextSteps', [])
        # Map scores to old categories format for backward compat
        cats = {}
        for key, score_data in scores.items():
            s = score_data.get('score', 0) if isinstance(score_data, dict) else score_data
            da = evaluation.get('detailedAnalysis', {}).get(key, {})
            cats[key] = {
                'score': s,
                'feedback': da.get('currentState', ''),
                'priority_actions': da.get('actionItems', [])
            }
        evaluation['categories'] = cats

        return evaluation
    except Exception as e:
        print(f"Evaluation error: {e}")
        traceback.print_exc()
        return None


def generate_pdf_report(proposal):
    """Generate PDF report for proposal using xhtml2pdf"""
    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)

    # Ensure numeric values are properly typed
    advance = evaluation.get('advance_estimate')
    if advance and isinstance(advance, dict):
        try:
            advance['low'] = float(advance.get('low', 0) or 0)
            advance['high'] = float(advance.get('high', 0) or 0)
        except (ValueError, TypeError):
            advance['low'] = 0
            advance['high'] = 0

    adv_est = evaluation.get('advanceEstimate')
    if adv_est and isinstance(adv_est, dict):
        try:
            adv_est['lowRange'] = float(adv_est.get('lowRange', 0) or 0)
            adv_est['highRange'] = float(adv_est.get('highRange', 0) or 0)
        except (ValueError, TypeError):
            adv_est['lowRange'] = 0
            adv_est['highRange'] = 0

    scores = evaluation.get('scores', {})
    for key, score_data in scores.items():
        if isinstance(score_data, dict):
            try:
                score_data['score'] = float(score_data.get('score', 0) or 0)
            except (ValueError, TypeError):
                score_data['score'] = 0

    categories = evaluation.get('categories', {})
    for key, cat in categories.items():
        if isinstance(cat, dict):
            try:
                cat['score'] = float(cat.get('score', 0) or 0)
            except (ValueError, TypeError):
                cat['score'] = 0

    html = render_template('report_pdf.html', proposal=proposal, evaluation=evaluation)
    pdf_buffer = BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=pdf_buffer)
    if pisa_status.err:
        print(f"PDF generation error: {pisa_status.err}")
        raise Exception(f"PDF generation failed with {pisa_status.err} errors")
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


def send_email(to_email, subject, html_content, attachments=None):
    """Send email via SMTP"""
    if not SMTP_USER or not SMTP_PASSWORD:
        print("Email not configured - skipping")
        return False
    
    try:
        msg = MIMEMultipart()
        msg['From'] = FROM_EMAIL
        msg['To'] = to_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(html_content, 'html'))
        
        if attachments:
            for filename, content in attachments:
                attachment = MIMEApplication(content)
                attachment.add_header('Content-Disposition', 'attachment', filename=filename)
                msg.attach(attachment)
        
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)
        
        return True
    except smtplib.SMTPAuthenticationError as e:
        print(f"Email SMTP auth error: {e}")
        print("HINT: Gmail requires an App Password (16-char) when 2FA is enabled.")
        print("Generate one at: https://myaccount.google.com/apppasswords")
        return False
    except Exception as e:
        print(f"Email error: {e}")
        return False


def send_author_notification(proposal):
    """Send evaluation results to the author"""
    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)

    score_display = f"{proposal.overall_score:.0f}" if proposal.overall_score is not None else "N/A"
    summary = evaluation.get('executiveSummary', '') or evaluation.get('summary', 'See attached report for details.')
    tier_desc = evaluation.get('tierDescription', '')
    app_url = APP_BASE_URL

    subject = f"Your Book Proposal Evaluation - {proposal.book_title}"

    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
        <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
            <div style="text-align: center; margin-bottom: 20px;">
                <h1 style="color: #2D1B69; margin-bottom: 5px;">Write It Great</h1>
                <p style="color: #666; font-style: italic;">Elite Ghostwriting &amp; Publishing Services</p>
            </div>

            <p>Dear {proposal.author_name},</p>

            <p>Thank you for submitting your book proposal for <strong>"{proposal.book_title}"</strong> to Write It Great. Our AI-powered evaluation system has completed a comprehensive analysis of your proposal.</p>

            <div style="background: #f8f6f9; padding: 20px; border-radius: 8px; margin: 20px 0; text-align: center;">
                <div style="font-size: 48px; font-weight: bold; color: {'#2e7d32' if proposal.tier == 'A' else '#1976d2' if proposal.tier == 'B' else '#f57c00' if proposal.tier == 'C' else '#d32f2f'};">{proposal.tier or 'N/A'}-Tier</div>
                <div style="font-size: 24px; color: #2D1B69; margin: 10px 0;">{score_display}/100</div>
                <div style="color: #666; font-style: italic;">{tier_desc}</div>
            </div>

            <h3 style="color: #2D1B69;">Executive Summary</h3>
            <p>{summary}</p>

            <div style="text-align: center; margin: 25px 0;">
                <a href="{app_url}/results/{proposal.submission_id}" style="display: inline-block; padding: 14px 28px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">View Your Full Report</a>
            </div>

            <p>Your complete evaluation report is also attached as a PDF for your records.</p>

            <p>A member of our team will reach out within 3-5 business days to discuss your results and next steps.</p>

            <hr style="border: none; border-top: 1px solid #eee; margin: 25px 0;">

            <p>Best regards,<br><strong>The Write It Great Team</strong><br><a href="https://www.writeitgreat.com" style="color: #2D1B69;">www.writeitgreat.com</a></p>
        </div>
    </body>
    </html>
    """

    # Try to generate PDF attachment, but send email even if PDF fails
    attachments = None
    try:
        pdf_content = generate_pdf_report(proposal)
        attachments = [(f"Book_Proposal_Evaluation_{proposal.submission_id}.pdf", pdf_content)]
    except Exception as e:
        print(f"PDF generation for author email failed (sending without attachment): {e}")

    try:
        return send_email(proposal.author_email, subject, html_content, attachments)
    except Exception as e:
        print(f"Author notification error: {e}")
        return False


def send_team_notification(proposal):
    """Send notification to team about new submission"""
    try:
        evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
        if evaluation:
            compute_advance_estimate(evaluation)

        score_display = f"{proposal.overall_score:.0f}" if proposal.overall_score is not None else "N/A"
        summary = evaluation.get('executiveSummary', '') or evaluation.get('summary', 'No summary')

        subject = f"[{proposal.tier or 'N/A'}-Tier] New Proposal: {proposal.book_title}"

        # Build score breakdown
        scores = evaluation.get('scores', {})
        score_rows = ""
        category_labels = {
            'marketing': 'Marketing & Platform', 'overview': 'Overview & Concept',
            'credentials': 'Author Credentials', 'comps': 'Comparative Titles',
            'writing': 'Sample Writing', 'outline': 'Book Outline', 'completeness': 'Completeness'
        }
        for key, label in category_labels.items():
            score_data = scores.get(key, {})
            s = score_data.get('score', 0) if isinstance(score_data, dict) else score_data
            score_rows += f"<tr><td>{label}</td><td style='text-align:center;font-weight:bold;'>{int(float(s or 0))}/100</td></tr>"

        html_content = f"""
        <html>
        <body style="font-family: Arial, sans-serif; line-height: 1.6;">
            <div style="background: #2D1B69; color: white; padding: 20px; text-align: center;">
                <h1 style="margin: 0;">Write It Great</h1>
                <p style="margin: 5px 0 0;">New Book Proposal Submission</p>
            </div>
            <div style="padding: 20px;">
                <div style="text-align: center; margin: 20px 0;">
                    <span style="display: inline-block; padding: 10px 20px; font-size: 24px; font-weight: bold; border-radius: 5px; color: white; background: {'#2e7d32' if proposal.tier == 'A' else '#1976d2' if proposal.tier == 'B' else '#f57c00' if proposal.tier == 'C' else '#d32f2f'};">
                        TIER {proposal.tier or 'N/A'}
                    </span>
                    <span style="font-size: 36px; font-weight: bold; color: #2D1B69; margin-left: 20px;">{score_display}/100</span>
                </div>

                <h2>Submission Details</h2>
                <table style="width: 100%; border-collapse: collapse;">
                    <tr><td style="padding: 8px; border-bottom: 1px solid #eee; font-weight: bold; width: 150px;">Author</td><td style="padding: 8px; border-bottom: 1px solid #eee;">{proposal.author_name}</td></tr>
                    <tr><td style="padding: 8px; border-bottom: 1px solid #eee; font-weight: bold;">Email</td><td style="padding: 8px; border-bottom: 1px solid #eee;"><a href="mailto:{proposal.author_email}">{proposal.author_email}</a></td></tr>
                    <tr><td style="padding: 8px; border-bottom: 1px solid #eee; font-weight: bold;">Book Title</td><td style="padding: 8px; border-bottom: 1px solid #eee;">{proposal.book_title}</td></tr>
                    <tr><td style="padding: 8px; border-bottom: 1px solid #eee; font-weight: bold;">Submission ID</td><td style="padding: 8px; border-bottom: 1px solid #eee;">{proposal.submission_id}</td></tr>
                </table>

                <h2>Executive Summary</h2>
                <p>{summary}</p>

                {'<h2>Score Breakdown</h2><table style="width: 100%; border-collapse: collapse; border: 1px solid #ddd;"><tr style="background: #2D1B69; color: white;"><th style="padding: 10px; text-align: left;">Category</th><th style="padding: 10px; text-align: center;">Score</th></tr>' + score_rows + '</table>' if score_rows else ''}

                <div style="margin-top: 20px;">
                    <a href="{APP_BASE_URL}/admin/proposal/{proposal.submission_id}" style="display: inline-block; padding: 12px 24px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 5px; font-weight: bold;">View Evaluation</a>
                    <a href="{APP_BASE_URL}/admin/proposal/{proposal.submission_id}/view-proposal" style="display: inline-block; padding: 12px 24px; background: #2D1B69; color: white; text-decoration: none; border-radius: 5px; font-weight: bold; margin-left: 10px;">Read Proposal Text</a>
                    <a href="{APP_BASE_URL}/admin/proposal/{proposal.submission_id}/download-proposal" style="display: inline-block; padding: 12px 24px; background: white; color: #2D1B69; text-decoration: none; border-radius: 5px; font-weight: bold; margin-left: 10px; border: 2px solid #2D1B69;">Download Original File</a>
                </div>
            </div>
        </body>
        </html>
        """

        success = True
        for email in TEAM_EMAILS:
            if email.strip():
                if not send_email(email.strip(), subject, html_content):
                    success = False

        return success
    except Exception as e:
        print(f"Team notification error: {e}")
        traceback.print_exc()
        return False


# ============================================================================
# BACKGROUND PROCESSING
# ============================================================================

def process_evaluation_background(app_obj, submission_id, proposal_text, proposal_type,
                                   author_name='', book_title='', platform_data=None):
    """Run OpenAI evaluation and email notifications in a background thread"""
    with app_obj.app_context():
        try:
            proposal = Proposal.query.filter_by(submission_id=submission_id).first()
            if not proposal:
                print(f"Background eval: proposal {submission_id} not found")
                return

            # Content-hash caching: reuse evaluation if same proposal was evaluated before
            c_hash = compute_content_hash(proposal_text, proposal_type)
            proposal.content_hash = c_hash

            cached = (Proposal.query
                      .filter(Proposal.content_hash == c_hash,
                              Proposal.evaluation_json.isnot(None),
                              Proposal.id != proposal.id)
                      .first())

            if cached and cached.evaluation_json:
                evaluation = json.loads(cached.evaluation_json)
                print(f"Background eval: using cached result for {submission_id} (matched {cached.submission_id})")
            else:
                evaluation = evaluate_proposal(proposal_text, proposal_type, author_name, book_title,
                                               platform_data=platform_data)

            if not evaluation:
                proposal.status = 'error'
                db.session.commit()
                print(f"Background eval: OpenAI evaluation failed for {submission_id}")
                return

            # Embed platform data so the advance calculator (and all downstream loaders)
            # can access it from evaluation_json without touching the Proposal model.
            evaluation['platform_data'] = platform_data or {}
            evaluation['marketing_text'] = proposal.marketing_strategy or ''

            # Always recompute advance estimate using the current submission's platform data
            compute_advance_estimate(evaluation)

            proposal.tier = evaluation.get('tier', 'C')
            proposal.overall_score = evaluation.get('total_score', evaluation.get('overall_score', 50))
            proposal.evaluation_json = json.dumps(evaluation)
            proposal.status = 'submitted'
            db.session.commit()

            # Send emails
            try:
                if send_author_notification(proposal):
                    proposal.author_email_sent = True
                if send_team_notification(proposal):
                    proposal.team_email_sent = True
                db.session.commit()
            except Exception as email_error:
                print(f"Email error (non-fatal): {email_error}")

            print(f"Background eval: completed for {submission_id}")

        except Exception as e:
            print(f"Background eval error for {submission_id}: {e}")
            traceback.print_exc()
            try:
                proposal = Proposal.query.filter_by(submission_id=submission_id).first()
                if proposal:
                    proposal.status = 'error'
                    db.session.commit()
            except Exception:
                pass


def send_author_milestone_email(proposal, new_status):
    """Send email to author when their proposal reaches a major milestone"""
    if new_status not in AUTHOR_EMAIL_MILESTONES:
        return False

    friendly_status = AUTHOR_STATUS_LABELS.get(new_status, new_status)
    app_url = APP_BASE_URL

    # Custom message per milestone
    messages = {
        'author_call_scheduled': 'We would love to schedule a call to discuss your proposal in more detail. A team member will reach out shortly with available times.',
        'contract_sent': 'We are excited to move forward! A contract has been sent for your review. Please check your email for the details.',
        'contract_signed': 'Your contract has been signed. Welcome aboard! We are thrilled to be working with you on this project.',
        'shopping': 'Great news! Your proposal is now being presented to publishers. We will keep you updated on any interest.',
        'publisher_interest': 'Exciting development! One or more publishers have expressed interest in your book. We will be in touch with more details soon.',
        'offer_received': 'Wonderful news! We have received an offer for your book. A team member will contact you to discuss the details.',
        'deal_closed': 'Congratulations! Your book deal has been finalized. A team member will be in touch with the next steps.',
        'declined': 'After careful consideration, we have decided not to move forward with your proposal at this time. We wish you the best in your publishing journey.',
    }

    milestone_msg = messages.get(new_status, f'Your proposal status has been updated to: {friendly_status}.')

    subject = f"Update on Your Proposal - {proposal.book_title}"
    html_content = f"""
    <html>
    <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
        <div style="max-width: 600px; margin: 0 auto; padding: 20px;">
            <div style="text-align: center; margin-bottom: 20px;">
                <h1 style="color: #2D1B69; margin-bottom: 5px;">Write It Great</h1>
                <p style="color: #666; font-style: italic;">Elite Ghostwriting &amp; Publishing Services</p>
            </div>

            <p>Dear {proposal.author_name},</p>

            <p>We have an update regarding your book proposal for <strong>"{proposal.book_title}"</strong>.</p>

            <div style="background: #f8f6f9; padding: 20px; border-radius: 8px; margin: 20px 0; text-align: center;">
                <div style="font-size: 18px; font-weight: bold; color: #2D1B69;">{friendly_status}</div>
            </div>

            <p>{milestone_msg}</p>

            <div style="text-align: center; margin: 25px 0;">
                <a href="{app_url}/author/dashboard" style="display: inline-block; padding: 14px 28px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">View Your Dashboard</a>
            </div>

            <hr style="border: none; border-top: 1px solid #eee; margin: 25px 0;">

            <p>Best regards,<br><strong>The Write It Great Team</strong><br><a href="https://www.writeitgreat.com" style="color: #2D1B69;">www.writeitgreat.com</a></p>
        </div>
    </body>
    </html>
    """

    try:
        return send_email(proposal.author_email, subject, html_content)
    except Exception as e:
        print(f"Milestone email error: {e}")
        return False


# ============================================================================
# COACHING EMAIL FUNCTIONS
# ============================================================================

def _coaching_email_header():
    return """
    <div style="text-align:center;margin-bottom:20px;">
        <h1 style="color:#2D1B69;margin-bottom:5px;">Write It Great</h1>
        <p style="color:#666;font-style:italic;">Proposal Coaching Program</p>
    </div>"""

def _coaching_email_footer():
    return """
    <hr style="border:none;border-top:1px solid #eee;margin:25px 0;">
    <p>Best regards,<br><strong>The Write It Great Team</strong><br>
    <a href="https://www.writeitgreat.com" style="color:#2D1B69;">www.writeitgreat.com</a></p>"""


def send_coaching_welcome_email(author, enrollment):
    """Welcome email when an author enrolls in the coaching program"""
    app_url = APP_BASE_URL
    subject = "Welcome to the Write It Great Coaching Program!"
    book_title_line = f" for <strong>\"{enrollment.book_title}\"</strong>" if enrollment.book_title else ""
    html_content = f"""
    <html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Dear {author.name},</p>
        <p>Welcome! You've just enrolled in the Write It Great Coaching Program{book_title_line}. We're excited to help you develop a publisher-ready book proposal, step by step.</p>
        <div style="background:#f8f6f9;padding:20px;border-radius:8px;margin:20px 0;">
            <h3 style="color:#2D1B69;margin-top:0;">What to expect</h3>
            <ul style="margin:0;padding-left:20px;">
                <li>7 focused modules covering every section of your proposal</li>
                <li>An AI coach available in each module to guide you through the material</li>
                <li>Homework assignments reviewed by AI with specific, actionable feedback</li>
                <li>Each module unlocks after the previous one is approved</li>
                <li>When all 7 modules are complete, your proposal is ready for professional evaluation</li>
            </ul>
        </div>
        <div style="text-align:center;margin:25px 0;">
            <a href="{app_url}/author/coaching" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">Start Module 1 →</a>
        </div>
        <p>Your first module — <strong>Your Book Concept &amp; Hook</strong> — is ready and waiting. Start whenever you're ready.</p>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, subject, html_content)
    except Exception as e:
        print(f"Coaching welcome email error: {e}")
        return False


def send_coaching_module_unlocked_email(author, module_info, enrollment):
    """Email when a new module unlocks"""
    app_url = APP_BASE_URL
    module_num = module_info['order']
    module_title = module_info['title']
    subject = f"Module {module_num} Unlocked: {module_title}"
    html_content = f"""
    <html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Dear {author.name},</p>
        <p>Great work completing the previous module! Your next module is now unlocked.</p>
        <div style="background:linear-gradient(135deg,#f0fdf4,#dcfce7);border:1px solid #86efac;padding:20px;border-radius:8px;margin:20px 0;">
            <div style="font-size:2rem;margin-bottom:8px;">{module_info['icon']}</div>
            <div style="font-size:0.8rem;color:#166534;font-weight:600;text-transform:uppercase;letter-spacing:0.05em;">Module {module_num} of {len(COACHING_MODULES)}</div>
            <div style="font-size:1.25rem;font-weight:700;color:#14532d;margin:4px 0;">{module_title}</div>
            <div style="font-size:0.9rem;color:#166534;">{module_info['subtitle']}</div>
        </div>
        <p>{module_info['description']}</p>
        <div style="text-align:center;margin:25px 0;">
            <a href="{app_url}/author/coaching/module/{module_num}" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">Start Module {module_num} →</a>
        </div>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, subject, html_content)
    except Exception as e:
        print(f"Module unlocked email error: {e}")
        return False


def send_coaching_homework_reminder_email(author, module_info, enrollment):
    """Reminder email when homework hasn't been submitted after several days"""
    app_url = APP_BASE_URL
    module_num = module_info['order']
    subject = f"Don't forget — Module {module_num} homework is waiting for you"
    html_content = f"""
    <html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Dear {author.name},</p>
        <p>We noticed you haven't submitted your homework for <strong>Module {module_num}: {module_info['title']}</strong> yet — and that's okay, life gets busy. Just a friendly nudge to keep the momentum going.</p>
        <div style="background:#f8f6f9;padding:16px 20px;border-radius:8px;border-left:4px solid #2D1B69;margin:20px 0;">
            <strong style="color:#2D1B69;">Your homework:</strong>
            <p style="margin:8px 0 0;">{module_info['homework_prompt']}</p>
        </div>
        <p>Remember, you don't need it to be perfect — the AI coach will give you specific feedback and you can revise. The only way to make progress is to start.</p>
        <div style="text-align:center;margin:25px 0;">
            <a href="{app_url}/author/coaching/module/{module_num}" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">Submit Homework →</a>
        </div>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, subject, html_content)
    except Exception as e:
        print(f"Homework reminder email error: {e}")
        return False


def send_coaching_homework_reviewed_email(author, module_info, submission):
    """Email when AI (or admin) has reviewed a homework submission"""
    app_url = APP_BASE_URL
    module_num = module_info['order']
    approved = submission.ai_approved
    subject = (f"Module {module_num} approved — next module unlocked!" if approved
               else f"Feedback on your Module {module_num} homework")
    status_block = (
        f'<div style="background:linear-gradient(135deg,#f0fdf4,#dcfce7);border:1px solid #86efac;padding:16px 20px;border-radius:8px;margin:20px 0;">'
        f'<strong style="color:#166534;">✓ Approved!</strong> Your homework for Module {module_num} has been approved. '
        f'{"Module " + str(module_num + 1) + " is now unlocked." if module_num < len(COACHING_MODULES) else "You have completed all modules — your proposal is ready for evaluation!"}'
        f'</div>'
    ) if approved else (
        f'<div style="background:#fef9c3;border:1px solid #fde047;padding:16px 20px;border-radius:8px;margin:20px 0;">'
        f'<strong style="color:#854d0e;">Revision needed</strong> — your homework has been reviewed and needs some adjustments before you can advance.'
        f'</div>'
    )
    html_content = f"""
    <html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Dear {author.name},</p>
        <p>Your homework for <strong>Module {module_num}: {module_info['title']}</strong> has been reviewed.</p>
        {status_block}
        <p><strong>Feedback:</strong></p>
        <p style="background:#f8f6f9;padding:16px;border-radius:8px;">{submission.ai_feedback or 'See your coaching dashboard for detailed feedback.'}</p>
        <div style="text-align:center;margin:25px 0;">
            <a href="{app_url}/author/coaching/module/{module_num}" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">{'Continue to Module ' + str(module_num + 1) + ' →' if approved and module_num < len(COACHING_MODULES) else 'View Feedback & Revise →'}</a>
        </div>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, subject, html_content)
    except Exception as e:
        print(f"Homework reviewed email error: {e}")
        return False


def send_coaching_complete_email(author, enrollment):
    """Email when the author completes all coaching modules"""
    app_url = APP_BASE_URL
    subject = "Congratulations — Your Proposal Is Ready for Evaluation!"
    book_title_line = f"for <strong>\"{enrollment.book_title}\"</strong> " if enrollment.book_title else ""
    html_content = f"""
    <html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Dear {author.name},</p>
        <p>This is a big moment. You've completed all 7 modules of the Write It Great Coaching Program {book_title_line}and built a complete book proposal from the ground up.</p>
        <div style="background:linear-gradient(135deg,#f0fdf4,#dcfce7);border:1px solid #86efac;padding:24px;border-radius:8px;margin:20px 0;text-align:center;">
            <div style="font-size:2.5rem;margin-bottom:8px;">🎉</div>
            <div style="font-size:1.25rem;font-weight:700;color:#14532d;">All 7 Modules Complete!</div>
            <div style="font-size:0.9rem;color:#166534;margin-top:6px;">Your proposal is ready for professional evaluation</div>
        </div>
        <p>The next step is to submit your completed proposal for a full professional evaluation. Our AI evaluation system will score your proposal across all 7 dimensions and give you a detailed report with your tier rating, score breakdown, advance estimate, and priority action plan.</p>
        <div style="text-align:center;margin:25px 0;">
            <a href="{app_url}/author/coaching" style="display:inline-block;padding:14px 28px;background:#B8F2B8;color:#1a3a1a;text-decoration:none;border-radius:8px;font-weight:bold;">Go to Your Dashboard →</a>
        </div>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, subject, html_content)
    except Exception as e:
        print(f"Coaching complete email error: {e}")
        return False


# ============================================================================
# ADMIN INVITE & ONE-PAGER EMAILS
# ============================================================================

def send_author_welcome_invite_email(author, token):
    """Admin-created account: send invite email with set-password link"""
    app_url = APP_BASE_URL
    set_url = f"{app_url}/author/reset-password/{token}"
    path_line = ''
    if author.assigned_path == 'one_pager':
        path_line = '<p>You\'ve been set up to start with the <strong>Quick One-Pager</strong> — a fast, guided experience to capture your book idea in one page.</p>'
    elif author.assigned_path == 'full_proposal':
        path_line = '<p>You\'ve been set up for the <strong>Full Proposal Program</strong> — a step-by-step guide to building a publisher-ready book proposal.</p>'
    html_content = f"""<html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        {_coaching_email_header()}
        <p>Hi {author.name},</p>
        <p>Welcome to <strong>Write It Great</strong>! Your account has been created and you're ready to get started.</p>
        {path_line}
        <p>Click the button below to set your password and access your account:</p>
        <p style="text-align:center;margin:2rem 0;">
            <a href="{set_url}" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">
                Set My Password &amp; Get Started →
            </a>
        </p>
        <p style="font-size:0.85rem;color:#666;">This link expires in 48 hours. If you have any questions, just reply to this email.</p>
        {_coaching_email_footer()}
    </div></body></html>"""
    try:
        return send_email(author.email, 'Your Write It Great Access Is Ready', html_content)
    except Exception as e:
        print(f"Invite email error: {e}")
        return False


def send_one_pager_submitted_notification(author, submission):
    """Notify the WIG team when an author submits their one-pager for review"""
    app_url = APP_BASE_URL
    admin_url = f"{app_url}/admin/one-pager/{submission.id}"
    answers = json.loads(submission.answers_json) if submission.answers_json else {}
    html_content = f"""<html><body style="font-family:Arial,sans-serif;line-height:1.6;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:20px;">
        <h2 style="color:#2D1B69;">📄 New One-Pager Submission</h2>
        <p><strong>{author.name}</strong> ({author.email}) has submitted their one-pager for review.</p>
        <table style="width:100%;border-collapse:collapse;margin:1rem 0;">
            <tr><td style="padding:6px;font-weight:bold;width:140px;">Book title:</td><td style="padding:6px;">{submission.book_title or '—'}</td></tr>
            <tr style="background:#f9f9f9;"><td style="padding:6px;font-weight:bold;">Problem solved:</td><td style="padding:6px;">{answers.get('problem','—')[:200]}</td></tr>
            <tr><td style="padding:6px;font-weight:bold;">Target reader:</td><td style="padding:6px;">{answers.get('reader','—')[:200]}</td></tr>
        </table>
        <p><a href="{admin_url}" style="display:inline-block;padding:12px 24px;background:#B8F2B8;color:#1a3a1a;text-decoration:none;border-radius:5px;font-weight:bold;">View Full One-Pager →</a></p>
    </div></body></html>"""
    for team_email in TEAM_EMAILS:
        if team_email.strip():
            send_email(team_email.strip(), f"One-Pager Submitted: {author.name}", html_content)


# ── Re-engagement email copy ──────────────────────────────────────────────────
REENGAGEMENT_PROMPTS = [
    "Who's the one person you'd hand this book to first, and why?",
    "What's the thing you wish someone had told you earlier that this book would say?",
    "If a journalist wrote a headline about your book's impact in 5 years, what would it say?",
]

REENGAGEMENT_TYPES = {
    'never_started':        'Your proposal is waiting — here\'s how to begin in 10 minutes',
    'stalled_one_pager':    'You\'ve already started something great — one question to get unstuck',
    'one_pager_to_full':    'Your one-pager showed real promise. Ready to take it further?',
    'stalled_full_proposal':'Look at the progress you\'ve made. Don\'t stop now.',
    'dormant_30_days':      'We saved your work. It\'s still here whenever you\'re ready.',
}


def send_reengagement_email(author, email_type, module_name='', completed_count=0):
    """Send a warm re-engagement email. Returns True if sent."""
    app_url = APP_BASE_URL
    subject_map = {
        'never_started':        'Your proposal is waiting — here\'s how to begin in 10 minutes',
        'stalled_one_pager':    'You\'ve already started something great ✨',
        'one_pager_to_full':    'Your one-pager showed real promise. Ready for the next step?',
        'stalled_full_proposal': ("You're " + str(completed_count) + " section" + ("s" if completed_count != 1 else "") + " in — don't stop now" if completed_count else "Keep going — don't stop now"),
        'dormant_30_days':      'We saved your work. It\'s still here.',
    }
    import random
    prompt = random.choice(REENGAGEMENT_PROMPTS)
    body_map = {
        'never_started': f"""<p>Hi {author.name},</p>
            <p>Your Write It Great account is all set up, and your book idea is waiting to come to life. Getting started is easier than you think — authors who try one question often can't stop.</p>
            <p><strong>Here's your first move:</strong> Pick up the One-Pager. It's five questions and about 45 minutes. That's all you need to turn your idea into something concrete.</p>
            <p>A question to get your pen moving: <em>"{prompt}"</em></p>""",
        'stalled_one_pager': f"""<p>Hi {author.name},</p>
            <p>You started something — and that takes guts. Your answers are saved and ready for you exactly where you left them.</p>
            <p>Here's a question to get unstuck: <em>"{prompt}"</em></p>
            <p>You don't need to have everything figured out. Just write what's true right now.</p>""",
        'one_pager_to_full': f"""<p>Hi {author.name},</p>
            <p>Your one-pager is done — and what you wrote showed genuine promise. A lot of authors stop there. The ones who get published don't.</p>
            <p>The Full Proposal Program takes you section by section with AI coaching at every step. Your one-pager content seeds it automatically.</p>""",
        'stalled_full_proposal': f"""<p>Hi {author.name},</p>
            <p>{"You've already completed " + str(completed_count) + " section" + ("s" if completed_count != 1 else "") + (" of your proposal" if completed_count else "") + "." if completed_count else "Your proposal draft is waiting."} That's real work. It would be a shame to leave it unfinished.</p>
            <p>Here's a question to get you back into it: <em>"{prompt}"</em></p>
            <p>{"Section " + str(completed_count + 1) if completed_count else "Your next section"} is unlocked and ready whenever you are.</p>""",
        'dormant_30_days': f"""<p>Hi {author.name},</p>
            <p>We haven't seen you in a while — and that's okay. Life gets busy. But your work is still here, exactly where you left it, waiting for you.</p>
            <p>Whenever you're ready, so are we.</p>""",
    }
    cta_map = {
        'never_started':        (f"{app_url}/author/coaching/quickstart", 'Start My One-Pager →'),
        'stalled_one_pager':    (f"{app_url}/author/coaching/quickstart", 'Pick Up Where I Left Off →'),
        'one_pager_to_full':    (f"{app_url}/author/coaching", 'Start the Full Program →'),
        'stalled_full_proposal':(f"{app_url}/author/coaching", 'Continue My Proposal →'),
        'dormant_30_days':      (f"{app_url}/author/coaching", 'Back to My Work →'),
    }
    body = body_map.get(email_type, '')
    cta_url, cta_text = cta_map.get(email_type, (f"{app_url}/author/coaching", 'Back to My Work →'))
    html_content = f"""<html><body style="font-family:Arial,sans-serif;line-height:1.7;color:#333;">
    <div style="max-width:600px;margin:0 auto;padding:24px;">
        {_coaching_email_header()}
        {body}
        <p style="text-align:center;margin:2rem 0;">
            <a href="{cta_url}" style="display:inline-block;padding:14px 28px;background:#2D1B69;color:white;text-decoration:none;border-radius:8px;font-weight:bold;">{cta_text}</a>
        </p>
        <p style="font-size:0.8rem;color:#999;text-align:center;">You're receiving this because you signed up at Write It Great. Reply any time — we read every email.</p>
        {_coaching_email_footer()}
    </div></body></html>"""
    subject = subject_map.get(email_type, 'A note from Write It Great')
    try:
        result = send_email(author.email, subject, html_content)
        if result:
            record = AuthorEngagementEmail(author_id=author.id, email_type=email_type)
            db.session.add(record)
            db.session.commit()
        return result
    except Exception as e:
        print(f"Re-engagement email error: {e}")
        return False


def check_reengagement_emails():
    """Check all authors and send re-engagement emails where triggered.
    Called from a background thread once per hour."""
    with app.app_context():
        try:
            now = datetime.utcnow()
            authors = Author.query.filter_by(pending_setup=False).all()
            for author in authors:
                # Hard cap: max 3 automated emails total
                sent_count = author.engagement_emails.count()
                if sent_count >= 3:
                    continue
                # Never send more than one email per 7-day window
                last = author.engagement_emails.order_by(
                    AuthorEngagementEmail.sent_at.desc()).first()
                if last and (now - last.sent_at).days < 7:
                    continue

                enrollment = CoachingEnrollment.query.filter_by(
                    author_id=author.id, status='active').first()
                one_pager = author.one_pager_submissions.order_by(
                    OnePagerSubmission.created_at.desc()).first()
                days_since_login = (
                    (now - author.last_login_at).days if author.last_login_at
                    else (now - author.created_at).days
                )

                # Stop if author has submitted for review
                if one_pager and one_pager.status == 'submitted':
                    continue

                email_type = None
                completed_count = 0

                if days_since_login >= 30:
                    email_type = 'dormant_30_days'
                elif not enrollment and not one_pager:
                    # Never started — account created but no activity
                    days_since_created = (now - author.created_at).days
                    if days_since_created >= 7:
                        email_type = 'never_started'
                elif one_pager and not enrollment:
                    # Has one-pager; check if stalled or done
                    if one_pager.summary_text and (now - one_pager.created_at).days >= 14:
                        email_type = 'one_pager_to_full'
                    elif not one_pager.summary_text and (now - one_pager.created_at).days >= 7:
                        email_type = 'stalled_one_pager'
                elif enrollment:
                    all_mp = list(enrollment.module_progress.all())
                    completed_count = sum(1 for mp in all_mp if mp.status == 'approved')
                    last_hw = HomeworkSubmission.query.filter_by(
                        enrollment_id=enrollment.id
                    ).order_by(HomeworkSubmission.submitted_at.desc()).first()
                    last_chat = CoachingChatMessage.query.filter_by(
                        enrollment_id=enrollment.id, role='user'
                    ).order_by(CoachingChatMessage.created_at.desc()).first()
                    dates = [d for d in [
                        last_hw.submitted_at if last_hw else None,
                        last_chat.created_at if last_chat else None,
                        enrollment.enrolled_at,
                    ] if d]
                    last_activity = max(dates) if dates else enrollment.enrolled_at
                    days_stalled = (now - last_activity).days
                    if days_stalled >= 7:
                        email_type = 'stalled_full_proposal'

                if email_type:
                    types_already_sent = {e.email_type for e in author.engagement_emails.all()}
                    if email_type not in types_already_sent:
                        send_reengagement_email(author, email_type, completed_count=completed_count)
        except Exception as e:
            print(f"Re-engagement check error: {e}")


def _start_reengagement_thread():
    """Background thread that checks re-engagement emails every hour."""
    import time
    def _loop():
        time.sleep(300)  # wait 5 min after startup before first run
        while True:
            check_reengagement_emails()
            time.sleep(3600)  # run every hour
    t = threading.Thread(target=_loop, daemon=True)
    t.start()


# ============================================================================
# PUBLIC ROUTES
# ============================================================================

@app.route('/')
def index():
    """Main submission form — requires login (author or team)"""
    if not current_user.is_authenticated:
        return redirect(url_for('author_login'))
    return render_template('index.html')


@app.route('/coach')
def coach():
    """Conversational AI coaching agent for authors developing a proposal."""
    if not current_user.is_authenticated:
        return redirect(url_for('author_login'))
    return render_template('coach.html')


COACH_SYSTEM_PROMPT = """You are an expert book proposal coach at Write It Great — an elite literary agency and ghostwriting consultancy that has placed books with major publishers worldwide.

Your role is to help aspiring nonfiction authors develop a compelling, publisher-ready book proposal through warm, focused conversation. Think of yourself as the best mentor the author has ever had: deeply knowledgeable, honest, encouraging, and specific.

YOUR COACHING APPROACH:
- Start by warmly welcoming the author and asking one open-ended question about their book idea
- Build the proposal section by section through natural conversation — do not present a list of steps upfront
- Ask ONE focused question at a time. Never ask two questions at once
- After the author responds, give 2-3 sentences of specific, expert feedback on what they said, then ask your next question
- When something is strong, say so clearly: "That's a strong hook — an editor would keep reading." Specificity in praise is as important as specificity in critique
- When something needs work, be honest and constructive: explain WHY it's a problem and HOW to fix it with a concrete example
- If the author seems stuck, offer two concrete options or a before/after example to unblock them
- Use the author's own words when reflecting back what you heard — it shows you're listening

THE 7 SECTIONS TO DEVELOP (guide through these naturally, not as a checklist):
1. Book Hook — the central idea in 1-2 compelling sentences (clarity, marketability, "why now")
2. Target Audience — specific ideal reader profile (age, situation, mindset — not "everyone")
3. Author Credentials — why they are the right person (platform numbers, expertise, access, authority)
4. Comparative Titles — 3-5 books published in the last 5 years that share their audience (realistic benchmarks, not aspirational outliers — if their platform is small, comps should reflect modest sales)
5. Chapter Outline — structure and logical arc (each chapter's purpose and contribution)
6. Sample Writing — 500-1,000 words of their best prose (voice, quality, fit for the reader)
7. Marketing & Platform — specific activities + real numbers (email list, social followers, speaking, bulk orders)

PUBLISHING INDUSTRY KNOWLEDGE TO APPLY:
- Nonfiction publishers buy the author's platform first, the book idea second — platform is the #1 factor
- A strong hook can open doors a weak platform can't; a weak hook closes doors a strong platform can't open
- Comp titles must be realistic: an author with 2,000 followers citing books that sold 500,000 copies will be laughed out of a pitch meeting — flag this diplomatically but clearly
- Editors read the marketing section last but it can override a positive impression from everything else
- Chapter outlines should show narrative arc, not just a table of contents — each chapter should build on the last
- Sample writing quality is non-negotiable: a weak sample kills proposals that are strong everywhere else

CONVERSATION STYLE:
- Warm, specific, expert, encouraging — like a trusted mentor, not a chatbot
- Keep responses to 3-4 short paragraphs maximum — this is a conversation, not a lecture
- Use concrete examples when explaining what "good" looks like (e.g., cite real or plausible book titles, real hooks)
- Never give generic advice ("be more specific", "write better") — always say HOW and WHY
- Acknowledge what the author said before responding to it
- Celebrate genuine strengths out loud — authors need to know what to protect

FINAL STEP — WHEN ALL SECTIONS ARE SOLID:
When you have covered all 7 sections and the author's answers are publisher-ready (or close to it), tell them clearly: summarise what they've built, highlight the 2-3 strongest elements, and suggest they submit for a full evaluation using the button below the chat. Do not suggest submission prematurely — only when you genuinely believe the material is ready."""


@app.route('/api/coach-chat', methods=['POST'])
def api_coach_chat():
    """Conversational coaching agent — accepts full message history, returns next agent turn."""
    try:
        data     = request.get_json(force=True) or {}
        messages = data.get('messages', [])
        book_title = data.get('book_title', '').strip()

        if not isinstance(messages, list) or not messages:
            return jsonify({'success': False, 'error': 'No messages provided.'})

        system = COACH_SYSTEM_PROMPT
        if book_title:
            system += f'\n\nThe author\'s working book title is: "{book_title}"'

        # Trim history to last 30 messages to keep context manageable
        history = [m for m in messages if isinstance(m, dict) and m.get('role') in ('user', 'assistant')]
        history = history[-30:]

        response = client.chat.completions.create(
            model='gpt-4o',
            messages=[{'role': 'system', 'content': system}] + history,
            temperature=0.8,
            max_tokens=700
        )

        reply = response.choices[0].message.content.strip()
        return jsonify({'success': True, 'message': reply})

    except Exception as e:
        print(f'/api/coach-chat error: {e}')
        return jsonify({'success': False, 'error': 'Could not get a response. Please try again.'})


@app.route('/api/coach-feedback', methods=['POST'])
def api_coach_feedback():
    """Return AI bullet-point feedback for a single coaching section."""
    try:
        data = request.get_json(force=True) or {}
        section    = data.get('section', '').strip()
        text       = data.get('text', '').strip()
        book_title = data.get('book_title', '').strip()
        book_hook  = data.get('book_hook', '').strip()

        if not text or len(text) < 10:
            return jsonify({'success': False, 'error': 'Please write something before requesting feedback.'})

        section_guides = {
            'hook': (
                'Book Hook',
                'Is the central idea crystal clear in 1-2 sentences? Does it immediately tell an editor what the book is about, who it\'s for, and why it matters now? Is it compelling and specific enough to stand out in a crowded market?'
            ),
            'audience': (
                'Target Audience',
                'Is the ideal reader described with real specificity (age, profession, situation, mindset)? Does it avoid vague phrases like "anyone who..." or "everyone"? Does it demonstrate understanding of an actual buying audience with clear market size?'
            ),
            'credentials': (
                'Author Credentials',
                'Does this establish clear authority to write this book? Does it mention measurable platform (followers, email list, speaking history, media appearances)? Does it explain why this author — and not someone else — should write this book?'
            ),
            'comps': (
                'Comparative Titles',
                'Are the comp titles published within the last 5 years? Are they in the same genre targeting the same reader? Do the cited comp sales figures align with the author\'s current platform size? Would an editor see these as realistic benchmarks, not aspirational outliers?'
            ),
            'outline': (
                'Chapter Outline',
                'Does the structure tell a logical story from beginning to end? Is each chapter\'s purpose and contribution to the arc clear? Does the sequence build momentum and deliver on the hook\'s promise?'
            ),
            'writing': (
                'Sample Writing',
                'Does the voice clearly match the intended reader? Is the quality strong enough that an editor would keep reading past the first paragraph? Is the sample in the 500-1,000 word range expected for a book proposal?'
            ),
            'marketing': (
                'Marketing & Platform',
                'Does the plan describe specific, concrete activities — not just "social media" or "word of mouth"? Are platform numbers cited? Is there a credible path to the first 1,000 buyers? Does the author\'s reach align with what their proposed comp titles actually sold?'
            ),
        }

        label, guide = section_guides.get(section, ('Section', 'Evaluate this section of the book proposal.'))

        context_parts = []
        if book_title:
            context_parts.append(f'Book: "{book_title}"')
        if book_hook and section != 'hook':
            context_parts.append(f'Hook: "{book_hook}"')
        context_line = ('Context — ' + ', '.join(context_parts) + '.\n') if context_parts else ''

        prompt = f"""You are a warm, experienced literary agent reviewing a section of a nonfiction book proposal.
{context_line}
The author has written the following for their "{label}" section:
---
{text[:3000]}
---

Criteria to evaluate against: {guide}

Give exactly 3-5 specific, actionable bullet points. Rules:
- The FIRST bullet must be a genuine strength — start it with **Strong:**
- Remaining bullets use: **Strong:**, **Improve:**, **Add:**, or **Watch out:** as appropriate
- Every bullet must reference something specific from what the author actually wrote
- Be encouraging and direct — tell them exactly what to keep, what to add, what to sharpen
- No generic advice — if it could apply to any proposal, rewrite it to be specific

Return a JSON object with a single key "bullets" containing an array of strings.
Example: {{"bullets": ["**Strong:** Your opening sentence immediately names a specific reader pain point...", "**Improve:** The credentials paragraph lists titles but doesn't explain why each qualifies you for THIS book..."]}}"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.7,
            max_tokens=600
        )

        result = json.loads(response.choices[0].message.content)
        bullets = result.get('bullets', [])
        if not isinstance(bullets, list):
            bullets = []
        return jsonify({'success': True, 'bullets': bullets[:5]})

    except Exception as e:
        print(f"/api/coach-feedback error: {e}")
        return jsonify({'success': False, 'error': 'Could not generate feedback. Please try again.'})


# ============================================================================
# COACHING PLATFORM ROUTES — AUTHOR
# ============================================================================

def _get_module_info(order):
    """Return module dict from COACHING_MODULES by order number (1-based)"""
    for m in COACHING_MODULES:
        if m['order'] == order:
            return m
    return None


def _get_or_create_module_progress(enrollment_id, module_order):
    """Return AuthorModuleProgress row, creating it if missing"""
    mp = AuthorModuleProgress.query.filter_by(
        enrollment_id=enrollment_id, module_order=module_order).first()
    if not mp:
        mp = AuthorModuleProgress(enrollment_id=enrollment_id, module_order=module_order)
        db.session.add(mp)
        db.session.flush()
    return mp


def _get_kb_context(module_order):
    """Return relevant knowledge base content to append to the system prompt."""
    try:
        docs = KnowledgeBaseDocument.query.filter(
            db.or_(
                KnowledgeBaseDocument.module_order == module_order,
                KnowledgeBaseDocument.module_order == None
            )
        ).order_by(KnowledgeBaseDocument.uploaded_at.desc()).limit(3).all()
        if not docs:
            return ''
        parts = ['\n\nKNOWLEDGE BASE — REFERENCE MATERIALS FOR THIS MODULE:']
        for doc in docs:
            if doc.content_text:
                snippet = doc.content_text[:1500].strip()
                label = f'[{doc.doc_type.upper()}] {doc.title}'
                parts.append(f'\n{label}:\n{snippet}')
        return '\n'.join(parts)
    except Exception:
        return ''


def _build_module_system_prompt(module_info, author_name, book_title):
    """Build a module-specific coaching system prompt"""
    order = module_info['order']

    # Only surface the working title in Module 1 (Hook & Concept).
    # In all other modules, referencing the title by name distracts from the
    # actual content being developed and can lead to off-topic conversation.
    if order == 1 and book_title:
        title_line = (
            f'The author\'s working book title is: "{book_title}". '
            f'Treat this as a working title — a starting point, not a fixed identity.\n\n'
        )
    else:
        title_line = ''

    # Module-specific guardrails injected before the shared coaching style block
    module_extra = ''
    if order == 2:
        module_extra = (
            '\nIMPORTANT — TARGET READER: Help the author describe the *type* of reader '
            'using real demographic and psychographic language (occupation, life stage, '
            'specific problem, motivation to buy). Do NOT suggest or use fictional named '
            'personas (e.g. "Meet Sarah, 42, a mid-career manager…"). Real book proposals '
            'describe reader types, not invented characters. If the author drifts toward '
            'naming a fictional person, gently steer them back to the category of reader.\n'
        )
    elif order == 3:
        module_extra = (
            '\nIMPORTANT — COMP TITLES: All primary comp titles must be published within '
            'the last 3 years. Titles older than 3 years may only be used as secondary '
            'comps and must be flagged as such. Sales figures are NOT required and should '
            'never be asked for or mentioned — focus entirely on audience overlap, tone, '
            'market positioning, and the specific gap the author\'s book fills. '
            'Use the format: **Title by Author (Year)** — one sentence on relevance, '
            'one sentence on how this book is different.\n'
        )

    return f"""You are an expert book proposal coach at Write It Great — a literary agency and ghostwriting consultancy that has helped authors land deals with major publishers worldwide.

{title_line}You are coaching {author_name} on Section {order} of 7: **{module_info['title']}** — {module_info['subtitle']}.
{module_extra}
WHAT THIS MODULE IS ABOUT:
{module_info['description']}

YOUR COACHING STYLE:
- Warm, encouraging, and specific — be the best mentor this author has ever had
- ALWAYS lead with what is working before addressing anything that needs improvement
- Use short bullet points for feedback, not long paragraphs of critique
- Ask ONE focused question at a time — never two questions at once
- After each author response: give 1-2 sentences of specific, genuine praise for what's strong, then one clear direction for what to strengthen next
- Keep your responses to 3-4 short paragraphs maximum — this is a conversation, not a lecture
- Reflect the author's own words back to them when you can
- If the author seems stuck or unsure, offer two concrete options or a quick before/after example
- Stay focused on {module_info['chat_context']} — do NOT raise topics that belong to other modules
- If the author asks about something outside this module (e.g. credentials, marketing numbers, comp titles when this isn't that module), respond warmly: "That's definitely worth working through — we'll give it proper attention in its own section. For now, let's keep building your {module_info['title']}."

HOMEWORK FOR THIS MODULE:
{module_info['homework_prompt']}

When the author is ready to write their homework, encourage them warmly. They can see the homework area below the chat — you don't need to direct them to it explicitly.{_get_kb_context(order)}"""


def _review_homework_with_ai(module_info, content, author_name, book_title):
    """Use AI to review homework submission.
    Returns (approved: bool, feedback: str, publisher_ready: bool).
    approved is always True so authors can always advance through the programme.
    publisher_ready reflects the AI's honest assessment of submission quality.
    """
    title_line = f'Book title: "{book_title}"\n' if book_title else ''
    prompt = f"""You are a warm, encouraging literary agent reviewing a homework submission for a book proposal coaching program.

Module {module_info['order']}: {module_info['title']}
{title_line}Author: {author_name}

HOMEWORK PROMPT GIVEN TO THE AUTHOR:
{module_info['homework_prompt']}

AUTHOR'S SUBMISSION:
---
{content[:30000]}
---

Review this submission with a constructive, encouraging eye. Your feedback must:
1. LEAD with what is working — be specific and genuine, not generic praise
2. Identify 1-2 specific areas to strengthen with clear direction
3. Be actionable — tell the author exactly what to add, change, or sharpen

Return a JSON object:
{{
    "publisher_ready": true or false,
    "feedback_bullets": ["bullet 1", "bullet 2", "bullet 3"],
    "word_count_adequate": true or false
}}

Rules:
- publisher_ready = true if an editor would find this section credible and compelling as written
- feedback_bullets = 3-5 concise bullets. The FIRST 1-2 bullets MUST begin with "✓ " and highlight what's genuinely strong. Remaining bullets begin with "→ " and give specific improvement direction. Reference what the author actually wrote — no generic advice.
- word_count_adequate = true if the submission is substantive enough for a real proposal section
- Tone: honest but warm — this author is doing the hard work of writing a real book"""

    try:
        response = client.chat.completions.create(
            model='gpt-4o-mini',
            messages=[{'role': 'user', 'content': prompt}],
            response_format={'type': 'json_object'},
            temperature=0.4,
            max_tokens=800
        )
        result = json.loads(response.choices[0].message.content)
        publisher_ready = bool(result.get('publisher_ready', False))
        bullets = result.get('feedback_bullets', [])
        if not bullets:
            bullets = ['Review complete. Keep refining and resubmit when ready.']
        # Always advance — authors can always move to the next module.
        # publisher_ready signals quality without being a gate.
        return True, bullets, publisher_ready
    except Exception as e:
        print(f"AI homework review error: {e}")
        return True, ['We encountered an issue reviewing your submission. Please try again.'], False


@app.route('/author/coaching')
def author_coaching_dashboard():
    """Coaching program dashboard — shows enrollment status and module progress"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))

    enrollment = CoachingEnrollment.query.filter_by(
        author_id=current_user.id, status='active').first()

    # If enrolled, check for homework reminders (fire if module unlocked 3+ days, no submission)
    if enrollment:
        current_mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment.id, module_order=enrollment.current_module).first()
        if current_mp and current_mp.status == 'in_progress' and current_mp.unlocked_at:
            days_since_unlock = (datetime.utcnow() - current_mp.unlocked_at).days
            last_reminder = current_mp.homework_reminder_sent_at
            reminder_overdue = (not last_reminder or
                                (datetime.utcnow() - last_reminder).days >= 3)
            # Check no approved submission for this module
            approved_hw = HomeworkSubmission.query.filter_by(
                enrollment_id=enrollment.id,
                module_order=enrollment.current_module,
                status='approved').first()
            if days_since_unlock >= 7 and reminder_overdue and not approved_hw:
                module_info = _get_module_info(enrollment.current_module)
                if module_info:
                    try:
                        send_coaching_homework_reminder_email(current_user, module_info, enrollment)
                        current_mp.homework_reminder_sent_at = datetime.utcnow()
                        db.session.commit()
                    except Exception:
                        pass

        all_progress = list(enrollment.module_progress.order_by(
            AuthorModuleProgress.module_order).all())

        # Auto-repair: if section N+1 is approved, section N must also be approved.
        # This fixes data inconsistencies where the previous section was left in
        # 'in_progress' when the next section was unlocked/approved.
        progress_by_order = {mp.module_order: mp for mp in all_progress}
        repaired = False
        for mp in all_progress:
            if mp.status == 'approved' and mp.module_order > 1:
                prev = progress_by_order.get(mp.module_order - 1)
                if prev and prev.status not in ('approved',):
                    prev.status = 'approved'
                    prev.completed_at = prev.completed_at or datetime.utcnow()
                    repaired = True
        if repaired:
            db.session.commit()

        completed_count = sum(1 for mp in all_progress if mp.status == 'approved')
        # Saved draft content per module
        saved_content = {}
        for mp in all_progress:
            mc = CoachingModuleContent.query.filter_by(
                enrollment_id=enrollment.id,
                module_order=mp.module_order
            ).first()
            saved_content[mp.module_order] = mc
        has_any_content = any(mc and mc.content for mc in saved_content.values())
    else:
        all_progress = []
        completed_count = 0
        saved_content = {}
        has_any_content = False

    return render_template(
        'author_coaching_dashboard.html',
        enrollment=enrollment,
        modules=COACHING_MODULES,
        all_progress=all_progress,
        completed_count=completed_count,
        saved_content=saved_content,
        has_any_content=has_any_content,
        total_modules=len(COACHING_MODULES),
    )


@app.route('/author/coaching/enroll', methods=['GET', 'POST'])
def author_coaching_enroll():
    """Enroll author in the coaching program"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))

    # The DB has a unique constraint on author_id so there can only be one
    # enrollment per author. Check for any existing enrollment regardless of status.
    existing = CoachingEnrollment.query.filter_by(author_id=current_user.id).first()
    if existing and existing.status == 'active':
        return redirect(url_for('author_coaching_dashboard'))

    if request.method == 'POST':
        book_title = request.form.get('book_title', '').strip()
        try:
            if existing:
                # Re-enrollment: reset the existing record in place (unique constraint
                # prevents creating a second row, so we update the existing one).
                existing.book_title = book_title or existing.book_title
                existing.status = 'active'
                existing.current_module = 1
                existing.completed_at = None
                existing.welcome_email_sent = False
                existing.complete_email_sent = False
                enrollment = existing
                # Remove stale progress rows so fresh ones are created below
                AuthorModuleProgress.query.filter_by(enrollment_id=existing.id).delete()
                db.session.flush()
            else:
                enrollment = CoachingEnrollment(
                    author_id=current_user.id,
                    book_title=book_title or None,
                    status='active',
                    current_module=1,
                )
                db.session.add(enrollment)
                db.session.flush()

            # Create module progress rows: all modules unlocked from the start
            now = datetime.utcnow()
            for m in COACHING_MODULES:
                mp = AuthorModuleProgress(
                    enrollment_id=enrollment.id,
                    module_order=m['order'],
                    status='in_progress',
                    unlocked_at=now,
                )
                db.session.add(mp)

            db.session.commit()
        except Exception as e:
            db.session.rollback()
            import traceback
            print(f"Enrollment error: {e}\n{traceback.format_exc()}")
            flash('Something went wrong creating your enrollment. Please try again.', 'error')
            return render_template('author_coaching_enroll.html')

        # Send welcome email
        try:
            if send_coaching_welcome_email(current_user, enrollment):
                enrollment.welcome_email_sent = True
                db.session.commit()
        except Exception:
            pass

        flash('Welcome to the coaching program! Module 1 is ready for you.', 'success')
        return redirect(url_for('author_coaching_onboarding'))

    return render_template('author_coaching_enroll.html')


@app.route('/author/coaching/onboarding')
def author_coaching_onboarding():
    """First-time onboarding screen: what a proposal is, how the platform works, button guide"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))
    enrollment = CoachingEnrollment.query.filter_by(
        author_id=current_user.id, status='active').first()
    return render_template('author_coaching_onboarding.html',
                           enrollment=enrollment,
                           modules=COACHING_MODULES)


@app.route('/author/coaching/quickstart', methods=['GET', 'POST'])
def author_coaching_quickstart():
    """Quick One-Pager Mode — 5-question entry point for early-stage authors."""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))

    # Load most recent draft submission for this author (if any)
    existing = current_user.one_pager_submissions.order_by(
        OnePagerSubmission.created_at.desc()).first()

    if request.method == 'POST':
        answers = {
            'problem':   request.form.get('problem', '').strip(),
            'reader':    request.form.get('reader', '').strip(),
            'different': request.form.get('different', '').strip(),
            'why_you':   request.form.get('why_you', '').strip(),
            'marketing': request.form.get('marketing', '').strip(),
            'book_title':request.form.get('book_title', '').strip(),
        }
        if not all([answers['problem'], answers['reader'], answers['why_you']]):
            flash('Please fill in at least questions 1, 2, and 4.', 'error')
            return render_template('author_coaching_quickstart.html', answers=answers,
                                   submission=existing)

        try:
            prompt = f"""You are an expert book proposal coach at Write It Great. An author has answered 5 focused questions about their nonfiction book. Generate a clean, compelling one-page proposal summary.

AUTHOR'S ANSWERS:
Working title: {answers['book_title'] or 'Not yet decided'}
1. What problem does your book solve? {answers['problem']}
2. Who is your target reader? {answers['reader']}
3. Why is your book different from what's already out there? {answers['different'] or 'Not provided'}
4. Why are you the right person to write it? {answers['why_you']}
5. How do you plan to market it? {answers['marketing'] or 'Not provided'}

Write a one-page proposal summary with these sections:
- **The Problem & Promise** (2-3 sentences: the gap this book fills, why it matters now)
- **The Reader** (1-2 sentences: specific target audience — demographic and psychographic)
- **What Makes This Book Different** (2-3 sentences: unique angle, methodology, or perspective)
- **The Author** (2-3 sentences: credentials and unique position to write this)
- **Marketing Potential** (1-2 sentences: platform, reach, and promotional opportunities)
- **Next Steps** (2-3 bullet points: what to develop further before a full submission)

Tone: professional, warm, and specific. Use the author's actual words and voice. Do not pad — be crisp."""

            response = client.chat.completions.create(
                model='gpt-4o-mini',
                messages=[{'role': 'user', 'content': prompt}],
                temperature=0.7,
                max_tokens=950,
            )
            summary = response.choices[0].message.content.strip()
        except Exception as e:
            print(f'Quickstart AI error: {e}')
            summary = None

        # Save / update the submission record
        try:
            if existing and existing.status == 'draft':
                submission = existing
            else:
                submission = OnePagerSubmission(author_id=current_user.id)
                db.session.add(submission)
            submission.book_title = answers['book_title'] or None
            submission.answers_json = json.dumps(answers)
            submission.summary_text = summary
            db.session.commit()
        except Exception as e:
            print(f'One-pager save error: {e}')
            submission = None

        return render_template('author_coaching_quickstart.html',
                               answers=answers,
                               summary=summary,
                               submission=submission)

    # GET: pre-fill from existing draft
    if existing and existing.answers_json:
        prefill = json.loads(existing.answers_json)
    else:
        prefill = {}
    return render_template('author_coaching_quickstart.html', answers=prefill,
                           summary=existing.summary_text if existing else None,
                           submission=existing)


@app.route('/author/coaching/quickstart/submit', methods=['POST'])
def author_quickstart_submit():
    """Author submits their one-pager to the WIG team for review."""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))
    submission_id = request.form.get('submission_id', type=int)
    submission = OnePagerSubmission.query.filter_by(
        id=submission_id, author_id=current_user.id).first()
    if not submission or not submission.summary_text:
        flash('No one-pager found to submit. Please generate it first.', 'error')
        return redirect(url_for('author_coaching_quickstart'))
    if submission.status == 'submitted':
        flash('Your one-pager has already been submitted.', 'info')
        return redirect(url_for('author_coaching_quickstart'))
    submission.status = 'submitted'
    submission.submitted_at = datetime.utcnow()
    db.session.commit()
    try:
        send_one_pager_submitted_notification(current_user, submission)
    except Exception:
        pass
    flash('Your one-pager has been sent to the Write It Great team. We\'ll be in touch!', 'success')
    return redirect(url_for('author_coaching_quickstart'))


@app.route('/author/coaching/module/<int:module_order>')
def author_coaching_module(module_order):
    """Individual coaching module page — chat + homework"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))

    enrollment = CoachingEnrollment.query.filter_by(
        author_id=current_user.id, status='active').first()
    if not enrollment:
        return redirect(url_for('author_coaching_dashboard'))

    module_info = _get_module_info(module_order)
    if not module_info:
        flash('Module not found.', 'error')
        return redirect(url_for('author_coaching_dashboard'))

    mp = _get_or_create_module_progress(enrollment.id, module_order)
    if not mp:
        flash('Module not found.', 'error')
        return redirect(url_for('author_coaching_dashboard'))

    # Load chat history for this module
    chat_messages = CoachingChatMessage.query.filter_by(
        enrollment_id=enrollment.id, module_order=module_order
    ).order_by(CoachingChatMessage.created_at).all()

    # Autosaved draft content
    saved_content = CoachingModuleContent.query.filter_by(
        enrollment_id=enrollment.id, module_order=module_order).first()

    # Latest submission (for feedback display after "Mark Complete")
    latest_submission = HomeworkSubmission.query.filter_by(
        enrollment_id=enrollment.id, module_order=module_order
    ).order_by(HomeworkSubmission.submitted_at.desc()).first()

    # Map module order → section key for feedback API
    section_keys = {1:'hook', 2:'audience', 3:'comps', 4:'bio', 5:'outline', 6:'writing', 7:'marketing'}

    # All module progress (for stepper nav + progress bar)
    all_module_progress = {
        p.module_order: p.status
        for p in AuthorModuleProgress.query.filter_by(enrollment_id=enrollment.id).all()
    }
    approved_count = sum(1 for s in all_module_progress.values() if s == 'approved')

    # Module 1 hook content — used by research agent on any module
    m1_content = CoachingModuleContent.query.filter_by(
        enrollment_id=enrollment.id, module_order=1).first()
    hook_content = (m1_content.content or '') if m1_content else ''

    return render_template(
        'author_coaching_module.html',
        enrollment=enrollment,
        module_info=module_info,
        module_progress=mp,
        chat_messages=chat_messages,
        saved_content=saved_content,
        latest_submission=latest_submission,
        section_key=section_keys.get(module_order, 'hook'),
        total_modules=len(COACHING_MODULES),
        all_module_progress=all_module_progress,
        approved_count=approved_count,
        hook_content=hook_content,
        module_resources=COACHING_MODULE_RESOURCES.get(module_order, {}),
        kb_docs=KnowledgeBaseDocument.query.filter(
            db.or_(
                KnowledgeBaseDocument.module_order == module_order,
                KnowledgeBaseDocument.module_order == None
            )
        ).order_by(KnowledgeBaseDocument.uploaded_at.desc()).all(),
    )


@app.route('/api/coaching/chat', methods=['POST'])
def api_coaching_chat():
    """Module-scoped coaching chat — persists messages server-side"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authenticated.'})

    try:
        data = request.get_json(force=True) or {}
        enrollment_id = data.get('enrollment_id')
        module_order = int(data.get('module_order', 0))
        user_message = (data.get('message') or '').strip()

        if not enrollment_id or not module_order or not user_message:
            return jsonify({'success': False, 'error': 'Missing required fields.'})

        enrollment = CoachingEnrollment.query.filter_by(
            id=enrollment_id, author_id=current_user.id).first()
        if not enrollment:
            return jsonify({'success': False, 'error': 'Enrollment not found.'})

        mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order).first()
        if not mp or mp.status == 'locked':
            return jsonify({'success': False, 'error': 'Module not accessible.'})

        module_info = _get_module_info(module_order)
        if not module_info:
            return jsonify({'success': False, 'error': 'Module not found.'})

        # Persist user message
        user_msg = CoachingChatMessage(
            enrollment_id=enrollment_id,
            module_order=module_order,
            role='user',
            content=user_message,
        )
        db.session.add(user_msg)
        db.session.flush()

        # Build history for API call (last 30 messages)
        history = CoachingChatMessage.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order
        ).order_by(CoachingChatMessage.created_at).all()
        api_messages = [{'role': m.role if m.role == 'user' else 'assistant',
                         'content': m.content} for m in history[-30:]]

        system_prompt = _build_module_system_prompt(
            module_info, current_user.name, enrollment.book_title or '')

        response = client.chat.completions.create(
            model='gpt-4o',
            messages=[{'role': 'system', 'content': system_prompt}] + api_messages,
            temperature=0.8,
            max_tokens=700,
        )
        reply = response.choices[0].message.content.strip()

        # Persist assistant reply
        assist_msg = CoachingChatMessage(
            enrollment_id=enrollment_id,
            module_order=module_order,
            role='assistant',
            content=reply,
        )
        db.session.add(assist_msg)
        db.session.commit()

        return jsonify({'success': True, 'message': reply})

    except Exception as e:
        print(f'/api/coaching/chat error: {e}')
        return jsonify({'success': False, 'error': 'Could not get a response. Please try again.'})


@app.route('/api/coaching/update-title', methods=['POST'])
def api_coaching_update_title():
    """Update the working book title on an enrollment (AJAX, author-only)"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authenticated.'})
    try:
        data = request.get_json(force=True) or {}
        enrollment_id = data.get('enrollment_id')
        new_title = (data.get('book_title') or '').strip()

        if not enrollment_id:
            return jsonify({'success': False, 'error': 'Missing enrollment_id.'})

        enrollment = CoachingEnrollment.query.filter_by(
            id=enrollment_id, author_id=current_user.id).first()
        if not enrollment:
            return jsonify({'success': False, 'error': 'Enrollment not found.'})

        enrollment.book_title = new_title or None
        db.session.commit()
        return jsonify({'success': True, 'book_title': enrollment.book_title or ''})
    except Exception as e:
        print(f'/api/coaching/update-title error: {e}')
        return jsonify({'success': False, 'error': 'Could not update title.'})


@app.route('/api/coaching/content/save', methods=['POST'])
def api_coaching_content_save():
    """Autosave draft content for a coaching module (upsert)"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authenticated.'})
    try:
        data = request.get_json(force=True) or {}
        enrollment_id = data.get('enrollment_id')
        module_order = int(data.get('module_order', 0))
        content = data.get('content', '')

        if not enrollment_id or not module_order:
            return jsonify({'success': False, 'error': 'Missing fields.'})

        enrollment = CoachingEnrollment.query.filter_by(
            id=enrollment_id, author_id=current_user.id).first()
        if not enrollment:
            return jsonify({'success': False, 'error': 'Enrollment not found.'})

        word_count = len(content.split()) if content else 0

        mc = CoachingModuleContent.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order).first()
        if mc:
            mc.content = content
            mc.word_count = word_count
            mc.last_saved_at = datetime.utcnow()
        else:
            mc = CoachingModuleContent(
                enrollment_id=enrollment_id,
                module_order=module_order,
                content=content,
                word_count=word_count,
            )
            db.session.add(mc)
        db.session.commit()
        return jsonify({'success': True, 'word_count': word_count})
    except Exception as e:
        print(f'/api/coaching/content/save error: {e}')
        db.session.rollback()
        return jsonify({'success': False, 'error': 'Save failed.'})


@app.route('/author/coaching/proposal')
def author_coaching_proposal():
    """Compiled draft proposal — all module content in one readable view"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return redirect(url_for('author_login'))

    enrollment = CoachingEnrollment.query.filter_by(
        author_id=current_user.id).order_by(
        CoachingEnrollment.enrolled_at.desc()).first()
    if not enrollment:
        return redirect(url_for('author_coaching_dashboard'))

    sections = []
    for m in COACHING_MODULES:
        mc = CoachingModuleContent.query.filter_by(
            enrollment_id=enrollment.id, module_order=m['order']).first()
        mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment.id, module_order=m['order']).first()
        sections.append({
            'module': m,
            'content': mc.content if mc else '',
            'word_count': mc.word_count if mc else 0,
            'status': mp.status if mp else 'locked',
        })

    total_words = sum(s['word_count'] for s in sections)
    return render_template('author_coaching_proposal.html',
                           enrollment=enrollment,
                           sections=sections,
                           total_words=total_words)


@app.route('/author/coaching/evaluate', methods=['POST'])
@login_required
def author_coaching_evaluate():
    """Submit coaching proposal for full evaluation without requiring a PDF upload.
    Assembles all saved module content, creates a Proposal record, and kicks off
    the background evaluation pipeline. Returns JSON with submission_id."""
    if not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authorised.'})

    enrollment = CoachingEnrollment.query.filter_by(
        author_id=current_user.id).order_by(
        CoachingEnrollment.enrolled_at.desc()).first()
    if not enrollment:
        return jsonify({'success': False, 'error': 'No coaching enrollment found.'})

    # Assemble all module content into a structured proposal text
    lines = [f"BOOK PROPOSAL — {enrollment.book_title or 'Untitled'}\n"]
    for m in COACHING_MODULES:
        mc = CoachingModuleContent.query.filter_by(
            enrollment_id=enrollment.id, module_order=m['order']).first()
        content_text = (mc.content or '').strip() if mc else ''
        lines.append(f"\n{'='*60}")
        lines.append(f"Section {m['order']}: {m['title']}")
        lines.append('='*60)
        lines.append(content_text if content_text else '[Not yet completed]')

    proposal_text = '\n'.join(lines)

    if len(proposal_text.strip()) < 200:
        return jsonify({'success': False, 'error': 'Please complete at least a few sections before submitting for evaluation.'})

    proposal = Proposal(
        submission_id=generate_submission_id(),
        author_id=current_user.id,
        author_name=current_user.name,
        author_email=current_user.email,
        book_title=enrollment.book_title or 'Untitled',
        proposal_type='full',
        ownership_confirmed=True,
        proposal_text=proposal_text[:50000],
        original_filename='coaching_proposal.txt',
        status='processing',
    )
    db.session.add(proposal)
    db.session.commit()

    thread = threading.Thread(
        target=process_evaluation_background,
        args=(app, proposal.submission_id, proposal_text,
              'full', current_user.name, enrollment.book_title or 'Untitled')
    )
    thread.daemon = True
    thread.start()

    return jsonify({
        'success': True,
        'submission_id': proposal.submission_id,
        'status_url': url_for('check_status', submission_id=proposal.submission_id),
        'results_url': url_for('author_proposal_detail', submission_id=proposal.submission_id),
    })


@app.route('/api/coaching/homework', methods=['POST'])
def api_coaching_homework_submit():
    """Submit homework for AI review"""
    if not current_user.is_authenticated or not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authenticated.'})

    try:
        data = request.get_json(force=True) or {}
        enrollment_id = data.get('enrollment_id')
        module_order = int(data.get('module_order', 0))
        content = (data.get('content') or '').strip()

        if not enrollment_id or not module_order or not content:
            return jsonify({'success': False, 'error': 'Missing required fields.'})

        if len(content) < 50:
            return jsonify({'success': False, 'error': 'Please write at least a few sentences.'})

        enrollment = CoachingEnrollment.query.filter_by(
            id=enrollment_id, author_id=current_user.id).first()
        if not enrollment:
            return jsonify({'success': False, 'error': 'Enrollment not found.'})

        mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order).first()
        if not mp or mp.status == 'locked':
            return jsonify({'success': False, 'error': 'Module not accessible.'})

        module_info = _get_module_info(module_order)
        if not module_info:
            return jsonify({'success': False, 'error': 'Module not found.'})

        # Determine revision number
        last_submission = HomeworkSubmission.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order
        ).order_by(HomeworkSubmission.revision_number.desc()).first()
        revision_number = (last_submission.revision_number + 1) if last_submission else 1

        # Create submission record
        submission = HomeworkSubmission(
            enrollment_id=enrollment_id,
            module_order=module_order,
            content=content,
            revision_number=revision_number,
            status='pending_review',
        )
        db.session.add(submission)
        db.session.flush()

        # AI review (synchronous) — approved is always True; publisher_ready is the quality signal
        approved, feedback, publisher_ready = _review_homework_with_ai(
            module_info, content, current_user.name, enrollment.book_title or '')

        # feedback is a list of bullet strings; store as newline-separated for DB
        submission.ai_feedback = '\n'.join(feedback) if isinstance(feedback, list) else feedback
        submission.ai_approved = publisher_ready   # store actual quality assessment
        submission.ai_reviewed_at = datetime.utcnow()
        submission.status = 'approved'             # always advance

        if approved:
            mp.status = 'approved'
            mp.completed_at = datetime.utcnow()

            next_order = module_order + 1
            if next_order <= len(COACHING_MODULES):
                # Unlock next module — create the row if it was somehow deleted
                next_mp = _get_or_create_module_progress(enrollment_id, next_order)
                next_mp.status = 'in_progress'
                next_mp.unlocked_at = datetime.utcnow()
                enrollment.current_module = next_order

            else:
                # All modules complete
                enrollment.status = 'completed'
                enrollment.completed_at = datetime.utcnow()
                try:
                    if send_coaching_complete_email(current_user, enrollment):
                        enrollment.complete_email_sent = True
                except Exception:
                    pass

        db.session.commit()

        return jsonify({
            'success': True,
            'approved': approved,
            'publisher_ready': publisher_ready,
            'feedback_bullets': feedback if isinstance(feedback, list) else [feedback],
            'status': submission.status,
            'next_module': next_order if approved and next_order <= len(COACHING_MODULES) else None,
            'program_complete': approved and module_order == len(COACHING_MODULES),
        })

    except Exception as e:
        print(f'/api/coaching/homework error: {e}')
        db.session.rollback()
        return jsonify({'success': False, 'error': 'Could not submit homework. Please try again.'})


@app.route('/api/coaching/research', methods=['POST'])
@login_required
def api_coaching_research():
    """Research agent: surfaces relevant research areas and viral angles for the author's book concept."""
    data = request.get_json() or {}
    hook = (data.get('hook') or '').strip()
    book_title = (data.get('book_title') or '').strip()
    module_order = int(data.get('module_order', 1))

    if len(hook) < 20:
        return jsonify({'success': False, 'error': 'Write a bit more about your book first so the research agent has something to work with.'})

    title_line = f'Book title: "{book_title}"\n' if book_title else ''
    prompt = f"""You are a publishing industry research agent. An author is developing a non-fiction book proposal and needs research support.

{title_line}Author's current book concept / hook:
"{hook}"

Your job: surface the most relevant, specific, and actionable research to strengthen this book proposal. Respond with a JSON object:

{{
    "research_areas": [
        "3-5 specific academic fields, journals, or bodies of research directly relevant to this topic"
    ],
    "viral_angles": [
        "3-5 proven viral content angles or trending cultural conversations this book can tap into — be specific, not generic"
    ],
    "comparable_titles": [
        "3-5 real published books that occupy nearby market territory — format each as: Title by Author (year) — one sentence on relevance"
    ],
    "market_insight": "2-3 sentences on the current market moment for this topic: why now, who's buying, what gap this book fills",
    "compelling_stats": [
        "2-4 real or widely-cited statistics, studies, or data points that would strengthen a proposal pitch for this book"
    ]
}}

Be specific and credible. Avoid generic advice. Think like a smart literary agent who has done their homework."""

    try:
        response = client.chat.completions.create(
            model='gpt-4o',
            messages=[{'role': 'user', 'content': prompt}],
            response_format={'type': 'json_object'},
            temperature=0.5,
            max_tokens=1000
        )
        result = json.loads(response.choices[0].message.content)
        return jsonify({'success': True, **result})
    except Exception as e:
        print(f"Research agent error: {e}")
        return jsonify({'success': False, 'error': 'Research agent encountered an error. Please try again.'})


@app.route('/api/coaching/save-continue', methods=['POST'])
@login_required
def api_coaching_save_continue():
    """Save current module content and send the author a 'come back' reminder email."""
    if not getattr(current_user, 'is_author', False):
        return jsonify({'success': False, 'error': 'Not authorised.'})

    data = request.get_json() or {}
    enrollment_id = data.get('enrollment_id')
    module_order = int(data.get('module_order', 0))
    content = (data.get('content') or '').strip()

    enrollment = CoachingEnrollment.query.filter_by(
        id=enrollment_id, author_id=current_user.id).first()
    if not enrollment:
        return jsonify({'success': False, 'error': 'Enrollment not found.'})

    # Upsert content
    mc = CoachingModuleContent.query.filter_by(
        enrollment_id=enrollment_id, module_order=module_order).first()
    if mc:
        mc.content = content
        mc.word_count = len(content.split()) if content else 0
        mc.last_saved_at = datetime.utcnow()
    elif content:
        mc = CoachingModuleContent(
            enrollment_id=enrollment_id,
            module_order=module_order,
            content=content,
            word_count=len(content.split()),
            last_saved_at=datetime.utcnow(),
        )
        db.session.add(mc)

    db.session.commit()

    # Send reminder email (uses existing homework reminder function, rate-limited to once per module)
    module_info = _get_module_info(module_order)
    mp = AuthorModuleProgress.query.filter_by(
        enrollment_id=enrollment_id, module_order=module_order).first()
    email_sent = False
    if module_info and mp:
        last_reminder = mp.homework_reminder_sent_at
        if not last_reminder or (datetime.utcnow() - last_reminder).days >= 1:
            try:
                send_coaching_homework_reminder_email(current_user, module_info, enrollment)
                mp.homework_reminder_sent_at = datetime.utcnow()
                db.session.commit()
                email_sent = True
            except Exception:
                pass

    return jsonify({'success': True, 'email_sent': email_sent})


@app.route('/api/evaluate', methods=['POST'])
def api_evaluate():
    """Handle proposal submission and evaluation"""
    try:
        # Use logged-in author's info if available, fall back to form data
        if current_user.is_authenticated and getattr(current_user, 'is_author', False):
            author_name = current_user.name
            author_email = current_user.email
            logged_in_author_id = current_user.id
        else:
            author_name = request.form.get('author_name', '').strip()
            author_email = request.form.get('author_email', '').strip()
            logged_in_author_id = None

        book_title = request.form.get('book_title', '').strip()
        proposal_type = request.form.get('proposal_type', 'full')

        # Extract structured platform numbers (Part A) and marketing strategy (Part B)
        platform_data_raw = request.form.get('platform_data', '').strip()
        marketing_strategy = request.form.get('marketing_strategy', '').strip()
        platform_data_dict = {}
        if platform_data_raw:
            try:
                platform_data_dict = json.loads(platform_data_raw)
            except (json.JSONDecodeError, ValueError):
                platform_data_dict = {}

        if not all([author_name, author_email, book_title]):
            return jsonify({'success': False, 'error': 'Please fill in all required fields.'})

        file = request.files.get('proposal_file')
        if not file or file.filename == '':
            return jsonify({'success': False, 'error': 'Please upload your proposal document.'})

        original_filename = secure_filename(file.filename)
        filename = original_filename.lower()
        if filename.endswith('.pdf'):
            # Read file bytes for storage, then reset for text extraction
            file_bytes = file.read()
            file.seek(0)
            proposal_text = extract_text_from_pdf(file)
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            file_bytes = file.read()
            file.seek(0)
            proposal_text = extract_text_from_docx(file)
        elif filename.endswith('.txt'):
            # Plain-text submissions from the guided coaching builder
            file_bytes = file.read()
            proposal_text = file_bytes.decode('utf-8', errors='ignore')
        else:
            return jsonify({'success': False, 'error': 'Please upload a PDF or Word document.'})

        if len(proposal_text.strip()) < 500:
            return jsonify({'success': False, 'error': 'Could not extract sufficient text from document.'})

        # Save proposal immediately with 'processing' status
        proposal = Proposal(
            submission_id=generate_submission_id(),
            author_id=logged_in_author_id,
            author_name=author_name,
            author_email=author_email,
            book_title=book_title,
            proposal_type=proposal_type,
            ownership_confirmed=True,
            proposal_text=proposal_text[:50000],
            original_filename=original_filename,
            original_file=file_bytes,
            status='processing',
            platform_data=platform_data_raw if platform_data_raw else None,
            marketing_strategy=marketing_strategy if marketing_strategy else None,
        )

        db.session.add(proposal)
        db.session.commit()

        # Run evaluation in background thread to avoid Heroku 30s timeout
        thread = threading.Thread(
            target=process_evaluation_background,
            args=(app, proposal.submission_id, proposal_text, proposal_type, author_name, book_title,
                  platform_data_dict)
        )
        thread.daemon = True
        thread.start()

        return jsonify({
            'success': True,
            'proposal_id': proposal.submission_id
        })

    except Exception as e:
        print(f"Error: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': 'An unexpected error occurred. Please try again.'})


@app.route('/api/status/<submission_id>')
def check_status(submission_id):
    """Check evaluation status (for polling from results page)"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    return jsonify({
        'status': proposal.status,
        'ready': proposal.status not in ('processing',)
    })


# ---------------------------------------------------------------------------
# External Submission API  (Wix integration)
# ---------------------------------------------------------------------------

# Simple in-memory rate limiter: max 10 submissions per IP per hour
_submit_rate = {}

def _check_rate_limit(ip, max_requests=10, window=3600):
    """Return True if the request is within rate limits."""
    import time
    now = time.time()
    hits = _submit_rate.get(ip, [])
    hits = [t for t in hits if now - t < window]
    if len(hits) >= max_requests:
        return False
    hits.append(now)
    _submit_rate[ip] = hits
    return True


def _cors_headers(response):
    """Add CORS headers scoped to the configured origin."""
    origin = request.headers.get('Origin', '')
    allowed = [o.strip() for o in CORS_ORIGIN.split(',') if o.strip()]
    if origin in allowed:
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type, X-API-Key'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
    return response


@app.route('/api/submit', methods=['POST', 'OPTIONS'])
def api_submit():
    """
    External submission endpoint for the Wix site.

    Accepts multipart/form-data with:
        - author_name, author_email, book_title  (required text fields)
        - proposal_type  (optional: 'full', 'marketing_only', 'no_marketing')
        - proposal_file  (required: PDF or DOCX, max 10 MB)

    Requires X-API-Key header matching the API_KEY env var.
    Returns JSON: { success, results_url, status_url }
    """
    # CORS preflight
    if request.method == 'OPTIONS':
        resp = app.make_default_options_response()
        return _cors_headers(resp)

    def _json(data, status=200):
        resp = jsonify(data)
        resp.status_code = status
        return _cors_headers(resp)

    # --- Auth ---
    if not API_KEY:
        return _json({'success': False, 'error': 'API not configured.'}, 503)

    provided_key = request.headers.get('X-API-Key', '')
    if provided_key != API_KEY:
        return _json({'success': False, 'error': 'Invalid or missing API key.'}, 401)

    # --- Rate limit ---
    client_ip = request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip()
    if not _check_rate_limit(client_ip):
        return _json({'success': False, 'error': 'Too many submissions. Please try again later.'}, 429)

    # --- Validate fields ---
    try:
        author_name = request.form.get('author_name', '').strip()
        author_email = request.form.get('author_email', '').strip()
        book_title = request.form.get('book_title', '').strip()
        proposal_type = request.form.get('proposal_type', 'full').strip()

        if not all([author_name, author_email, book_title]):
            return _json({'success': False, 'error': 'Author name, email, and book title are required.'}, 400)

        if proposal_type not in ('full', 'marketing_only', 'no_marketing'):
            proposal_type = 'full'

        # --- Validate file ---
        file = request.files.get('proposal_file')
        if not file or file.filename == '':
            return _json({'success': False, 'error': 'Please upload a proposal document (PDF or DOCX).'}, 400)

        original_filename = secure_filename(file.filename)
        filename_lower = original_filename.lower()

        if not (filename_lower.endswith('.pdf') or filename_lower.endswith('.docx') or filename_lower.endswith('.doc')):
            return _json({'success': False, 'error': 'Only PDF and Word documents are accepted.'}, 400)

        file_bytes = file.read()

        # 10 MB limit for external submissions
        if len(file_bytes) > 10 * 1024 * 1024:
            return _json({'success': False, 'error': 'File size must be under 10 MB.'}, 400)

        file.seek(0)

        # --- Extract text ---
        if filename_lower.endswith('.pdf'):
            proposal_text = extract_text_from_pdf(file)
        else:
            proposal_text = extract_text_from_docx(file)

        if len(proposal_text.strip()) < 500:
            return _json({'success': False, 'error': 'Could not extract enough text from the document. Please check the file.'}, 400)

        # --- Create proposal ---
        proposal = Proposal(
            submission_id=generate_submission_id(),
            author_name=author_name,
            author_email=author_email,
            book_title=book_title,
            proposal_type=proposal_type,
            ownership_confirmed=True,
            proposal_text=proposal_text[:50000],
            original_filename=original_filename,
            original_file=file_bytes,
            status='processing'
        )
        db.session.add(proposal)
        db.session.commit()

        # --- Kick off background evaluation ---
        thread = threading.Thread(
            target=process_evaluation_background,
            args=(app, proposal.submission_id, proposal_text, proposal_type, author_name, book_title)
        )
        thread.daemon = True
        thread.start()

        app_url = APP_BASE_URL
        return _json({
            'success': True,
            'submission_id': proposal.submission_id,
            'results_url': f"{app_url}/results/{proposal.submission_id}",
            'status_url': f"{app_url}/api/status/{proposal.submission_id}",
        })

    except Exception as e:
        print(f"/api/submit error: {e}")
        traceback.print_exc()
        return _json({'success': False, 'error': 'An unexpected error occurred. Please try again.'}, 500)


@app.route('/results/<submission_id>')
def results(submission_id):
    """Public results page for authors"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    processing = proposal.status == 'processing'
    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)
    word_count = len(proposal.proposal_text.split()) if proposal.proposal_text else 0
    return render_template('results.html', proposal=proposal, evaluation=evaluation,
                           processing=processing, word_count=word_count)


@app.route('/download/<submission_id>')
def download_pdf(submission_id):
    """Download PDF report"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()

    try:
        pdf_content = generate_pdf_report(proposal)
        pdf_buffer = BytesIO(pdf_content)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=f"Book_Proposal_Evaluation_{proposal.submission_id}.pdf",
            mimetype='application/pdf'
        )
    except Exception as e:
        print(f"PDF download error: {e}")
        traceback.print_exc()
        flash('Error generating PDF report. Please try again later.', 'error')
        return redirect(url_for('results', submission_id=submission_id))


# ============================================================================
# AUTHOR PORTAL ROUTES
# ============================================================================

@app.route('/author/register', methods=['GET', 'POST'])
def author_register():
    """Author registration"""
    if current_user.is_authenticated:
        if getattr(current_user, 'is_author', False):
            return redirect(url_for('author_dashboard'))
        # Team member visiting author register — log them out first
        logout_user()
        session.pop('user_type', None)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if not all([name, email, password]):
            flash('All fields are required.', 'error')
            return render_template('author_register.html')

        if password != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('author_register.html')

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('author_register.html')

        existing = Author.query.filter_by(email=email).first()
        if existing:
            flash('An account with this email already exists. Please log in.', 'error')
            return redirect(url_for('author_login'))

        author = Author(email=email, name=name)
        author.set_password(password)
        db.session.add(author)
        db.session.commit()

        # Link any existing proposals submitted with this email
        Proposal.query.filter_by(author_email=email, author_id=None).update({'author_id': author.id})
        db.session.commit()

        session['user_type'] = 'author'
        login_user(author)
        flash(f'Welcome, {name}! Your account has been created.', 'success')
        return redirect(url_for('author_dashboard'))

    return render_template('author_register.html')


@app.route('/author/login', methods=['GET', 'POST'])
def author_login():
    """Author login"""
    if current_user.is_authenticated:
        if getattr(current_user, 'is_author', False):
            return redirect(url_for('author_dashboard'))
        # Team member visiting author login — log them out first
        logout_user()
        session.pop('user_type', None)

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')

        author = Author.query.filter_by(email=email).first()
        if author and author.check_password(password):
            session['user_type'] = 'author'
            login_user(author)
            author.last_login_at = datetime.utcnow()
            if author.pending_setup:
                author.pending_setup = False
            db.session.commit()
            next_url = _safe_next(request.args.get('next'))
            if author.assigned_path == 'one_pager' and not next_url:
                return redirect(url_for('author_coaching_quickstart'))
            return redirect(next_url or url_for('author_dashboard'))

        flash('Invalid email or password.', 'error')

    return render_template('author_login.html')


@app.route('/author/logout')
@login_required
def author_logout():
    """Author logout"""
    logout_user()
    session.pop('user_type', None)
    return redirect(url_for('author_login'))


@app.route('/author/dashboard')
@author_login_required
def author_dashboard():
    """Author dashboard showing their proposals"""
    proposals = Proposal.query.filter_by(author_id=current_user.id).order_by(Proposal.submitted_at.desc()).all()
    return render_template('author_dashboard.html',
                         proposals=proposals,
                         status_labels=AUTHOR_STATUS_LABELS)


@app.route('/author/proposal/<submission_id>')
@author_login_required
def author_proposal_detail(submission_id):
    """Author view of a specific proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id, author_id=current_user.id).first_or_404()
    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)
    friendly_status = AUTHOR_STATUS_LABELS.get(proposal.status, proposal.status)
    return render_template('author_proposal.html',
                         proposal=proposal,
                         evaluation=evaluation,
                         friendly_status=friendly_status)


@app.route('/author/forgot-password', methods=['GET', 'POST'])
def author_forgot_password():
    """Author forgot password — sends reset link"""
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        author = Author.query.filter_by(email=email).first()
        if author:
            token = author.generate_reset_token()
            db.session.commit()

            app_url = APP_BASE_URL
            reset_link = f"{app_url}/author/reset-password/{token}"
            html_content = f"""
            <html><body style="font-family: Arial, sans-serif; color: #333;">
                <div style="max-width: 500px; margin: 0 auto; padding: 20px;">
                    <h2 style="color: #2D1B69;">Password Reset</h2>
                    <p>Hi {author.name},</p>
                    <p>Click the link below to reset your password. This link expires in 1 hour.</p>
                    <p><a href="{reset_link}" style="display: inline-block; padding: 12px 24px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">Reset Password</a></p>
                    <p style="color: #999; font-size: 0.875rem;">If you didn't request this, you can safely ignore this email.</p>
                </div>
            </body></html>
            """
            try:
                send_email(email, 'Reset Your Password - Write It Great', html_content)
            except Exception as e:
                print(f"Author password reset email error: {e}")

        # Always show success to avoid email enumeration
        flash('If an account exists with that email, a reset link has been sent.', 'success')
        return redirect(url_for('author_login'))

    return render_template('author_forgot_password.html')


@app.route('/author/reset-password/<token>', methods=['GET', 'POST'])
def author_reset_password(token):
    """Author password reset"""
    author = Author.query.filter_by(password_reset_token=token).first()
    if not author or not author.verify_reset_token(token):
        flash('Invalid or expired reset link.', 'error')
        return redirect(url_for('author_forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if password != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('author_reset_password.html', token=token)

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('author_reset_password.html', token=token)

        author.set_password(password)
        author.password_reset_token = None
        author.password_reset_expires = None
        db.session.commit()
        flash('Password reset successfully. Please log in.', 'success')
        return redirect(url_for('author_login'))

    return render_template('author_reset_password.html', token=token)


# ============================================================================
# PUBLISHER ROUTES
# ============================================================================

@app.route('/publisher/register', methods=['GET', 'POST'])
def publisher_register():
    """Publisher self-registration (pending admin approval)"""
    if current_user.is_authenticated:
        if getattr(current_user, 'is_publisher', False):
            return redirect(url_for('publisher_dashboard'))
        logout_user()
        session.pop('user_type', None)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        company = request.form.get('company', '').strip()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if not all([name, email, password]):
            flash('All fields are required.', 'error')
            return render_template('publisher_register.html')

        if password != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('publisher_register.html')

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('publisher_register.html')

        existing = Publisher.query.filter_by(email=email).first()
        if existing:
            flash('An account with this email already exists.', 'error')
            return redirect(url_for('publisher_login'))

        publisher = Publisher(email=email, name=name, company=company)
        publisher.set_password(password)
        db.session.add(publisher)
        db.session.commit()

        flash('Account created! Your account is pending approval by our team. You will be able to log in once approved.', 'success')
        return redirect(url_for('publisher_login'))

    return render_template('publisher_register.html')


@app.route('/publisher/login', methods=['GET', 'POST'])
def publisher_login():
    """Publisher login (no 2FA)"""
    if current_user.is_authenticated:
        if getattr(current_user, 'is_publisher', False):
            return redirect(url_for('publisher_dashboard'))
        logout_user()
        session.pop('user_type', None)

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')

        publisher = Publisher.query.filter_by(email=email).first()

        if publisher and not publisher.is_active_account:
            flash('This account has been deactivated. Contact our team for assistance.', 'error')
            return render_template('publisher_login.html')

        if publisher and not publisher.is_approved:
            if publisher.check_password(password):
                flash('Your account is pending approval. We will notify you once approved.', 'error')
            else:
                flash('Invalid email or password.', 'error')
            return render_template('publisher_login.html')

        if publisher and publisher.check_password(password):
            session['user_type'] = 'publisher'
            login_user(publisher)
            return redirect(url_for('publisher_dashboard'))

        flash('Invalid email or password.', 'error')

    return render_template('publisher_login.html')


@app.route('/publisher/logout')
@login_required
def publisher_logout():
    """Publisher logout"""
    logout_user()
    session.pop('user_type', None)
    return redirect(url_for('publisher_login'))


@app.route('/publisher/dashboard')
@publisher_login_required
def publisher_dashboard():
    """Publisher dashboard — shows proposals shared with them"""
    shared = PublisherProposal.query.filter_by(publisher_id=current_user.id)\
        .join(Proposal).order_by(PublisherProposal.shared_at.desc()).all()
    return render_template('publisher_dashboard.html',
                         shared_proposals=shared,
                         publisher_status_labels=PUBLISHER_STATUS_LABELS)


@app.route('/publisher/proposal/<submission_id>')
@publisher_login_required
def publisher_proposal_detail(submission_id):
    """Publisher view of a shared proposal — full evaluation"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    # Verify this proposal is shared with the current publisher
    shared = PublisherProposal.query.filter_by(
        publisher_id=current_user.id, proposal_id=proposal.id).first()
    if not shared:
        flash('You do not have access to this proposal.', 'error')
        return redirect(url_for('publisher_dashboard'))

    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)
    return render_template('publisher_proposal.html',
                         proposal=proposal,
                         evaluation=evaluation,
                         shared=shared,
                         publisher_status_options=PUBLISHER_STATUS_OPTIONS)


@app.route('/publisher/proposal/<submission_id>/update-status', methods=['POST'])
@publisher_login_required
def publisher_update_status(submission_id):
    """Publisher updates their status on a shared proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    shared = PublisherProposal.query.filter_by(
        publisher_id=current_user.id, proposal_id=proposal.id).first()
    if not shared:
        flash('You do not have access to this proposal.', 'error')
        return redirect(url_for('publisher_dashboard'))

    new_status = request.form.get('publisher_status', '').strip()
    valid_statuses = {key for key, _ in PUBLISHER_STATUS_OPTIONS}
    if new_status not in valid_statuses:
        flash('Invalid status.', 'error')
        return redirect(url_for('publisher_proposal_detail', submission_id=submission_id))

    old_status = shared.publisher_status
    shared.publisher_status = new_status
    shared.status_updated_at = datetime.utcnow()
    db.session.commit()

    # Log the change in the proposal activity
    note = ProposalNote(
        proposal_id=proposal.id,
        user_name=f'{current_user.name} ({current_user.company or "Publisher"})',
        action='publisher_status_change',
        old_value=PUBLISHER_STATUS_LABELS.get(old_status, old_status),
        new_value=PUBLISHER_STATUS_LABELS.get(new_status, new_status),
        content=f'Publisher status updated'
    )
    db.session.add(note)
    db.session.commit()

    flash(f'Status updated to "{PUBLISHER_STATUS_LABELS.get(new_status, new_status)}".', 'success')
    return redirect(url_for('publisher_proposal_detail', submission_id=submission_id))


@app.route('/publisher/profile', methods=['GET', 'POST'])
@publisher_login_required
def publisher_profile():
    """Publisher profile editing — topics, genres, bio"""
    if request.method == 'POST':
        current_user.name = request.form.get('name', current_user.name).strip()
        current_user.company = request.form.get('company', '').strip() or None
        current_user.bio = request.form.get('bio', '').strip() or None
        current_user.preferred_topics = request.form.get('preferred_topics', '').strip() or None
        current_user.website = request.form.get('website', '').strip() or None

        # Genres come as a multi-select
        selected_genres = request.form.getlist('preferred_genres')
        current_user.preferred_genres = json.dumps(selected_genres) if selected_genres else None

        db.session.commit()
        flash('Profile updated successfully.', 'success')
        return redirect(url_for('publisher_profile'))

    # Parse stored genres for template
    stored_genres = []
    if current_user.preferred_genres:
        try:
            stored_genres = json.loads(current_user.preferred_genres)
        except (json.JSONDecodeError, TypeError):
            stored_genres = []

    return render_template('publisher_profile.html',
                         genre_options=GENRE_OPTIONS,
                         stored_genres=stored_genres)


@app.route('/publisher/forgot-password', methods=['GET', 'POST'])
def publisher_forgot_password():
    """Publisher forgot password"""
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        publisher = Publisher.query.filter_by(email=email).first()
        if publisher:
            token = publisher.generate_reset_token()
            db.session.commit()

            app_url = APP_BASE_URL
            reset_link = f"{app_url}/publisher/reset-password/{token}"
            html_content = f"""
            <html><body style="font-family: Arial, sans-serif; color: #333;">
                <div style="max-width: 500px; margin: 0 auto; padding: 20px;">
                    <h2 style="color: #2D1B69;">Password Reset</h2>
                    <p>Hi {publisher.name},</p>
                    <p>Click the link below to reset your publisher account password. This link expires in 1 hour.</p>
                    <p><a href="{reset_link}" style="display: inline-block; padding: 12px 24px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">Reset Password</a></p>
                    <p style="color: #999; font-size: 0.875rem;">If you didn't request this, you can safely ignore this email.</p>
                </div>
            </body></html>
            """
            try:
                send_email(email, 'Reset Your Password - Write It Great', html_content)
            except Exception as e:
                print(f"Publisher password reset email error: {e}")

        flash('If an account exists with that email, a reset link has been sent.', 'success')
        return redirect(url_for('publisher_login'))

    return render_template('publisher_forgot_password.html')


@app.route('/publisher/reset-password/<token>', methods=['GET', 'POST'])
def publisher_reset_password(token):
    """Publisher password reset"""
    publisher = Publisher.query.filter_by(password_reset_token=token).first()
    if not publisher or not publisher.verify_reset_token(token):
        flash('Invalid or expired reset link.', 'error')
        return redirect(url_for('publisher_forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if password != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('publisher_reset_password.html', token=token)

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('publisher_reset_password.html', token=token)

        publisher.set_password(password)
        publisher.password_reset_token = None
        publisher.password_reset_expires = None
        db.session.commit()
        flash('Password reset successfully. Please log in.', 'success')
        return redirect(url_for('publisher_login'))

    return render_template('publisher_reset_password.html', token=token)


# ============================================================================
# ADMIN ROUTES
# ============================================================================

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    """Team login with 2FA support and lockout protection"""
    if current_user.is_authenticated:
        if getattr(current_user, 'is_team_member', False):
            return redirect(url_for('admin_dashboard'))
        # Author visiting team login — log them out first
        logout_user()
        session.pop('user_type', None)

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')

        user = AdminUser.query.filter_by(email=email).first()

        if user and not user.is_active_account:
            flash('This account has been deactivated. Contact an admin.', 'error')
            return render_template('admin_login.html')

        if user and user.is_locked():
            remaining = int((user.locked_until - datetime.utcnow()).total_seconds() / 60) + 1
            flash(f'Account locked due to too many failed attempts. Try again in {remaining} minutes.', 'error')
            return render_template('admin_login.html')

        if user and user.check_password(password):
            user.record_successful_login()
            db.session.commit()

            # If 2FA is enabled, redirect to verification
            if user.totp_enabled:
                session['pending_2fa_user_id'] = user.id
                return redirect(url_for('admin_verify_2fa'))

            # 2FA not yet enabled — always require setup (mandatory)
            # Reset any stale secret so user gets a fresh QR code
            user.totp_secret = None
            db.session.commit()
            session['setup_2fa_user_id'] = user.id
            return redirect(url_for('admin_setup_2fa'))

        if user:
            user.record_failed_login()
            db.session.commit()

        flash('Invalid email or password.', 'error')

    return render_template('admin_login.html')


@app.route('/admin/logout')
@login_required
def admin_logout():
    """Admin logout"""
    logout_user()
    session.pop('user_type', None)
    return redirect(url_for('admin_login'))


@app.route('/admin')
@team_required
def admin_dashboard():
    """Admin dashboard showing all proposals"""
    page = request.args.get('page', 1, type=int)
    tier_filter = request.args.get('tier', '')
    status_filter = request.args.get('status', '')
    search = request.args.get('search', '')
    view = request.args.get('view', '')  # '' = active, 'archive' = archived

    query = Proposal.query

    # Filter by archive status
    if view == 'archive':
        query = query.filter(Proposal.is_archived == True)
    else:
        query = query.filter(db.or_(Proposal.is_archived == False, Proposal.is_archived == None))

    if tier_filter:
        query = query.filter_by(tier=tier_filter)
    if status_filter:
        query = query.filter_by(status=status_filter)
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            db.or_(
                Proposal.author_name.ilike(search_term),
                Proposal.book_title.ilike(search_term),
                Proposal.author_email.ilike(search_term)
            )
        )

    proposals = query.order_by(Proposal.submitted_at.desc()).paginate(page=page, per_page=20)

    # Stats always reflect active (non-archived) proposals
    active_query = Proposal.query.filter(db.or_(Proposal.is_archived == False, Proposal.is_archived == None))
    stats = {
        'total': active_query.count(),
        'a_tier': active_query.filter_by(tier='A').count(),
        'b_tier': active_query.filter_by(tier='B').count(),
        'c_tier': active_query.filter_by(tier='C').count(),
        'd_tier': active_query.filter_by(tier='D').count(),
        'submitted': active_query.filter_by(status='submitted').count(),
        'shopping': active_query.filter_by(status='shopping').count(),
        'archived': Proposal.query.filter(Proposal.is_archived == True).count(),
    }

    return render_template('admin_dashboard.html',
                         proposals=proposals,
                         stats=stats,
                         status_options=STATUS_OPTIONS,
                         current_filters={'tier': tier_filter, 'status': status_filter, 'search': search, 'view': view})


@app.route('/admin/proposal/<submission_id>', methods=['GET', 'POST'])
@team_required
def admin_proposal_detail(submission_id):
    """View and edit individual proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    
    if request.method == 'POST':
        user_name = current_user.name if current_user.is_authenticated else 'System'

        # Track status changes
        new_status = request.form.get('status')
        if new_status and new_status in [s[0] for s in STATUS_OPTIONS] and new_status != proposal.status:
            status_labels = dict(STATUS_OPTIONS)
            note = ProposalNote(
                proposal_id=proposal.id,
                user_name=user_name,
                action='status_change',
                old_value=status_labels.get(proposal.status, proposal.status),
                new_value=status_labels.get(new_status, new_status),
            )
            db.session.add(note)
            old_status = proposal.status
            proposal.status = new_status

            # Send milestone email to author if applicable
            if new_status in AUTHOR_EMAIL_MILESTONES:
                try:
                    send_author_milestone_email(proposal, new_status)
                except Exception as email_err:
                    print(f"Milestone email error (non-fatal): {email_err}")

        # Add note if provided
        note_text = request.form.get('notes', '').strip()
        if note_text:
            note = ProposalNote(
                proposal_id=proposal.id,
                user_name=user_name,
                action='note',
                content=note_text,
            )
            db.session.add(note)

        db.session.commit()
        flash('Proposal updated successfully', 'success')
        return redirect(url_for('admin_proposal_detail', submission_id=submission_id))
    
    evaluation = json.loads(proposal.evaluation_json) if proposal.evaluation_json else {}
    if evaluation:
        compute_advance_estimate(evaluation)
    activity = proposal.activity_log.order_by(ProposalNote.created_at.desc()).all()

    # Publisher sharing data
    shared_publishers = PublisherProposal.query.filter_by(proposal_id=proposal.id).all()
    available_publishers = Publisher.query.filter_by(is_approved=True, is_active_account=True).order_by(Publisher.name).all()
    # Exclude already-shared publishers
    shared_ids = {sp.publisher_id for sp in shared_publishers}
    available_publishers = [p for p in available_publishers if p.id not in shared_ids]

    return render_template('admin_proposal.html',
                         proposal=proposal,
                         evaluation=evaluation,
                         activity=activity,
                         status_options=STATUS_OPTIONS,
                         shared_publishers=shared_publishers,
                         available_publishers=available_publishers,
                         publisher_status_labels=PUBLISHER_STATUS_LABELS)


@app.route('/admin/proposal/<submission_id>/view-proposal')
@team_required
def view_proposal_text(submission_id):
    """View the original submitted proposal with formatting preserved"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()

    formatted_html = None
    embed_pdf = False
    fname = (proposal.original_filename or '').lower()

    if proposal.original_file and fname.endswith('.pdf'):
        embed_pdf = True
    elif proposal.original_file and fname.endswith('.docx'):
        formatted_html = convert_docx_to_html(proposal.original_file)

    return render_template('admin_view_proposal.html',
                           proposal=proposal,
                           formatted_html=formatted_html,
                           embed_pdf=embed_pdf)


@app.route('/admin/proposal/<submission_id>/embed-proposal')
@team_required
def embed_proposal_file(submission_id):
    """Serve the original file inline for embedding (PDF viewer)"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    if not proposal.original_file or not proposal.original_filename:
        return 'No file available', 404
    file_buffer = BytesIO(proposal.original_file)
    fname = proposal.original_filename.lower()
    if fname.endswith('.pdf'):
        mimetype = 'application/pdf'
    else:
        mimetype = 'application/octet-stream'
    return send_file(file_buffer, mimetype=mimetype, download_name=proposal.original_filename)


@app.route('/admin/proposal/<submission_id>/download-proposal')
@team_required
def download_proposal_text(submission_id):
    """Download the original uploaded file, or extracted text as .txt fallback"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()

    # Serve original file if available
    if proposal.original_file and proposal.original_filename:
        file_buffer = BytesIO(proposal.original_file)
        fname = proposal.original_filename.lower()
        if fname.endswith('.pdf'):
            mimetype = 'application/pdf'
        elif fname.endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif fname.endswith('.doc'):
            mimetype = 'application/msword'
        else:
            mimetype = 'application/octet-stream'
        return send_file(
            file_buffer,
            as_attachment=True,
            download_name=proposal.original_filename,
            mimetype=mimetype
        )

    # Fallback to extracted text
    if not proposal.proposal_text:
        flash('No proposal text available for this submission.', 'error')
        return redirect(url_for('admin_proposal_detail', submission_id=submission_id))

    text_buffer = BytesIO(proposal.proposal_text.encode('utf-8'))
    filename = f"{proposal.author_name.replace(' ', '_')}_{proposal.submission_id}_proposal.txt"
    return send_file(
        text_buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='text/plain'
    )


@app.route('/admin/proposal/<submission_id>/resend-author-email', methods=['POST'])
@team_required
def resend_author_email(submission_id):
    """Resend evaluation email to author"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    
    if send_author_notification(proposal):
        proposal.author_email_sent = True
        db.session.commit()
        flash('Email sent to author successfully', 'success')
    else:
        flash('Failed to send email. Check email configuration.', 'error')
    
    return redirect(url_for('admin_proposal_detail', submission_id=submission_id))


@app.route('/admin/proposal/<submission_id>/delete', methods=['POST'])
@team_required
def admin_delete_proposal(submission_id):
    """Delete a proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    title = proposal.book_title
    db.session.delete(proposal)
    db.session.commit()
    flash(f'Proposal "{title}" has been deleted.', 'success')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/proposal/<submission_id>/archive', methods=['POST'])
@team_required
def admin_toggle_archive(submission_id):
    """Archive or unarchive a proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    proposal.is_archived = not proposal.is_archived
    action_label = 'archived' if proposal.is_archived else 'unarchived'

    user_name = current_user.name if current_user.is_authenticated else 'System'
    note = ProposalNote(
        proposal_id=proposal.id,
        user_name=user_name,
        action='status_change',
        old_value='Archived' if not proposal.is_archived else 'Active',
        new_value='Archived' if proposal.is_archived else 'Active',
    )
    db.session.add(note)
    db.session.commit()
    flash(f'Proposal "{proposal.book_title}" has been {action_label}.', 'success')

    if proposal.is_archived:
        return redirect(url_for('admin_dashboard'))
    return redirect(url_for('admin_proposal_detail', submission_id=submission_id))


@app.route('/admin/proposals/bulk-action', methods=['POST'])
@team_required
def admin_bulk_action():
    """Apply bulk actions to selected proposals"""
    action = request.form.get('bulk_action', '')
    proposal_ids = request.form.getlist('proposal_ids')

    if not proposal_ids:
        flash('No proposals selected.', 'error')
        return redirect(url_for('admin_dashboard'))

    proposals = Proposal.query.filter(Proposal.submission_id.in_(proposal_ids)).all()

    if action == 'delete':
        count = len(proposals)
        for p in proposals:
            db.session.delete(p)
        db.session.commit()
        flash(f'{count} proposal(s) deleted.', 'success')
    elif action == 'archive':
        count = len(proposals)
        for p in proposals:
            p.is_archived = True
        db.session.commit()
        flash(f'{count} proposal(s) archived.', 'success')
    elif action == 'unarchive':
        count = len(proposals)
        for p in proposals:
            p.is_archived = False
        db.session.commit()
        flash(f'{count} proposal(s) restored from archive.', 'success')
    elif action in [s[0] for s in STATUS_OPTIONS]:
        count = len(proposals)
        for p in proposals:
            p.status = action
        db.session.commit()
        flash(f'{count} proposal(s) set to "{action.replace("_", " ").title()}".', 'success')
    else:
        flash('Invalid action.', 'error')

    return redirect(url_for('admin_dashboard'))


@app.route('/admin/proposals/add', methods=['GET', 'POST'])
@team_required
def admin_add_proposal():
    """Manually add a proposal"""
    if request.method == 'POST':
        author_name = request.form.get('author_name', '').strip()
        author_email = request.form.get('author_email', '').strip()
        book_title = request.form.get('book_title', '').strip()
        tier = request.form.get('tier', '').strip()
        score = request.form.get('score', '').strip()
        status = request.form.get('status', 'submitted')
        notes = request.form.get('notes', '').strip()

        if not all([author_name, author_email, book_title]):
            flash('Author name, email, and book title are required.', 'error')
            return render_template('admin_add_proposal.html', status_options=STATUS_OPTIONS)

        proposal = Proposal(
            submission_id=str(uuid.uuid4())[:12],
            author_name=author_name,
            author_email=author_email,
            book_title=book_title,
            tier=tier if tier in ('A', 'B', 'C', 'D') else None,
            overall_score=float(score) if score else None,
            status=status,
            notes=notes,
            proposal_type='full'
        )
        db.session.add(proposal)
        db.session.commit()
        flash(f'Proposal "{book_title}" added successfully.', 'success')
        return redirect(url_for('admin_proposal_detail', submission_id=proposal.submission_id))

    return render_template('admin_add_proposal.html', status_options=STATUS_OPTIONS)


ALLOWED_EMAIL_DOMAIN = 'writeitgreat.com'


def is_valid_team_email(email):
    """Strictly validate that email is a @writeitgreat.com address"""
    import re
    email = (email or '').strip().lower()
    # Must match: localpart@writeitgreat.com (exactly)
    pattern = r'^[a-zA-Z0-9._%+-]+@writeitgreat\.com$'
    return bool(re.match(pattern, email))


@app.route('/admin/register', methods=['GET', 'POST'])
def admin_register():
    """Team registration - strictly restricted to @writeitgreat.com emails"""
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        name = request.form.get('name', '').strip()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if not all([email, name, password, confirm]):
            flash('All fields are required.', 'error')
            return render_template('admin_register.html')

        if not is_valid_team_email(email):
            flash(f'Only @{ALLOWED_EMAIL_DOMAIN} email addresses can register.', 'error')
            return render_template('admin_register.html')

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('admin_register.html')

        if password != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('admin_register.html')

        if AdminUser.query.filter_by(email=email).first():
            flash('An account with this email already exists. Try logging in or resetting your password.', 'error')
            return render_template('admin_register.html')

        user = AdminUser(email=email, name=name, password_hash='temp')
        user.set_password(password)
        db.session.add(user)
        db.session.commit()

        # Don't auto-login — redirect to 2FA setup
        session['setup_2fa_user_id'] = user.id
        return redirect(url_for('admin_setup_2fa'))

    return render_template('admin_register.html')


@app.route('/admin/forgot-password', methods=['GET', 'POST'])
def admin_forgot_password():
    """Request a password reset email"""
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        user = AdminUser.query.filter_by(email=email).first()

        # Always show success message to prevent email enumeration
        flash('If an account exists with that email, a password reset link has been sent.', 'success')

        if user:
            token = user.generate_reset_token()
            db.session.commit()

            app_url = APP_BASE_URL
            reset_url = f"{app_url}/admin/reset-password/{token}"

            reset_html = f"""
            <html>
            <body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
                <div style="max-width: 500px; margin: 0 auto; padding: 20px;">
                    <div style="text-align: center; margin-bottom: 20px;">
                        <h1 style="color: #2D1B69; margin-bottom: 5px;">Write It Great</h1>
                    </div>
                    <p>Hi {user.name},</p>
                    <p>You requested a password reset for your Write It Great dashboard account.</p>
                    <div style="text-align: center; margin: 25px 0;">
                        <a href="{reset_url}" style="display: inline-block; padding: 14px 28px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">Reset Your Password</a>
                    </div>
                    <p style="font-size: 0.875rem; color: #666;">This link expires in 1 hour. If you didn't request this, you can safely ignore this email.</p>
                    <hr style="border: none; border-top: 1px solid #eee; margin: 25px 0;">
                    <p style="font-size: 0.8rem; color: #999;">Write It Great &middot; <a href="https://www.writeitgreat.com" style="color: #2D1B69;">writeitgreat.com</a></p>
                </div>
            </body>
            </html>
            """
            send_email(user.email, 'Password Reset - Write It Great Dashboard', reset_html)

        return redirect(url_for('admin_forgot_password'))

    return render_template('admin_forgot_password.html')


@app.route('/admin/reset-password/<token>', methods=['GET', 'POST'])
def admin_reset_password(token):
    """Reset password using a valid token"""
    if current_user.is_authenticated:
        return redirect(url_for('admin_dashboard'))

    user = AdminUser.query.filter_by(password_reset_token=token).first()
    if not user or not user.verify_reset_token(token):
        flash('This reset link is invalid or has expired. Please request a new one.', 'error')
        return redirect(url_for('admin_forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm = request.form.get('confirm_password', '')

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
        elif password != confirm:
            flash('Passwords do not match.', 'error')
        else:
            user.set_password(password)
            user.password_reset_token = None
            user.password_reset_expires = None
            db.session.commit()
            flash('Your password has been reset. Please log in.', 'success')
            return redirect(url_for('admin_login'))

    return render_template('admin_reset_password.html', token=token)


@app.route('/admin/setup-2fa', methods=['GET', 'POST'])
def admin_setup_2fa():
    """Set up TOTP 2FA after registration or first login"""
    user_id = session.get('setup_2fa_user_id')
    if not user_id:
        return redirect(url_for('admin_login'))

    user = AdminUser.query.get(user_id)
    if not user:
        session.pop('setup_2fa_user_id', None)
        return redirect(url_for('admin_login'))

    # Generate TOTP secret if not already set
    if not user.totp_secret:
        user.setup_totp()
        db.session.commit()

    if request.method == 'POST':
        code = request.form.get('totp_code', '').strip()
        if user.verify_totp(code):
            user.totp_enabled = True
            db.session.commit()
            session.pop('setup_2fa_user_id', None)
            session['user_type'] = 'admin'
            login_user(user)
            flash(f'Welcome, {user.name}! Two-factor authentication is now active.', 'success')
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid code. Please try again with a new code from your authenticator app.', 'error')

    # Generate QR code as base64 image (black on white for maximum scan reliability)
    totp_uri = user.get_totp_uri()
    qr = qrcode.QRCode(version=1, box_size=10, border=4)
    qr.add_data(totp_uri)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color='black', back_color='white')
    qr_buffer = BytesIO()
    qr_img.save(qr_buffer, format='PNG')
    qr_buffer.seek(0)
    qr_b64 = base64.b64encode(qr_buffer.getvalue()).decode('utf-8')

    return render_template('admin_setup_2fa.html',
                           qr_code=qr_b64,
                           totp_secret=user.totp_secret,
                           user_name=user.name)


@app.route('/admin/verify-2fa', methods=['GET', 'POST'])
def admin_verify_2fa():
    """Verify TOTP code during login"""
    user_id = session.get('pending_2fa_user_id')
    if not user_id:
        return redirect(url_for('admin_login'))

    user = AdminUser.query.get(user_id)
    if not user:
        session.pop('pending_2fa_user_id', None)
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        code = request.form.get('totp_code', '').strip()
        if user.verify_totp(code):
            session.pop('pending_2fa_user_id', None)
            session['user_type'] = 'admin'
            login_user(user)
            return redirect(url_for('admin_dashboard'))
        else:
            user.record_failed_login()
            db.session.commit()
            if user.is_locked():
                session.pop('pending_2fa_user_id', None)
                flash('Account locked due to too many failed attempts.', 'error')
                return redirect(url_for('admin_login'))
            flash('Invalid code. Please try again.', 'error')

    return render_template('admin_verify_2fa.html')


# ============================================================================
# TEAM MANAGEMENT (admin only)
# ============================================================================

@app.route('/admin/team')
@admin_required
def admin_team():
    """Team management page - admin only"""
    members = AdminUser.query.order_by(AdminUser.created_at.desc()).all()
    return render_template('admin_team.html', members=members, role_choices=ROLE_CHOICES)


@app.route('/admin/team/<int:user_id>/update-role', methods=['POST'])
@admin_required
def admin_update_role(user_id):
    """Change a team member's role"""
    target = AdminUser.query.get_or_404(user_id)
    new_role = request.form.get('role', ROLE_MEMBER)

    if target.id == current_user.id:
        flash("You can't change your own role.", 'error')
        return redirect(url_for('admin_team'))

    if new_role not in [r[0] for r in ROLE_CHOICES]:
        flash('Invalid role.', 'error')
        return redirect(url_for('admin_team'))

    target.role = new_role
    db.session.commit()
    flash(f'{target.name} is now {new_role}.', 'success')
    return redirect(url_for('admin_team'))


@app.route('/admin/team/<int:user_id>/toggle-active', methods=['POST'])
@admin_required
def admin_toggle_active(user_id):
    """Activate or deactivate a team member"""
    target = AdminUser.query.get_or_404(user_id)

    if target.id == current_user.id:
        flash("You can't deactivate your own account.", 'error')
        return redirect(url_for('admin_team'))

    target.is_active_account = not target.is_active_account
    db.session.commit()
    status = 'activated' if target.is_active_account else 'deactivated'
    flash(f'{target.name} has been {status}.', 'success')
    return redirect(url_for('admin_team'))


@app.route('/admin/team/<int:user_id>/reset-2fa', methods=['POST'])
@admin_required
def admin_reset_2fa(user_id):
    """Reset a team member's 2FA (they'll need to set it up again on next login)"""
    target = AdminUser.query.get_or_404(user_id)
    target.totp_secret = None
    target.totp_enabled = False
    db.session.commit()
    flash(f"2FA has been reset for {target.name}. They'll set it up again on next login.", 'success')
    return redirect(url_for('admin_team'))


@app.route('/admin/team/<int:user_id>/unlock', methods=['POST'])
@admin_required
def admin_unlock_account(user_id):
    """Unlock a locked account"""
    target = AdminUser.query.get_or_404(user_id)
    target.failed_login_attempts = 0
    target.locked_until = None
    db.session.commit()
    flash(f'{target.name} has been unlocked.', 'success')
    return redirect(url_for('admin_team'))


@app.route('/admin/team/<int:user_id>/delete', methods=['POST'])
@admin_required
def admin_delete_member(user_id):
    """Permanently delete a team member account"""
    target = AdminUser.query.get_or_404(user_id)

    if target.id == current_user.id:
        flash("You can't delete your own account.", 'error')
        return redirect(url_for('admin_team'))

    name = target.name
    db.session.delete(target)
    db.session.commit()
    flash(f'{name} has been permanently removed.', 'success')
    return redirect(url_for('admin_team'))


# ============================================================================
# ADMIN PUBLISHER MANAGEMENT
# ============================================================================

@app.route('/admin/publishers')
@team_required
def admin_publishers():
    """Manage publisher accounts — list with filters"""
    # Filters
    search = request.args.get('search', '').strip()
    genre_filter = request.args.get('genre', '').strip()
    status_filter = request.args.get('status', '').strip()  # pending, active, deactivated

    query = Publisher.query

    if search:
        search_term = f'%{search}%'
        query = query.filter(
            db.or_(
                Publisher.name.ilike(search_term),
                Publisher.email.ilike(search_term),
                Publisher.company.ilike(search_term),
            )
        )

    if status_filter == 'pending':
        query = query.filter_by(is_approved=False)
    elif status_filter == 'active':
        query = query.filter_by(is_approved=True, is_active_account=True)
    elif status_filter == 'deactivated':
        query = query.filter_by(is_active_account=False)

    if genre_filter:
        # Match publishers whose preferred_genres JSON contains the genre
        query = query.filter(Publisher.preferred_genres.ilike(f'%{genre_filter}%'))

    publishers = query.order_by(Publisher.created_at.desc()).all()

    # Stats (unfiltered)
    total = Publisher.query.count()
    approved = Publisher.query.filter_by(is_approved=True).count()
    pending = Publisher.query.filter_by(is_approved=False).count()
    active = Publisher.query.filter_by(is_approved=True, is_active_account=True).count()

    return render_template('admin_publishers.html',
                         publishers=publishers,
                         genre_options=GENRE_OPTIONS,
                         publisher_status_labels=PUBLISHER_STATUS_LABELS,
                         stats={'total': total, 'approved': approved, 'pending': pending, 'active': active},
                         current_filters={'search': search, 'genre': genre_filter, 'status': status_filter})


@app.route('/admin/publishers/<int:publisher_id>')
@team_required
def admin_publisher_detail(publisher_id):
    """View a publisher's full profile and shared proposal history"""
    publisher = Publisher.query.get_or_404(publisher_id)

    # Parse genres
    stored_genres = []
    if publisher.preferred_genres:
        try:
            stored_genres = json.loads(publisher.preferred_genres)
        except (json.JSONDecodeError, TypeError):
            stored_genres = []

    # Get shared proposals with their publisher status
    shared = PublisherProposal.query.filter_by(publisher_id=publisher.id)\
        .join(Proposal).order_by(PublisherProposal.shared_at.desc()).all()

    return render_template('admin_publisher_detail.html',
                         publisher=publisher,
                         stored_genres=stored_genres,
                         shared_proposals=shared,
                         publisher_status_labels=PUBLISHER_STATUS_LABELS)


@app.route('/admin/publishers/add', methods=['GET', 'POST'])
@team_required
def admin_add_publisher():
    """Manually create a publisher account (pre-approved)"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        company = request.form.get('company', '').strip() or None
        password = request.form.get('password', '')

        if not name or not email or not password:
            flash('Name, email, and password are required.', 'error')
            return render_template('admin_add_publisher.html', genre_options=GENRE_OPTIONS)

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
            return render_template('admin_add_publisher.html', genre_options=GENRE_OPTIONS)

        existing = Publisher.query.filter_by(email=email).first()
        if existing:
            flash(f'A publisher with email {email} already exists.', 'error')
            return render_template('admin_add_publisher.html', genre_options=GENRE_OPTIONS)

        publisher = Publisher(
            email=email,
            name=name,
            company=company,
            is_approved=True,
            is_active_account=True,
            bio=request.form.get('bio', '').strip() or None,
            preferred_topics=request.form.get('preferred_topics', '').strip() or None,
            website=request.form.get('website', '').strip() or None,
        )
        publisher.set_password(password)

        selected_genres = request.form.getlist('preferred_genres')
        publisher.preferred_genres = json.dumps(selected_genres) if selected_genres else None

        db.session.add(publisher)
        db.session.commit()

        flash(f'{name} has been added as an approved publisher.', 'success')
        return redirect(url_for('admin_publisher_detail', publisher_id=publisher.id))

    return render_template('admin_add_publisher.html', genre_options=GENRE_OPTIONS)


@app.route('/admin/publishers/<int:publisher_id>/approve', methods=['POST'])
@team_required
def admin_approve_publisher(publisher_id):
    """Approve a pending publisher account"""
    publisher = Publisher.query.get_or_404(publisher_id)
    publisher.is_approved = True
    db.session.commit()

    # Send approval notification email
    app_url = APP_BASE_URL
    html_content = f"""
    <html><body style="font-family: Arial, sans-serif; color: #333;">
        <div style="max-width: 500px; margin: 0 auto; padding: 20px;">
            <h2 style="color: #2D1B69;">Account Approved</h2>
            <p>Hi {publisher.name},</p>
            <p>Your publisher account at Write It Great has been approved! You can now log in to view proposals shared with you.</p>
            <p><a href="{app_url}/publisher/login" style="display: inline-block; padding: 12px 24px; background: #B8F2B8; color: #1a3a1a; text-decoration: none; border-radius: 8px; font-weight: bold;">Log In Now</a></p>
        </div>
    </body></html>
    """
    try:
        send_email(publisher.email, 'Your Publisher Account Has Been Approved - Write It Great', html_content)
    except Exception as e:
        print(f"Publisher approval email error: {e}")

    flash(f'{publisher.name} has been approved.', 'success')
    return redirect(request.referrer or url_for('admin_publishers'))


@app.route('/admin/publishers/<int:publisher_id>/toggle-active', methods=['POST'])
@team_required
def admin_toggle_publisher_active(publisher_id):
    """Activate or deactivate a publisher"""
    publisher = Publisher.query.get_or_404(publisher_id)
    publisher.is_active_account = not publisher.is_active_account
    db.session.commit()
    status = 'activated' if publisher.is_active_account else 'deactivated'
    flash(f'{publisher.name} has been {status}.', 'success')
    return redirect(request.referrer or url_for('admin_publishers'))


@app.route('/admin/publishers/<int:publisher_id>/delete', methods=['POST'])
@team_required
def admin_delete_publisher(publisher_id):
    """Delete a publisher account and their shared proposal links"""
    publisher = Publisher.query.get_or_404(publisher_id)
    name = publisher.name
    PublisherProposal.query.filter_by(publisher_id=publisher.id).delete()
    db.session.delete(publisher)
    db.session.commit()
    flash(f'{name} has been permanently removed.', 'success')
    return redirect(url_for('admin_publishers'))


@app.route('/admin/proposal/<submission_id>/share', methods=['POST'])
@team_required
def admin_share_proposal(submission_id):
    """Share a proposal with one or more publishers"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    publisher_ids = request.form.getlist('publisher_ids')

    shared_count = 0
    for pid in publisher_ids:
        existing = PublisherProposal.query.filter_by(
            publisher_id=int(pid), proposal_id=proposal.id).first()
        if not existing:
            sp = PublisherProposal(
                publisher_id=int(pid),
                proposal_id=proposal.id,
                shared_by=current_user.name if current_user.is_authenticated else 'System'
            )
            db.session.add(sp)
            shared_count += 1

    if shared_count > 0:
        db.session.commit()
        # Log the share action
        note = ProposalNote(
            proposal_id=proposal.id,
            user_name=current_user.name if current_user.is_authenticated else 'System',
            action='shared',
            content=f'Shared with {shared_count} publisher(s)'
        )
        db.session.add(note)
        db.session.commit()
        flash(f'Proposal shared with {shared_count} publisher(s).', 'success')
    else:
        flash('Proposal is already shared with the selected publisher(s).', 'error')

    return redirect(url_for('admin_proposal_detail', submission_id=submission_id))


@app.route('/admin/proposal/<submission_id>/unshare/<int:publisher_id>', methods=['POST'])
@team_required
def admin_unshare_proposal(submission_id, publisher_id):
    """Remove a publisher's access to a proposal"""
    proposal = Proposal.query.filter_by(submission_id=submission_id).first_or_404()
    sp = PublisherProposal.query.filter_by(
        publisher_id=publisher_id, proposal_id=proposal.id).first_or_404()
    publisher_name = sp.publisher.name
    db.session.delete(sp)

    note = ProposalNote(
        proposal_id=proposal.id,
        user_name=current_user.name if current_user.is_authenticated else 'System',
        action='unshared',
        content=f'Removed access for {publisher_name}'
    )
    db.session.add(note)
    db.session.commit()
    flash(f'Access removed for {publisher_name}.', 'success')
    return redirect(url_for('admin_proposal_detail', submission_id=submission_id))


# ============================================================================
# COACHING PLATFORM ROUTES — ADMIN
# ============================================================================

@app.route('/admin/coaching')
@team_required
def admin_coaching_list():
    """Admin overview of all coaching enrollments"""
    enrollments = (CoachingEnrollment.query
                   .join(Author)
                   .order_by(CoachingEnrollment.enrolled_at.desc())
                   .all())

    enrollment_data = []
    for enr in enrollments:
        all_mp = list(enr.module_progress.order_by(AuthorModuleProgress.module_order).all())
        completed = sum(1 for mp in all_mp if mp.status == 'approved')
        # Last activity: last chat message or homework submission
        last_chat = (CoachingChatMessage.query
                     .filter_by(enrollment_id=enr.id, role='user')
                     .order_by(CoachingChatMessage.created_at.desc()).first())
        last_hw = (HomeworkSubmission.query
                   .filter_by(enrollment_id=enr.id)
                   .order_by(HomeworkSubmission.submitted_at.desc()).first())
        last_dates = [d for d in [
            last_chat.created_at if last_chat else None,
            last_hw.submitted_at if last_hw else None,
            enr.enrolled_at
        ] if d]
        last_active = max(last_dates) if last_dates else enr.enrolled_at
        enrollment_data.append({
            'enrollment': enr,
            'completed': completed,
            'total': len(COACHING_MODULES),
            'pct': int((completed / len(COACHING_MODULES)) * 100),
            'last_active': last_active,
        })

    return render_template('admin_coaching_list.html',
                           enrollment_data=enrollment_data,
                           modules=COACHING_MODULES)


@app.route('/admin/coaching/<int:enrollment_id>')
@team_required
def admin_coaching_detail(enrollment_id):
    """Admin detailed view of a single author's coaching progress"""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    all_mp = list(enrollment.module_progress.order_by(AuthorModuleProgress.module_order).all())
    # Build per-module data: progress + all submissions + chat history
    module_data = []
    for mp in all_mp:
        m_info = _get_module_info(mp.module_order)
        if not m_info:
            continue
        submissions = (HomeworkSubmission.query
                       .filter_by(enrollment_id=enrollment_id, module_order=mp.module_order)
                       .order_by(HomeworkSubmission.submitted_at.desc()).all())
        chats = (CoachingChatMessage.query
                 .filter_by(enrollment_id=enrollment_id, module_order=mp.module_order)
                 .order_by(CoachingChatMessage.created_at).all())
        module_data.append({
            'module': m_info,
            'progress': mp,
            'submissions': submissions,
            'chats': chats,
        })

    completed_count = sum(1 for mp in all_mp if mp.status == 'approved')
    return render_template('admin_coaching_detail.html',
                           enrollment=enrollment,
                           module_data=module_data,
                           completed_count=completed_count,
                           total_modules=len(COACHING_MODULES))


@app.route('/admin/coaching/<int:enrollment_id>/module/<int:module_order>/unlock', methods=['POST'])
@team_required
def admin_coaching_unlock_module(enrollment_id, module_order):
    """Admin manually unlocks a coaching module for an author"""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    mp = AuthorModuleProgress.query.filter_by(
        enrollment_id=enrollment_id, module_order=module_order).first_or_404()

    # Always auto-approve the previous section when unlocking, regardless of
    # whether this module is already in progress (fixes gap where early-return
    # prevented the previous section from being marked approved).
    if module_order > 1:
        prev_mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment_id, module_order=module_order - 1).first()
        if prev_mp and prev_mp.status not in ('approved',):
            prev_mp.status = 'approved'
            prev_mp.completed_at = prev_mp.completed_at or datetime.utcnow()

    if mp.status not in ('locked', 'revision_requested'):
        # Module already accessible — just ensure prev section is approved (done above)
        db.session.commit()
        flash('Module is already in progress or approved.', 'info')
        return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))

    mp.status = 'in_progress'
    mp.unlocked_at = mp.unlocked_at or datetime.utcnow()

    enrollment.current_module = max(enrollment.current_module, module_order)
    db.session.commit()

    module_info = _get_module_info(module_order)
    author = enrollment.author
    if module_info:
        try:
            send_coaching_module_unlocked_email(author, module_info, enrollment)
        except Exception:
            pass

    flash(f'Module {module_order} unlocked for {author.name}.', 'success')
    return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))


@app.route('/admin/coaching/<int:enrollment_id>/module/<int:module_order>/approve', methods=['POST'])
@team_required
def admin_coaching_approve_module(enrollment_id, module_order):
    """Admin manually marks a module as approved/complete without unlocking the next one"""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    mp = AuthorModuleProgress.query.filter_by(
        enrollment_id=enrollment_id, module_order=module_order).first_or_404()

    if mp.status == 'approved':
        flash('Section is already marked complete.', 'info')
        return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))

    mp.status = 'approved'
    mp.completed_at = mp.completed_at or datetime.utcnow()
    db.session.commit()

    author = enrollment.author
    flash(f'Section {module_order} marked complete for {author.name}.', 'success')
    return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))


@app.route('/admin/coaching/<int:enrollment_id>/homework/<int:submission_id>/review', methods=['POST'])
@team_required
def admin_coaching_review_homework(enrollment_id, submission_id):
    """Admin adds feedback to a homework submission and optionally approves it"""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    submission = HomeworkSubmission.query.filter_by(
        id=submission_id, enrollment_id=enrollment_id).first_or_404()

    admin_feedback = request.form.get('admin_feedback', '').strip()
    action = request.form.get('action', '')  # 'approve' or 'request_revision'

    submission.admin_feedback = admin_feedback
    submission.admin_reviewed_by = current_user.name
    submission.admin_reviewed_at = datetime.utcnow()

    if action == 'approve':
        submission.status = 'approved'
        submission.ai_approved = True
        mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment_id,
            module_order=submission.module_order).first()
        if mp:
            mp.status = 'approved'
            mp.completed_at = mp.completed_at or datetime.utcnow()

        next_order = submission.module_order + 1
        if next_order <= len(COACHING_MODULES):
            next_mp = AuthorModuleProgress.query.filter_by(
                enrollment_id=enrollment_id, module_order=next_order).first()
            if next_mp and next_mp.status == 'locked':
                next_mp.status = 'in_progress'
                next_mp.unlocked_at = datetime.utcnow()
            enrollment.current_module = max(enrollment.current_module, next_order)
            next_module_info = _get_module_info(next_order)
            if next_module_info:
                try:
                    send_coaching_module_unlocked_email(enrollment.author, next_module_info, enrollment)
                except Exception:
                    pass
        elif submission.module_order == len(COACHING_MODULES):
            enrollment.status = 'completed'
            enrollment.completed_at = enrollment.completed_at or datetime.utcnow()
            try:
                send_coaching_complete_email(enrollment.author, enrollment)
            except Exception:
                pass

        flash('Homework approved and next module unlocked.', 'success')

    elif action == 'request_revision':
        submission.status = 'revision_requested'
        submission.ai_approved = False
        mp = AuthorModuleProgress.query.filter_by(
            enrollment_id=enrollment_id,
            module_order=submission.module_order).first()
        if mp:
            mp.status = 'revision_requested'
        flash('Revision requested. Author has been notified via email.', 'success')
        try:
            module_info = _get_module_info(submission.module_order)
            if module_info:
                send_coaching_homework_reviewed_email(enrollment.author, module_info, submission)
        except Exception:
            pass

    db.session.commit()
    return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))


@app.route('/admin/coaching/<int:enrollment_id>/delete-author', methods=['POST'])
@team_required
def admin_coaching_delete_author(enrollment_id):
    """Permanently delete an author account and all associated data."""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    author = enrollment.author
    author_name = author.name
    author_email = author.email

    # Delete all coaching data for every enrollment this author has
    for enr in CoachingEnrollment.query.filter_by(author_id=author.id).all():
        CoachingChatMessage.query.filter_by(enrollment_id=enr.id).delete()
        HomeworkSubmission.query.filter_by(enrollment_id=enr.id).delete()
        CoachingModuleContent.query.filter_by(enrollment_id=enr.id).delete()
        AuthorModuleProgress.query.filter_by(enrollment_id=enr.id).delete()
        db.session.delete(enr)

    # Delete proposals and their child records
    for proposal in Proposal.query.filter_by(author_id=author.id).all():
        ProposalNote.query.filter_by(proposal_id=proposal.id).delete()
        PublisherProposal.query.filter_by(proposal_id=proposal.id).delete()
        db.session.delete(proposal)

    # Delete the author account itself
    db.session.delete(author)
    db.session.commit()

    flash(f'Author account for {author_name} ({author_email}) has been permanently deleted.', 'success')
    return redirect(url_for('admin_coaching_list'))


@app.route('/admin/author/<int:author_id>/delete', methods=['POST'])
@team_required
def admin_delete_author(author_id):
    """Delete an author account by author_id (works even with no enrollment)."""
    author = Author.query.get_or_404(author_id)
    author_name = author.name
    author_email = author.email

    for enr in CoachingEnrollment.query.filter_by(author_id=author.id).all():
        CoachingChatMessage.query.filter_by(enrollment_id=enr.id).delete()
        HomeworkSubmission.query.filter_by(enrollment_id=enr.id).delete()
        CoachingModuleContent.query.filter_by(enrollment_id=enr.id).delete()
        AuthorModuleProgress.query.filter_by(enrollment_id=enr.id).delete()
        db.session.delete(enr)

    for proposal in Proposal.query.filter_by(author_id=author.id).all():
        ProposalNote.query.filter_by(proposal_id=proposal.id).delete()
        PublisherProposal.query.filter_by(proposal_id=proposal.id).delete()
        db.session.delete(proposal)

    OnePagerSubmission.query.filter_by(author_id=author.id).delete()
    AuthorEngagementEmail.query.filter_by(author_id=author.id).delete()

    db.session.delete(author)
    db.session.commit()

    flash(f'Author account for {author_name} ({author_email}) has been permanently deleted.', 'success')
    return redirect(url_for('admin_pipeline'))


@app.route('/admin/coaching/<int:enrollment_id>/reset', methods=['POST'])
@team_required
def admin_coaching_reset_enrollment(enrollment_id):
    """Reset an author's coaching enrollment — wipes all progress, content, chat, homework.
    Use this to let a tester start fresh without deleting the account."""
    enrollment = CoachingEnrollment.query.get_or_404(enrollment_id)
    author_name = enrollment.author.name

    # Delete all child records
    CoachingChatMessage.query.filter_by(enrollment_id=enrollment_id).delete()
    HomeworkSubmission.query.filter_by(enrollment_id=enrollment_id).delete()
    CoachingModuleContent.query.filter_by(enrollment_id=enrollment_id).delete()
    AuthorModuleProgress.query.filter_by(enrollment_id=enrollment_id).delete()

    # Reset enrollment state
    enrollment.current_module = 1
    enrollment.status = 'active'
    enrollment.completed_at = None
    enrollment.complete_email_sent = False
    # welcome_email_sent kept True so the welcome email isn't re-sent

    # Re-create clean module progress rows — all unlocked
    now_dt = datetime.utcnow()
    for m in COACHING_MODULES:
        mp = AuthorModuleProgress(
            enrollment_id=enrollment_id,
            module_order=m['order'],
            status='in_progress',
            unlocked_at=now_dt,
        )
        db.session.add(mp)

    db.session.commit()
    flash(f'Enrollment for {author_name} has been fully reset to Module 1.', 'success')
    return redirect(url_for('admin_coaching_detail', enrollment_id=enrollment_id))


# ============================================================================
# ADMIN — CREATE AUTHOR ACCOUNT
# ============================================================================

@app.route('/admin/authors/add', methods=['GET', 'POST'])
@team_required
def admin_add_author():
    """Admin creates an author account on behalf of a prospect."""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('email', '').strip().lower()
        assigned_path = request.form.get('assigned_path', '') or None

        if not name or not email:
            flash('Name and email are required.', 'error')
            return render_template('admin_add_author.html')

        if Author.query.filter_by(email=email).first():
            flash(f'An account for {email} already exists.', 'error')
            return render_template('admin_add_author.html')

        # Create account with placeholder password; author will set their own
        author = Author(
            email=email,
            name=name,
            pending_setup=True,
            admin_created=True,
            assigned_path=assigned_path,
        )
        author.set_password(uuid.uuid4().hex)   # random unusable password
        token = author.generate_reset_token()
        # Give them 48 hours to set their password
        author.password_reset_expires = datetime.utcnow() + timedelta(hours=48)
        db.session.add(author)
        db.session.commit()

        sent = send_author_welcome_invite_email(author, token)
        if sent:
            flash(f'Account created and invite sent to {email}.', 'success')
        else:
            app_url = APP_BASE_URL
            setup_url = f"{app_url}/author/reset-password/{token}"
            flash(f'Account created for {email}. Email could not be sent — share this link manually: {setup_url}', 'warning')

        return redirect(url_for('admin_pipeline'))

    return render_template('admin_add_author.html')


# ============================================================================
# ADMIN — AUTHOR PIPELINE TRACKER
# ============================================================================

@app.route('/admin/pipeline')
@team_required
def admin_pipeline():
    """Author pipeline tracker — both One-Pager and Full Proposal paths."""
    now = datetime.utcnow()

    # All authors
    authors = Author.query.order_by(Author.created_at.desc()).all()

    rows = []
    for author in authors:
        enrollment = CoachingEnrollment.query.filter_by(
            author_id=author.id, status='active').first()
        one_pager = author.one_pager_submissions.order_by(
            OnePagerSubmission.created_at.desc()).first()

        # Determine path
        if enrollment:
            path = 'Full Proposal'
        elif author.assigned_path == 'one_pager' or one_pager:
            path = 'One-Pager'
        elif author.assigned_path == 'full_proposal':
            path = 'Full Proposal'
        else:
            path = '—'

        # Current step
        if enrollment:
            all_mp = list(enrollment.module_progress.order_by(
                AuthorModuleProgress.module_order).all())
            completed = sum(1 for mp in all_mp if mp.status == 'approved')
            current_step = f'Section {enrollment.current_module} of {len(COACHING_MODULES)}'
        else:
            completed = 0
            current_step = 'One-Pager' if one_pager else '—'

        # Last active date
        dates = [author.created_at]
        if author.last_login_at:
            dates.append(author.last_login_at)
        if enrollment:
            last_chat = CoachingChatMessage.query.filter_by(
                enrollment_id=enrollment.id, role='user'
            ).order_by(CoachingChatMessage.created_at.desc()).first()
            last_hw = HomeworkSubmission.query.filter_by(
                enrollment_id=enrollment.id
            ).order_by(HomeworkSubmission.submitted_at.desc()).first()
            if last_chat:
                dates.append(last_chat.created_at)
            if last_hw:
                dates.append(last_hw.submitted_at)
        if one_pager:
            dates.append(one_pager.created_at)
            if one_pager.submitted_at:
                dates.append(one_pager.submitted_at)
        last_active = max(dates)
        days_inactive = (now - last_active).days

        # Stage label
        if author.pending_setup:
            stage = 'Pending Setup'
        elif one_pager and one_pager.status == 'submitted':
            stage = 'Submitted for Review'
        elif enrollment and enrollment.status == 'completed':
            stage = 'Proposal Complete'
        elif days_inactive >= 7 and (enrollment or one_pager):
            stage = 'Stalled'
        elif enrollment or one_pager:
            stage = 'In Progress'
        else:
            stage = 'Not Started'

        rows.append({
            'author': author,
            'path': path,
            'current_step': current_step,
            'completed': completed,
            'last_active': last_active,
            'days_inactive': days_inactive,
            'stage': stage,
            'one_pager': one_pager,
            'enrollment': enrollment,
        })

    # Summary stats
    total = len(authors)
    active_7 = sum(1 for r in rows if r['days_inactive'] < 7 and r['stage'] not in ('Pending Setup', 'Not Started'))
    stalled = sum(1 for r in rows if r['stage'] == 'Stalled')
    one_pager_done = sum(1 for r in rows if r['one_pager'] and r['one_pager'].status == 'submitted')
    full_done = sum(1 for r in rows if r['enrollment'] and r['enrollment'].status == 'completed')
    # Conversion: one-pager authors who also enrolled in full program
    one_pager_authors = {r['author'].id for r in rows if r['one_pager']}
    converted = sum(1 for r in rows if r['author'].id in one_pager_authors and r['enrollment'])
    conversion_rate = int((converted / len(one_pager_authors)) * 100) if one_pager_authors else 0

    stats = {
        'total': total,
        'active_7': active_7,
        'stalled': stalled,
        'one_pager_done': one_pager_done,
        'full_done': full_done,
        'conversion_rate': conversion_rate,
    }

    return render_template('admin_pipeline.html', rows=rows, stats=stats)


# ============================================================================
# ADMIN — ONE-PAGER REVIEW
# ============================================================================

@app.route('/admin/one-pager/<int:submission_id>', methods=['GET', 'POST'])
@team_required
def admin_one_pager_detail(submission_id):
    """Admin view of a submitted one-pager."""
    submission = OnePagerSubmission.query.get_or_404(submission_id)
    if request.method == 'POST':
        submission.admin_notes = request.form.get('admin_notes', '').strip()
        db.session.commit()
        flash('Notes saved.', 'success')
    answers = json.loads(submission.answers_json) if submission.answers_json else {}
    return render_template('admin_one_pager_detail.html', submission=submission, answers=answers)


@app.route('/admin/authors/<int:author_id>/one-pager', methods=['GET', 'POST'])
@team_required
def admin_author_one_pager(author_id):
    """Admin fills out / edits an author's one-pager on their behalf (e.g. during a sales call).
    Saves to the same OnePagerSubmission the author sees when they log in."""
    author = Author.query.get_or_404(author_id)
    submission = author.one_pager_submissions.order_by(
        OnePagerSubmission.created_at.desc()).first()

    summary = None
    answers = {}

    if request.method == 'POST':
        action = request.form.get('action', 'save')   # 'save' | 'generate'
        answers = {
            'problem':    request.form.get('problem', '').strip(),
            'reader':     request.form.get('reader', '').strip(),
            'different':  request.form.get('different', '').strip(),
            'why_you':    request.form.get('why_you', '').strip(),
            'marketing':  request.form.get('marketing', '').strip(),
            'book_title': request.form.get('book_title', '').strip(),
        }

        if action == 'generate' and all([answers['problem'], answers['reader'], answers['why_you']]):
            try:
                prompt = f"""You are an expert book proposal coach at Write It Great. An author has answered 5 focused questions about their nonfiction book. Generate a clean, compelling one-page proposal summary.

AUTHOR'S ANSWERS:
Working title: {answers['book_title'] or 'Not yet decided'}
1. What problem does your book solve? {answers['problem']}
2. Who is your target reader? {answers['reader']}
3. Why is your book different from what's already out there? {answers['different'] or 'Not provided'}
4. Why are you the right person to write it? {answers['why_you']}
5. How do you plan to market it? {answers['marketing'] or 'Not provided'}

Write a one-page proposal summary with these sections:
- **The Problem & Promise** (2-3 sentences: the gap this book fills, why it matters now)
- **The Reader** (1-2 sentences: specific target audience — demographic and psychographic)
- **What Makes This Book Different** (2-3 sentences: unique angle, methodology, or perspective)
- **The Author** (2-3 sentences: credentials and unique position to write this)
- **Marketing Potential** (1-2 sentences: platform, reach, and promotional opportunities)
- **Next Steps** (2-3 bullet points: what to develop further before a full submission)

Tone: professional, warm, and specific. Use the author's actual words and voice. Do not pad — be crisp."""
                response = client.chat.completions.create(
                    model='gpt-4o-mini',
                    messages=[{'role': 'user', 'content': prompt}],
                    temperature=0.7,
                    max_tokens=950,
                )
                summary = response.choices[0].message.content.strip()
            except Exception as e:
                print(f'Admin one-pager AI error: {e}')
                flash('AI generation failed — answers were saved.', 'warning')

        # Save / update submission
        if not submission or submission.status == 'submitted':
            submission = OnePagerSubmission(author_id=author.id)
            db.session.add(submission)
        submission.book_title = answers['book_title'] or None
        submission.answers_json = json.dumps(answers)
        if summary:
            submission.summary_text = summary
        db.session.commit()

        if action == 'save':
            flash(f'Saved for {author.name}.', 'success')
        elif summary:
            flash(f'Summary generated and saved for {author.name}.', 'success')
        return redirect(url_for('admin_author_one_pager', author_id=author.id))

    # GET — load existing
    if submission and submission.answers_json:
        answers = json.loads(submission.answers_json)
        summary = submission.summary_text
    return render_template('admin_author_one_pager.html',
                           author=author,
                           submission=submission,
                           answers=answers,
                           summary=summary)


# ============================================================================
# ADMIN — KNOWLEDGE BASE
# ============================================================================

@app.route('/admin/knowledge-base')
@team_required
def admin_knowledge_base():
    """Admin knowledge base — upload training material per module."""
    docs = KnowledgeBaseDocument.query.order_by(
        db.case({None: 9999}, value=KnowledgeBaseDocument.module_order, else_=KnowledgeBaseDocument.module_order).asc(),
        KnowledgeBaseDocument.uploaded_at.desc()
    ).all()
    return render_template('admin_knowledge_base.html', docs=docs, modules=COACHING_MODULES)


@app.route('/admin/knowledge-base/upload', methods=['POST'])
@team_required
def admin_knowledge_base_upload():
    """Upload a knowledge base document."""
    title = request.form.get('title', '').strip()
    module_order = request.form.get('module_order', type=int)
    doc_type = request.form.get('doc_type', 'resource')
    file = request.files.get('doc_file')

    if not title or not file or not file.filename:
        flash('Title and file are required.', 'error')
        return redirect(url_for('admin_knowledge_base'))

    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    if ext not in ('pdf', 'docx', 'txt', 'doc'):
        flash('Only PDF, Word, and TXT files are supported.', 'error')
        return redirect(url_for('admin_knowledge_base'))

    file_bytes = file.read()
    content_text = ''
    try:
        if ext == 'txt':
            content_text = file_bytes.decode('utf-8', errors='ignore')
        elif ext in ('docx', 'doc'):
            from io import BytesIO as _BytesIO
            content_text = '\n'.join(
                p.text for p in Document(_BytesIO(file_bytes)).paragraphs if p.text.strip()
            )
        elif ext == 'pdf':
            content_text = extract_text_from_pdf(BytesIO(file_bytes))
    except Exception as e:
        print(f'KB extract error: {e}')

    doc = KnowledgeBaseDocument(
        title=title,
        filename=filename,
        content_text=content_text[:50000],   # cap at 50k chars
        file_data=file_bytes,
        file_type=ext,
        module_order=module_order or None,
        doc_type=doc_type,
        uploaded_by=current_user.email,
    )
    db.session.add(doc)
    db.session.commit()
    flash(f'"{title}" uploaded successfully.', 'success')
    return redirect(url_for('admin_knowledge_base'))


@app.route('/admin/knowledge-base/<int:doc_id>/delete', methods=['POST'])
@team_required
def admin_knowledge_base_delete(doc_id):
    """Delete a knowledge base document."""
    doc = KnowledgeBaseDocument.query.get_or_404(doc_id)
    db.session.delete(doc)
    db.session.commit()
    flash(f'"{doc.title}" deleted.', 'success')
    return redirect(url_for('admin_knowledge_base'))


@app.route('/admin/knowledge-base/<int:doc_id>/download')
@team_required
def admin_knowledge_base_download(doc_id):
    """Download original KB document file."""
    doc = KnowledgeBaseDocument.query.get_or_404(doc_id)
    return send_file(
        BytesIO(doc.file_data),
        download_name=doc.filename,
        as_attachment=True,
    )


# ============================================================================
# DATABASE MIGRATIONS
# ============================================================================

def run_migrations():
    """Ensure all DB columns and tables exist. Safe to run multiple times.

    Each ALTER TABLE runs in its own transaction with IF NOT EXISTS so a
    single failure cannot block other columns from being added.
    """
    from sqlalchemy import text

    is_pg = 'postgresql' in str(db.engine.url)

    def _add(table, col_def):
        """Add one column to a table — each call is an independent transaction."""
        if is_pg:
            try:
                with db.engine.begin() as conn:
                    conn.execute(text(f'ALTER TABLE {table} ADD COLUMN IF NOT EXISTS {col_def}'))
            except Exception as e:
                print(f'Migration ({table}): {e}')
        else:
            # SQLite < 3.37 has no IF NOT EXISTS on ADD COLUMN — check first
            from sqlalchemy import inspect as _insp
            try:
                existing = [c['name'] for c in _insp(db.engine).get_columns(table)]
                col_name = col_def.split()[0]
                if col_name not in existing:
                    with db.engine.begin() as conn:
                        conn.execute(text(f'ALTER TABLE {table} ADD COLUMN {col_def}'))
            except Exception as e:
                print(f'Migration ({table}): {e}')

    blob = 'BYTEA' if is_pg else 'BLOB'

    # ── proposal ──────────────────────────────────────────────────────────────
    _add('proposal', 'original_filename VARCHAR(500)')
    _add('proposal', f'original_file {blob}')
    _add('proposal', 'content_hash VARCHAR(64)')
    _add('proposal', 'is_archived BOOLEAN DEFAULT FALSE')
    _add('proposal', 'platform_data TEXT')
    _add('proposal', 'marketing_strategy TEXT')
    _add('proposal', 'author_id INTEGER')

    # ── admin_user ─────────────────────────────────────────────────────────────
    _add('admin_user', 'password_reset_token VARCHAR(100)')
    _add('admin_user', 'password_reset_expires TIMESTAMP')
    _add('admin_user', 'totp_secret VARCHAR(64)')
    _add('admin_user', 'totp_enabled BOOLEAN DEFAULT FALSE')
    _add('admin_user', 'failed_login_attempts INTEGER DEFAULT 0')
    _add('admin_user', 'locked_until TIMESTAMP')
    _add('admin_user', f"role VARCHAR(20) DEFAULT '{ROLE_MEMBER}'")
    _add('admin_user', 'is_active_account BOOLEAN DEFAULT TRUE')
    try:
        with db.engine.begin() as conn:
            conn.execute(text(
                f"UPDATE admin_user SET role = '{ROLE_ADMIN}'"
                f" WHERE email = 'anna@writeitgreat.com' AND (role IS NULL OR role != '{ROLE_ADMIN}')"
            ))
    except Exception:
        pass

    # ── publisher ──────────────────────────────────────────────────────────────
    _add('publisher', 'bio TEXT')
    _add('publisher', 'preferred_genres TEXT')
    _add('publisher', 'preferred_topics TEXT')
    _add('publisher', 'website VARCHAR(300)')

    # ── publisher_proposal ─────────────────────────────────────────────────────
    _add('publisher_proposal', "publisher_status VARCHAR(50) DEFAULT 'new'")
    _add('publisher_proposal', 'status_updated_at TIMESTAMP')

    # ── coaching_enrollment ────────────────────────────────────────────────────
    _add('coaching_enrollment', 'book_title VARCHAR(500)')
    _add('coaching_enrollment', 'completed_at TIMESTAMP')
    _add('coaching_enrollment', 'current_module INTEGER DEFAULT 1')
    _add('coaching_enrollment', 'welcome_email_sent BOOLEAN DEFAULT FALSE')
    _add('coaching_enrollment', 'complete_email_sent BOOLEAN DEFAULT FALSE')

    # ── author_module_progress ─────────────────────────────────────────────────
    _add('author_module_progress', 'unlocked_at TIMESTAMP')
    _add('author_module_progress', 'completed_at TIMESTAMP')
    _add('author_module_progress', 'admin_notes TEXT')
    _add('author_module_progress', 'module_unlock_email_sent BOOLEAN DEFAULT FALSE')
    _add('author_module_progress', 'homework_reminder_sent_at TIMESTAMP')

    # ── homework_submission ────────────────────────────────────────────────────
    _add('homework_submission', 'revision_number INTEGER DEFAULT 1')
    _add('homework_submission', 'ai_approved BOOLEAN')
    _add('homework_submission', 'ai_reviewed_at TIMESTAMP')
    _add('homework_submission', 'admin_reviewed_by VARCHAR(200)')
    _add('homework_submission', 'admin_reviewed_at TIMESTAMP')
    _add('homework_submission', 'review_email_sent BOOLEAN DEFAULT FALSE')
    _add('homework_submission', 'ai_feedback TEXT')
    _add('homework_submission', 'admin_feedback TEXT')
    _add('homework_submission', "status VARCHAR(30) DEFAULT 'pending_review'")
    # Back-fill status for rows inserted before this column existed
    try:
        with db.engine.begin() as conn:
            conn.execute(text("UPDATE homework_submission SET status = 'pending_review' WHERE status IS NULL"))
    except Exception:
        pass

    # ── Repair: ensure every active enrollment has a row for each module ───────
    try:
        active_enrollments = CoachingEnrollment.query.filter_by(status='active').all()
        repaired = 0
        for enr in active_enrollments:
            existing_orders = {
                mp.module_order
                for mp in AuthorModuleProgress.query.filter_by(enrollment_id=enr.id).all()
            }
            for m in COACHING_MODULES:
                order = m['order']
                if order not in existing_orders:
                    if order < enr.current_module:
                        row_status, unlocked_at = 'approved', datetime.utcnow()
                    elif order == enr.current_module:
                        row_status, unlocked_at = 'in_progress', datetime.utcnow()
                    else:
                        row_status, unlocked_at = 'locked', None
                    db.session.add(AuthorModuleProgress(
                        enrollment_id=enr.id,
                        module_order=order,
                        status=row_status,
                        unlocked_at=unlocked_at,
                    ))
                    repaired += 1
        if repaired:
            db.session.commit()
            print(f'Migration: created {repaired} missing module progress row(s)')
    except Exception as e:
        db.session.rollback()
        print(f'Migration repair warning: {e}')

    # ── author (new fields) ────────────────────────────────────────────────────
    _add('author', 'pending_setup BOOLEAN DEFAULT FALSE')
    _add('author', 'admin_created BOOLEAN DEFAULT FALSE')
    _add('author', 'assigned_path VARCHAR(30)')
    _add('author', 'last_login_at TIMESTAMP')

    # ── one_pager_submission (new table) ───────────────────────────────────────
    # SQLAlchemy db.create_all() handles new tables; _add is only for existing ones
    try:
        db.create_all()
    except Exception as e:
        print(f'Migration create_all: {e}')


# ============================================================================
# CLI COMMANDS
# ============================================================================

@app.cli.command('init-db')
def init_db():
    """Initialize the database"""
    db.create_all()
    run_migrations()
    print("Database initialized!")


@app.cli.command('create-admin')
def create_admin():
    """Create admin user"""
    import getpass
    
    email = input("Enter admin email [anna@writeitgreat.com]: ").strip() or "anna@writeitgreat.com"
    name = input("Enter admin name [Anna]: ").strip() or "Anna"
    password = getpass.getpass("Enter password: ")
    
    existing = AdminUser.query.filter_by(email=email).first()
    if existing:
        print(f"User {email} already exists. Updating password...")
        existing.set_password(password)
    else:
        user = AdminUser(email=email, name=name)
        user.set_password(password)
        db.session.add(user)
    
    db.session.commit()
    print(f"Admin user '{email}' created/updated successfully!")


# ============================================================================
# FIRST-RUN BOOTSTRAP
# ============================================================================

@app.route('/admin/bootstrap', methods=['GET', 'POST'])
def admin_bootstrap():
    """First-run setup (no admins) OR emergency password reset (admins exist).

    Password-reset mode is activated by setting the ADMIN_RESET_SECRET env var
    in Heroku config vars and visiting /admin/bootstrap?secret=<value>.
    Remove the env var after use to disable this route again.
    """
    reset_enabled = bool(os.environ.get('ADMIN_RESET_SECRET', '').strip())
    has_admins = AdminUser.query.count() > 0

    # Gate: no admins → open for first-run setup
    #       has admins → only accessible when ADMIN_RESET_SECRET env var is set
    if has_admins and not reset_enabled:
        abort(404)

    error = None
    if request.method == 'POST':
        email    = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm', '')

        if not email or not password:
            error = 'Email and password are required.'
        elif password != confirm:
            error = 'Passwords do not match.'
        elif len(password) < 8:
            error = 'Password must be at least 8 characters.'
        else:
            if has_admins:
                # Reset mode: update existing user's password
                user = AdminUser.query.filter_by(email=email).first()
                if not user:
                    error = f'No admin account found for {email}.'
                else:
                    user.set_password(password)
                    # Clear any lockout / 2FA so they can log in fresh
                    user.failed_login_attempts = 0
                    user.locked_until = None
                    user.totp_enabled = False
                    user.totp_secret = None
                    db.session.commit()
                    flash('Password reset. Please log in.', 'success')
                    return redirect(url_for('admin_login'))
            else:
                # First-run mode: create the first admin
                name = request.form.get('name', '').strip()
                if not name:
                    error = 'Name is required.'
                else:
                    user = AdminUser(email=email, name=name, role='admin')
                    user.set_password(password)
                    db.session.add(user)
                    db.session.commit()
                    flash('Admin account created. Please log in.', 'success')
                    return redirect(url_for('admin_login'))

    return render_template('admin_bootstrap.html', error=error, has_admins=has_admins)


# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.errorhandler(404)
def not_found(e):
    return render_template('error.html', error_code=404, error_message="Page not found"), 404


@app.errorhandler(500)
def server_error(e):
    return render_template('error.html', error_code=500, error_message="Server error"), 500


with app.app_context():
    db.create_all()
    try:
        run_migrations()
    except Exception as e:
        print(f"Migration note: {e}")

_start_reengagement_thread()

if __name__ == '__main__':
    app.run(debug=True, port=5000)
